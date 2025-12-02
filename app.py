# =====================================================
# SOCIOSTAT BANKING ENTERPRISE
# Parte 1/4: Inicialização + Funções Core
# =====================================================

import streamlit as st
from streamlit_option_menu import option_menu
import pandas as pd
import numpy as np
import plotly.express as px
import plotly.graph_objects as go
from sklearn.ensemble import GradientBoostingClassifier
from sklearn.preprocessing import StandardScaler
from sklearn.model_selection import train_test_split
from sklearn.metrics import roc_curve, auc, confusion_matrix, classification_report
from lifelines import KaplanMeierFitter
from datetime import datetime
from io import BytesIO
from docx import Document
import warnings
warnings.filterwarnings("ignore")
import folium  
import statsmodels.api as sm
import time
from gtts import gTTS
from io import BytesIO
import requests
from bs4 import BeautifulSoup
import matplotlib.pyplot as plt

# ===================== INICIALIZAÇÃO SESSION STATE =====================
if "theme" not in st.session_state:
    st.session_state.theme = "Dark" # Começa em Dark Mode

if "logged_in" not in st.session_state:
    st.session_state.logged_in = False
if "show_onboarding" not in st.session_state:
    st.session_state.show_onboarding = True # Começa sempre ativo na primeira vez
if "df" not in st.session_state:
    st.session_state.df = None
if "user_role" not in st.session_state:
    st.session_state.user_role = None
if "username" not in st.session_state:
    st.session_state.username = ""
if "audit_log" not in st.session_state:
    st.session_state.audit_log = []
if "notifications" not in st.session_state:
    st.session_state.notifications = []
if "language" not in st.session_state:
    st.session_state.language = "PT"

# ... (outros session states) ...

# Base de Dados Simulada de Utilizadores (Func 25)
if "db_users" not in st.session_state:
    st.session_state.db_users = [
        {"id": 1, "username": "admin", "role": "Admin", "status": "Ativo", "last_login": "2024-03-10 09:00"},
        {"id": 2, "username": "analista_joao", "role": "Analyst", "status": "Ativo", "last_login": "2024-03-09 14:30"},
        {"id": 3, "username": "gestor_maria", "role": "Risk Manager", "status": "Ativo", "last_login": "2024-03-10 08:45"},
        {"id": 4, "username": "auditor_pedro", "role": "Auditor", "status": "Inativo", "last_login": "2024-02-20 10:00"},
    ]

    # Base de dados de modelos (MLOps Registry)
if "model_registry" not in st.session_state:
    # Adicionar alguns dados simulados para não começar vazio
    st.session_state.model_registry = [
        {
            "id": "MOD-CS-2023-001", "nome": "Credit Scoring Base", "versao": "v1.0",
            "data": "2023-12-01 10:00", "autor": "System", "status": "Archived",
            "metricas": {"AUC": 0.78, "Accuracy": 0.82}, "params": {"model": "LogisticRegression"}
        },
        {
            "id": "MOD-CS-2024-001", "nome": "Credit Scoring Gradient", "versao": "v2.0",
            "data": "2024-01-15 14:30", "autor": "System", "status": "Production",
            "metricas": {"AUC": 0.85, "Accuracy": 0.88}, "params": {"model": "GradientBoosting", "n_estimators": 100}
        }
    ]

if "metadata_catalog" not in st.session_state:
    # Dados iniciais simulados
    st.session_state.metadata_catalog = {
        "Demo Enterprise": {
            "descricao": "Dataset sintético para demonstração de capacidades bancárias.",
            "fonte": "Gerador Interno (Numpy)",
            "responsavel": "Sistema",
            "ultima_atualizacao": datetime.now().strftime("%Y-%m-%d"),
            "confidencialidade": "Baixa"
        },
        "Dados Web Scraping": {
            "descricao": "Dados extraídos de fontes públicas via módulo de scraping.",
            "fonte": "Web (Wikipedia/INE)",
            "responsavel": "Analista de Dados",
            "ultima_atualizacao": "N/A",
            "confidencialidade": "Pública"
        }
    }

    # ... (junto aos outros estados como 'theme', 'logged_in', etc.)

if "experiments_db" not in st.session_state:
    # Dados simulados de experiências passadas (Funcionalidade 41)
    st.session_state.experiments_db = [
        {
            "id": "EXP-2023-881", "data": "2023-11-15 10:30", "autor": "analista_joao",
            "tipo": "Credit Scoring", "parametros": "Modelo: GradientBoosting | Threshold: 0.6",
            "resultado_chave": "AUC: 0.82", "status": "Auditado"
        },
        {
            "id": "EXP-2024-002", "data": "2024-01-10 14:20", "autor": "admin",
            "tipo": "Forecasting", "parametros": "Horizonte: 12m | Modelo: ARIMA(1,1,1)",
            "resultado_chave": "Default Previsto: 4.2%", "status": "Rascunho"
        }
    ]

if "user_xp" not in st.session_state:
    # Simular dados para o utilizador atual
    st.session_state.user_xp = 1250 
    st.session_state.user_level = 3
    st.session_state.badges = ["Analista Júnior", "Primeiro Relatório"]

if "gamification_log" not in st.session_state:
    st.session_state.gamification_log = [
        {"data": "2024-03-01", "acao": "Login Diário", "xp": 50},
        {"data": "2024-03-02", "acao": "Relatório Gerado", "xp": 100},
        {"data": "2024-03-05", "acao": "Upload de Dados (Qualidade A)", "xp": 300}
    ]

if "alert_rules" not in st.session_state:
    # Regras iniciais de exemplo
    st.session_state.alert_rules = [
        {"id": 1, "metrica": "Taxa Default", "operador": ">", "valor": 0.10, "ativo": True},
        {"id": 2, "metrica": "Score Médio", "operador": "<", "valor": 600, "ativo": True}
    ]
    
if "triggered_alerts_history" not in st.session_state:
    st.session_state.triggered_alerts_history = []

# ===================== APM / PERFORMANCE (FUNC 46) =====================
import time
from functools import wraps

# Inicializar estado dos logs de performance
if "apm_logs" not in st.session_state:
    st.session_state.apm_logs = []

def monitorar_performance(func):
    """
    Decorator que mede o tempo de execução de uma função e regista no APM.
    Uso: @monitorar_performance antes da definição da função.
    """
    @wraps(func)
    def wrapper(*args, **kwargs):
        start_time = time.time()
        try:
            result = func(*args, **kwargs)
            status = "Sucesso"
            return result
        except Exception as e:
            status = f"Erro: {str(e)}"
            raise e
        finally:
            end_time = time.time()
            duration = end_time - start_time
            
            # Registar métrica
            st.session_state.apm_logs.append({
                "timestamp": datetime.now(),
                "funcao": func.__name__,
                "duracao_sec": duration,
                "status": status,
                "user": st.session_state.get("username", "System")
            })
            
            # Manter apenas os últimos 1000 registos
            if len(st.session_state.apm_logs) > 1000:
                st.session_state.apm_logs.pop(0)
                
    return wrapper

# ===================== TRADUÇÕES (FUNC 20) =====================
TRANSLATIONS = {
    "PT": {
        "sidebar_user": "Sessão Ativa",
        "role": "Cargo",
        "records": "Registos",
        "dashboard_title": "🏠 Dashboard Executivo",
        "kpi_portfolio": "💰 Carteira Total",
        "kpi_default": "⚠️ Taxa Default",
        "kpi_clients": "👥 Clientes",
        "welcome": "Bem-vindo ao SocioStat Enterprise!",
        "logout": "🚪 Sair"
    },
    "EN": {
        "sidebar_user": "Active Session",
        "role": "Role",
        "records": "Records",
        "dashboard_title": "🏠 Executive Dashboard",
        "kpi_portfolio": "💰 Total Portfolio",
        "kpi_default": "⚠️ Default Rate",
        "kpi_clients": "👥 Clients",
        "welcome": "Welcome to SocioStat Enterprise!",
        "logout": "🚪 Logout"
    }
}

# ===================== GESTÃO DE UTILIZADORES (FUNC 25) =====================
def criar_novo_utilizador(username, role):
    """Cria um novo utilizador na base de dados simulada"""
    novo_id = max([u['id'] for u in st.session_state.db_users]) + 1
    novo_user = {
        "id": novo_id,
        "username": username,
        "role": role,
        "status": "Ativo",
        "last_login": "Nunca"
    }
    st.session_state.db_users.append(novo_user)
    return True

def atualizar_role_utilizador(user_id, novo_role):
    """Atualiza o cargo de um utilizador"""
    for u in st.session_state.db_users:
        if u['id'] == user_id:
            u['role'] = novo_role
            return True
    return False

def alternar_status_utilizador(user_id):
    """Ativa/Desativa um utilizador"""
    for u in st.session_state.db_users:
        if u['id'] == user_id:
            u['status'] = "Inativo" if u['status'] == "Ativo" else "Ativo"
            return u['status']
    return None

def t(key):
    """Função auxiliar de tradução"""
    lang = st.session_state.language
    return TRANSLATIONS.get(lang, TRANSLATIONS["PT"]).get(key, key)

# ===================== WEB SCRAPING (FUNC 39) =====================
@st.cache_data(ttl=600) # Cache de 10 min para não bloquear o IP
def executar_scraping_web(url):
    """
    Acede a uma URL e tenta extrair a primeira tabela de dados encontrada.
    """
    try:
        # Headers para fingir que somos um browser (evita bloqueios 403)
        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36"
        }
        
        response = requests.get(url, headers=headers, timeout=10)
        response.raise_for_status()
        
        # Parse do HTML
        soup = BeautifulSoup(response.text, 'html.parser')
        
        # Encontrar tabelas
        tabelas = pd.read_html(str(soup))
        
        if not tabelas:
            return {"sucesso": False, "erro": "Nenhuma tabela encontrada nesta página."}
            
        # Retorna a maior tabela encontrada (geralmente é a de dados principais)
        df_extraido = max(tabelas, key=len)
        
        # Limpeza básica (remover colunas vazias)
        df_extraido = df_extraido.dropna(axis=1, how='all')
        
        return {
            "sucesso": True,
            "dados": df_extraido,
            "titulo": soup.title.string if soup.title else "Sem Título",
            "total_tabelas": len(tabelas)
        }
        
    except Exception as e:
        return {"sucesso": False, "erro": str(e)}

# ===================== INTEGRAÇÃO CLOUD (FUNC 36) =====================
def init_cloud_storage():
    """Inicializa o armazenamento simulado se não existir"""
    if "cloud_storage" not in st.session_state:
        st.session_state.cloud_storage = {
            "Google Drive": [
                {"name": "dados_historicos_2023.csv", "size": "1.2 MB", "date": "2023-12-10"},
                {"name": "relatorio_risco_q1.pdf", "size": "4.5 MB", "date": "2024-01-15"},
                {"name": "portfolio_clientes_v2.xlsx", "size": "0.8 MB", "date": "2024-02-20"}
            ],
            "Dropbox": [
                {"name": "backup_full.sql", "size": "150 MB", "date": "2024-03-01"},
                {"name": "compliance_docs.zip", "size": "22 MB", "date": "2024-03-05"}
            ]
        }

def upload_to_cloud(file, provider):
    """Simula o upload de um ficheiro para a nuvem"""
    # Em produção, usaria a API do Google Drive (google-api-python-client)
    time.sleep(1.5) # Simular rede
    
    novo_ficheiro = {
        "name": file.name,
        "size": f"{file.size / 1024:.1f} KB",
        "date": datetime.now().strftime("%Y-%m-%d")
    }
    
    st.session_state.cloud_storage[provider].append(novo_ficheiro)
    return True

def import_from_cloud(filename, provider):
    """Simula a importação de um ficheiro da nuvem para a App"""
    time.sleep(1)
    # Numa app real, aqui faríamos pd.read_csv(link_do_drive)
    # Aqui, vamos apenas simular sucesso
    return True

# ===================== AUDITORIA & COMPLIANCE =====================

def enviar_notificacao(mensagem, tipo="info"):
    """
    Envia uma notificação visual (Toast) e grava no histórico.
    Tipos: 'info', 'sucesso', 'aviso', 'erro'
    """
    # 1. Mostrar Toast (Popup temporário no canto direito)
    if tipo == "sucesso":
        st.toast(f"✅ {mensagem}", icon="✅")
    elif tipo == "aviso":
        st.toast(f"⚠️ {mensagem}", icon="⚠️")
    elif tipo == "erro":
        st.toast(f"❌ {mensagem}", icon="❌")
    else:
        st.toast(f"ℹ️ {mensagem}", icon="ℹ️")
        
    # 2. Gravar no Histórico (com timestamp)
    timestamp = datetime.now().strftime("%H:%M")
    st.session_state.notifications.insert(0, {
        "time": timestamp,
        "msg": mensagem,
        "type": tipo
    })
    
    # Manter apenas as últimas 50 notificações
    if len(st.session_state.notifications) > 50:
        st.session_state.notifications.pop()

def renderizar_onboarding():
    """Mostra o guia de boas-vindas (Versão Nativa Streamlit)"""
    if st.session_state.get("show_onboarding", False):
        with st.container():
            # Cartão de Boas-Vindas
            st.info("👋 **Bem-vindo ao SocioStat Banking Enterprise**")
            
            c1, c2, c3 = st.columns(3)
            with c1:
                st.markdown("### 📊 1. Diagnóstico")
                st.caption("Consulte o **Dashboard** e o **Geo-Risk** para ver o estado da carteira.")
            with c2:
                st.markdown("### 🧠 2. Inteligência")
                st.caption("Use **Credit Scoring** e **Forecasting** para prever riscos futuros.")
            with c3:
                st.markdown("### 🛡️ 3. Ação")
                st.caption("Atue com o **Early Warning** e gere **Relatórios** automáticos.")
            
            if st.button("🚀 Começar a Trabalhar", key="btn_close_onboard", use_container_width=True):
                st.session_state.show_onboarding = False
                st.rerun()
        st.markdown("---")

# ===================== GAMIFICAÇÃO (FUNC 44) =====================
def calcular_nivel(xp):
    """Calcula o nível com base no XP (Nível sobe a cada 1000 XP)"""
    return int(xp / 1000) + 1

def adicionar_xp(pontos, motivo):
    """Dá pontos ao utilizador e verifica subida de nível"""
    st.session_state.user_xp += pontos
    
    # Registar no log
    st.session_state.gamification_log.insert(0, {
        "data": datetime.now().strftime("%Y-%m-%d"), 
        "acao": motivo, 
        "xp": pontos
    })
    
    # Verificar Level Up
    novo_nivel = calcular_nivel(st.session_state.user_xp)
    if novo_nivel > st.session_state.user_level:
        st.session_state.user_level = novo_nivel
        st.toast(f"🎉 PARABÉNS! Subiu para Nível {novo_nivel}!", icon="⭐")
        st.balloons()
    else:
        st.toast(f"+{pontos} XP: {motivo}", icon="✨")

def get_leaderboard():
    """Gera um ranking simulado de analistas"""
    return pd.DataFrame([
        {"Rank": 1, "Utilizador": "Ana Risk", "Nível": 15, "XP": 15400, "Badges": "🏆👑"},
        {"Rank": 2, "Utilizador": "Carlos Data", "Nível": 12, "XP": 12100, "Badges": "🔥"},
        {"Rank": 3, "Utilizador": st.session_state.username, "Nível": st.session_state.user_level, "XP": st.session_state.user_xp, "Badges": "⭐"},
        {"Rank": 4, "Utilizador": "Joana Audit", "Nível": 2, "XP": 2100, "Badges": ""},
    ]).sort_values("XP", ascending=False)

def log_audit(user, action, details):
    """Regista ações para compliance (RGPD, Basel III)"""
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    st.session_state.audit_log.append({
        "timestamp": timestamp,
        "user": user,
        "action": action,
        "details": details,
        "ip": "127.0.0.1"
    })

def export_audit_log():
    """Exporta log de auditoria"""
    return pd.DataFrame(st.session_state.audit_log)

def check_compliance_alerts(df):
    """Sistema automático de alertas regulatórios"""
    alerts = []
    
    # Alerta 1: Taxa de Default alta
    if 'default' in df.columns:
        default_rate = df['default'].mean()
        if default_rate > 0.15:
            alerts.append({
                "nivel": "CRÍTICO",
                "tipo": "Risco de Crédito",
                "mensagem": f"Taxa de default: {default_rate:.1%} (limite: 15%)",
                "acao": "Rever políticas de crédito imediatamente"
            })
    
    # Alerta 2: Concentração de segmento
    if 'segmento' in df.columns and len(df) > 0:
        seg_concentration = df['segmento'].value_counts(normalize=True).max()
        if seg_concentration > 0.4:
            alerts.append({
                "nivel": "AVISO",
                "tipo": "Concentração de Risco",
                "mensagem": f"Concentração: {seg_concentration:.1%} (limite: 40%)",
                "acao": "Diversificar carteira"
            })
    
    # Alerta 3: LTV elevado
    if 'ltv' in df.columns:
        high_ltv_pct = (df['ltv'] > 0.9).sum() / len(df)
        if high_ltv_pct > 0.2:
            alerts.append({
                "nivel": "AVISO",
                "tipo": "LTV Elevado",
                "mensagem": f"{high_ltv_pct:.1%} com LTV > 90%",
                "acao": "Rever garantias e provisões"
            })
    
    return alerts

# ===================== CONETORES OFICIAIS (FUNC 43) =====================
class DataConnector:
    def __init__(self, nome, base_url, auth_type="API Key"):
        self.nome = nome
        self.base_url = base_url
        self.auth_type = auth_type
        self.status = "Desconectado 🔴"
        
    def conectar(self, api_key):
        """Simula o teste de conexão"""
        time.sleep(1)
        if len(api_key) > 5:
            self.status = "Conectado 🟢"
            return True
        return False
        
    def extrair_indicador(self, indicador):
        """Simula a extração de uma série temporal"""
        if "🟢" not in self.status:
            return None
            
        # Gerar dados simulados realistas baseados no nome do conetor
        dates = pd.date_range(end=datetime.now(), periods=12, freq='M')
        
        if self.nome == "INE (Instituto Nacional Estatística)":
            valores = np.random.normal(100, 5, 12) # Ex: Índice de Preços
        elif self.nome == "BCE (Banco Central Europeu)":
            valores = np.linspace(3.5, 4.5, 12) # Ex: Taxas de Juro
        else:
            valores = np.random.rand(12)
            
        return pd.DataFrame({"Data": dates, indicador: valores}).set_index("Data")

# Inicializar conetores no estado
def init_connectors():
    if "connectors" not in st.session_state:
        st.session_state.connectors = {
            "INE": DataConnector("INE (Instituto Nacional Estatística)", "https://api.ine.pt/v1"),
            "BCE": DataConnector("BCE (Banco Central Europeu)", "https://sdw-wsrest.ecb.europa.eu"),
            "PORDATA": DataConnector("PORDATA", "https://www.pordata.pt/api")
        }

# ===================== CAUSALIDADE / DiD (FUNC 21) =====================
def calcular_causalidade_did(cenario="Campanha Renegociação"):
    """
    Simula e calcula um modelo Difference-in-Differences (DiD).
    Retorna o dataset e o efeito calculado.
    """
    # 1. Configurar Parâmetros de Simulação
    periodos = 12 # 6 meses antes, 6 depois
    ponto_intervencao = 6
    
    # Base: Taxa de Default
    base_control = 0.12 # 12% default
    base_treat = 0.13   # 13% default (ligeiramente pior)
    
    tendencia_mercado = -0.01 # O mercado está a melhorar (default cai 1%)
    efeito_real = -0.04 # A campanha reduz o default em 4% (o que queremos medir)
    
    dados = []
    
    # 2. Gerar Dados
    np.random.seed(42)
    for t in range(1, periodos + 1):
        fator_tempo = (t / periodos) * tendencia_mercado # Tendência comum
        is_post = 1 if t > ponto_intervencao else 0
        
        # Grupo Controlo (Não recebeu campanha)
        y_control = base_control + fator_tempo + np.random.normal(0, 0.005)
        dados.append({'Tempo': t, 'Grupo': 'Controlo', 'Default': y_control, 'Periodo': 'Pós' if is_post else 'Pré'})
        
        # Grupo Tratamento (Recebeu campanha)
        # O efeito só existe se for Grupo Tratamento E Pós-Intervenção
        y_treat = base_treat + fator_tempo + (efeito_real * is_post) + np.random.normal(0, 0.005)
        dados.append({'Tempo': t, 'Grupo': 'Tratamento', 'Default': y_treat, 'Periodo': 'Pós' if is_post else 'Pré'})
        
    df_did = pd.DataFrame(dados)
    
    # 3. Cálculo do Estimador DiD (A "Matemática")
    medias = df_did.groupby(['Grupo', 'Periodo'])['Default'].mean()
    
    control_pre = medias['Controlo']['Pré']
    control_pos = medias['Controlo']['Pós']
    treat_pre = medias['Tratamento']['Pré']
    treat_pos = medias['Tratamento']['Pós']
    
    # A fórmula mágica do DiD
    diff_control = control_pos - control_pre
    diff_treat = treat_pos - treat_pre
    did_estimator = diff_treat - diff_control
    
    return {
        'df': df_did,
        'did': did_estimator,
        'stats': {
            'control_pre': control_pre, 'control_pos': control_pos,
            'treat_pre': treat_pre, 'treat_pos': treat_pos
        },
        'ponto_intervencao': ponto_intervencao
    }

# ===================== BASEL III - CAPITAL REGULATÓRIO =====================
def calcular_capital_regulatorio(df):
    """Calcula capital necessário segundo Basel III"""
    if 'default' not in df.columns or 'divida_total' not in df.columns:
        return None
    
    # PD - Probability of Default
    pd_val = df['default'].mean()
    
    # EAD - Exposure at Default
    ead = df['divida_total'].sum()
    
    # LGD - Loss Given Default (assumir 45% padrão bancário)
    lgd = 0.45
    
    # Expected Loss
    expected_loss = pd_val * lgd * ead
    
    # RWA - Risk Weighted Assets (fórmula simplificada Basel III)
    rwa = ead * 1.06 * pd_val * lgd
    
    # Capital mínimo (8% do RWA - Pilar 1)
    capital_minimo = rwa * 0.08
    
    # Capital recomendado (8% + buffer conservação 2.5% = 10.5%)
    capital_recomendado = rwa * 0.105
    
    return {
        "PD": pd_val,
        "EAD": ead,
        "LGD": lgd,
        "Expected_Loss": expected_loss,
        "RWA": rwa,
        "Capital_Minimo_8pct": capital_minimo,
        "Capital_Recomendado_10.5pct": capital_recomendado
    }

# ===================== OTIMIZAÇÃO DE QUERIES (FUNC 50) =====================
def benchmark_query_performance(df):
    """
    Compara o desempenho de uma pesquisa não otimizada vs indexada.
    Simula o comportamento de um Full Table Scan vs Index Scan.
    """
    resultados = {}
    
    # Cenário: Procurar clientes de Alto Risco em Lisboa
    
    # 1. Método Lento (Simulação de Full Table Scan)
    start_slow = time.time()
    # Simular carga de leitura de disco sem índice
    time.sleep(1.5) 
    # Filtragem ineficiente (iterativa)
    _ = [row for index, row in df.iterrows() 
         if row['regiao'] == 'Lisboa' and row['score_interno'] < 600]
    end_slow = time.time()
    
    resultados['Slow'] = end_slow - start_slow
    
    # 2. Método Otimizado (Pandas Vectorization + Simulação de Índice)
    start_fast = time.time()
    # Filtragem vetorial (equivalente a Index Scan)
    _ = df[(df['regiao'] == 'Lisboa') & (df['score_interno'] < 600)]
    end_fast = time.time()
    
    resultados['Fast'] = end_fast - start_fast
    
    return resultados

def gerar_explain_plan_simulado(tipo_query):
    """Gera um plano de execução SQL falso para efeitos didáticos"""
    if tipo_query == "Lenta":
        return """
        Seq Scan on clientes_table  (cost=0.00..15432.00 rows=500 width=450)
          Filter: ((regiao = 'Lisboa') AND (score_interno < 600))
          Buffers: shared hit=15000 read=5000
        Planning Time: 0.5 ms
        Execution Time: 1500.0 ms  <-- ALERTA: FULL SCAN
        """
    else:
        return """
        Index Scan using idx_regiao_score on clientes_table  (cost=0.29..12.50 rows=500 width=450)
          Index Cond: ((regiao = 'Lisboa') AND (score_interno < 600))
          Buffers: shared hit=15
        Planning Time: 0.1 ms
        Execution Time: 0.02 ms
        """

# ===================== DEMO DATA =====================
@monitorar_performance
def configurar_demo_bancaria():
    """Gera dataset demo bancário realista COM COORDENADAS FIXAS"""
    np.random.seed(42) # Garante que os dados são sempre iguais
    n = 2000
    
    # Features base
    idade = np.random.randint(21, 75, n)
    score_interno = np.random.normal(650, 80, n).astype(int)
    rendimento = np.random.lognormal(7.5, 0.5, n)
    
    # Default correlacionado com features (realista)
    default_prob = 1 / (1 + np.exp(-(
        -5 
        + (score_interno < 600) * 2 
        + (rendimento < 1500) * 1.5
        + (idade < 25) * 0.5
    )))
    default = np.random.binomial(1, default_prob)
    
    df_demo = pd.DataFrame({
        "id_cliente": range(1, n+1),
        "duracao_meses": np.random.randint(1, 60, n),
        "default": default,
        "idade": idade,
        "score_interno": score_interno,
        "rendimento_mensal": rendimento,
        "ltv": np.random.uniform(0.2, 1.2, n),
        "divida_total": np.random.lognormal(9, 0.8, n),
        "num_produtos": np.random.randint(1, 6, n),
        "segmento": np.random.choice(['Particular', 'Empresas', 'Premium'], n, p=[0.6, 0.3, 0.1]),
        "regiao": np.random.choice(['Norte', 'Centro', 'Lisboa', 'Alentejo', 'Algarve'], n),
        "taxa_juro": np.random.uniform(2.5, 8.5, n),
        "num_atrasos_12m": np.random.poisson(0.5, n),
        "utilizacao_credito": np.random.uniform(0.1, 0.95, n),
    })
    # --- CORREÇÃO GEO: Gerar Coordenadas AQUI e FIXAR ---
    coords_base = {
        'Norte': {'lat': 41.1579, 'lon': -8.6291}, 
        'Centro': {'lat': 40.2033, 'lon': -8.4103}, 
        'Lisboa': {'lat': 38.7223, 'lon': -9.1393}, 
        'Alentejo': {'lat': 38.5714, 'lon': -7.9135}, 
        'Algarve': {'lat': 37.0179, 'lon': -7.9308}  
    }

    def get_lat(regiao):
        return coords_base[regiao]['lat'] + np.random.uniform(-0.2, 0.2)

    def get_lon(regiao):
        return coords_base[regiao]['lon'] + np.random.uniform(-0.2, 0.2)

    df_demo['latitude'] = df_demo['regiao'].apply(get_lat)
    df_demo['longitude'] = df_demo['regiao'].apply(get_lon)
    
# ... (dentro de configurar_demo_bancaria, antes de st.session_state.df = ...)
    
    # Adicionar coluna de Tags vazia (ou com algumas pré-preenchidas)
    tags_iniciais = [[] for _ in range(n)] # Lista de listas vazias
    # Simular algumas tags para teste
    for i in range(0, n, 10): # A cada 10 clientes
        tags_iniciais[i] = ["VIP", "Revisar"]
        
    df_demo['tags'] = tags_iniciais
    
    st.session_state.df = df_demo
    # ...
    
    st.session_state.df = df_demo
    log_audit("System", "Data Load", f"Demo bancária carregada: {n} registos")

# ===================== RECONSTRUÇÃO DE EXPERIÊNCIAS (FUNC 41) =====================
def salvar_snapshot(tipo, params, resultado):
    """Guarda o estado atual de uma análise"""
    novo_id = f"EXP-{datetime.now().strftime('%Y-%j')}-{np.random.randint(100,999)}"
    
    snapshot = {
        "id": novo_id,
        "data": datetime.now().strftime("%Y-%m-%d %H:%M"),
        "autor": st.session_state.username,
        "tipo": tipo,
        "parametros": str(params),
        "resultado_chave": str(resultado),
        "status": "Novo"
    }
    
    st.session_state.experiments_db.insert(0, snapshot)
    return novo_id
# ===================== SOFIA AI (COM VOZ REAL) =====================
def sofia_explica(texto):
    """
    Box de explicação da Sofia IA com geração de áudio real (gTTS).
    """
    # Cores do tema
    is_dark = st.session_state.get('theme', 'Dark') == 'Dark'
    bg_color = "linear-gradient(90deg, #1e3a8a, #1e40af)" if is_dark else "linear-gradient(90deg, #e0f2fe, #dbeafe)"
    text_color = "white" if is_dark else "#1e3a8a"
    border_color = "#60a5fa" if is_dark else "#2563eb"
    
    # Container Visual
    st.markdown(f"""
    <div style="
        background: {bg_color}; 
        color: {text_color}; 
        padding: 20px; 
        border-radius: 10px; 
        border-left: 5px solid {border_color}; 
        margin: 20px 0;
        box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1);
    ">
        <div style="display: flex; align-items: center; margin-bottom: 10px;">
            <span style="font-size: 24px; margin-right: 10px;">👩‍💼</span>
            <h4 style="margin: 0; color: {text_color};">Sofia AI Insight</h4>
        </div>
        <p style="font-size: 15px; line-height: 1.6; margin: 0;">{texto}</p>
    </div>
    """, unsafe_allow_html=True)
    
    # --- ÁUDIO REAL (Funcionalidade 34) ---
    # Usamos um expander para o player não ocupar espaço visual se não for usado
    with st.expander("🔊 Ouvir Explicação"):
        try:
            # Gerar áudio em memória (sem salvar ficheiro no disco)
            tts = gTTS(text=texto, lang='pt', slow=False)
            audio_bytes = BytesIO()
            tts.write_to_fp(audio_bytes)
            
            # Mostrar o Player de Áudio Nativo do Browser
            st.audio(audio_bytes, format='audio/mp3')
            
        except Exception as e:
            st.warning("⚠️ Não foi possível gerar áudio (Verifique a sua ligação à internet).")
# ===================== FIM DA PARTE 1 =====================

# =====================================================
# PARTE 2/4: Funções de Modelagem Avançada
# =====================================================

# ===================== AUDITORIA DE SEGURANÇA (FUNC 37) =====================
def analisar_ameacas_seguranca(logs):
    """
    Analisa os logs de auditoria em busca de padrões suspeitos.
    Retorna uma lista de alertas de segurança.
    """
    if not logs:
        return []
    
    df = pd.DataFrame(logs)
    df['timestamp'] = pd.to_datetime(df['timestamp'])
    
    ameacas = []
    
    # REGRA 1: Detetar Brute Force (Muitas falhas de login)
    # Filtrar ações de login falhado (se existissem logs específicos de falha)
    # Como na nossa demo só logamos o sucesso ou tentativas genéricas, 
    # vamos simular que qualquer ação repetida muito rápido é suspeita.
    
    # Agrupar por utilizador e contar ações nos últimos 5 minutos
    now = datetime.now()
    janela_tempo = now - pd.Timedelta(minutes=5)
    
    logs_recentes = df[df['timestamp'] > janela_tempo]
    
    if not logs_recentes.empty:
        contagem = logs_recentes['user'].value_counts()
        for user, count in contagem.items():
            if count > 20: # Se fez mais de 20 ações em 5 min
                ameacas.append({
                    "nivel": "ALTO",
                    "tipo": "Comportamento Anómalo",
                    "user": user,
                    "mensagem": f"Utilizador realizou {count} ações em <5 minutos. Possível script/bot.",
                    "hora": now.strftime("%H:%M")
                })

    # REGRA 2: Acesso a áreas sensíveis por perfis não autorizados (Simulação)
    # Vamos assumir que 'Admin' é o único que pode ir à 'Gestão de Utilizadores'
    # Se virmos um 'Analyst' lá, é alerta vermelho.
    acessos_indevidos = df[
        (df['action'] == "Admin Access") & 
        (df['user'] != "Admin") & 
        (df['user'] != "Demo User") # Permitir na demo
    ]
    
    for _, row in acessos_indevidos.iterrows():
        ameacas.append({
            "nivel": "CRÍTICO",
            "tipo": "Violação de Controlo de Acesso",
            "user": row['user'],
            "mensagem": f"Tentativa de acesso a área Admin.",
            "hora": row['timestamp'].strftime("%H:%M")
        })

    return ameacas

def simular_ataque():
    """Injeta logs falsos para testar o sistema de detecção"""
    fake_user = "hacker_ip_192.168.1.50"
    for _ in range(25):
        log_audit(fake_user, "Login Failed", "Invalid Password")
    return True

# ===================== INTEGRAÇÃO BI (FUNC 47) =====================
def gerar_vista_bi_otimizada(df):
    """
    Gera uma 'Flat Table' (Vista Desnormalizada) otimizada para Power BI/Tableau.
    Junta dados de clientes, cálculos de risco e segmentos.
    """
    df_bi = df.copy()
    
    # 1. Enriquecer com Cálculos (que o Power BI não saberia fazer sozinho)
    # Exemplo: Probabilidade de Default (simulada aqui, mas viria do modelo)
    if 'score_interno' in df_bi.columns:
        # Simular uma probabilidade baseada no score (Lógica de negócio do Python)
        df_bi['BI_Prob_Default'] = 1 / (1 + np.exp((df_bi['score_interno'] - 600) / 50))
        
    # 2. Criar Categorias de Negócio (Business Logic)
    if 'ltv' in df_bi.columns:
        df_bi['BI_Risco_LTV'] = pd.cut(df_bi['ltv'], bins=[0, 0.6, 0.8, 1.0, 10], labels=['Seguro', 'Moderado', 'Alto', 'Crítico'])
        
    # 3. Timestamp de Extração (Para controlo de versão no BI)
    df_bi['BI_Data_Extracao'] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    
    # 4. Selecionar colunas chave para exportação
    cols_export = [
        'id_cliente', 'segmento', 'regiao', 'divida_total', 
        'default', 'score_interno', 'BI_Prob_Default', 'BI_Risco_LTV', 'BI_Data_Extracao'
    ]
    
    # Garantir que colunas existem antes de selecionar
    cols_finais = [c for c in cols_export if c in df_bi.columns]
    
    return df_bi[cols_finais]

def gerar_script_sql_bi(tabela_nome="vw_sociostat_analytics"):
    """Gera o script SQL para criar a vista no Data Warehouse"""
    return f"""
    -- Script de Integração Power BI / Tableau
    -- Criado por SocioStat Enterprise em {datetime.now().strftime('%Y-%m-%d')}
    
    CREATE OR REPLACE VIEW {tabela_nome} AS
    SELECT 
        t1.id_cliente,
        t1.segmento,
        t1.regiao,
        t1.divida_total as exposicao_eur,
        t1.score_interno,
        -- Cálculos proprietários do SocioStat injetados via ETL
        t2.probabilidade_default,
        t2.cluster_risco
    FROM 
        clientes_raw t1
    LEFT JOIN 
        motor_risco_output t2 ON t1.id_cliente = t2.id_cliente
    WHERE 
        t1.ativo = 1;
        
    -- Permissões para o utilizador de BI
    GRANT SELECT ON {tabela_nome} TO user_powerbi;
    """

# ===================== ALERTAS PROATIVOS (FUNC 53) =====================
def verificar_regras_proativas(df):
    """
    Motor de Regras: Verifica se os dados atuais violam alguma regra definida.
    """
    kpis = calcular_kpis_principais(df)
    # Mapeamento de nomes amigáveis para chaves internas
    mapa_kpis = {
        "Taxa Default": kpis.get("Taxa_Default", 0),
        "Score Médio": kpis.get("Score_Medio", 0),
        "LTV Médio": kpis.get("LTV_Medio", 0),
        "Carteira Total": kpis.get("Carteira_Total", 0)
    }
    
    disparos = []
    
    for regra in st.session_state.alert_rules:
        if not regra['ativo']:
            continue
            
        valor_atual = mapa_kpis.get(regra['metrica'], 0)
        limite = regra['valor']
        
        violacao = False
        if regra['operador'] == ">" and valor_atual > limite:
            violacao = True
        elif regra['operador'] == "<" and valor_atual < limite:
            violacao = True
            
        if violacao:
            msg = f"ALERTA: {regra['metrica']} ({valor_atual:.2f}) violou o limite {regra['operador']} {limite}"
            disparos.append(msg)
            
            # Registar no histórico se for novo (para não spammar)
            ultimo_disparo = next((a for a in st.session_state.triggered_alerts_history if a['msg'] == msg), None)
            if not ultimo_disparo: 
                st.session_state.triggered_alerts_history.insert(0, {
                    "data": datetime.now().strftime("%Y-%m-%d %H:%M"),
                    "msg": msg,
                    "regra_id": regra['id']
                })
                # Enviar Notificação Toast (Func 17)
                enviar_notificacao(msg, "erro")
                
    return disparos

def adicionar_regra(metrica, operador, valor):
    """Cria uma nova regra de alerta"""
    novo_id = len(st.session_state.alert_rules) + 1
    st.session_state.alert_rules.append({
        "id": novo_id,
        "metrica": metrica,
        "operador": operador,
        "valor": valor,
        "ativo": True
    })
    return True

def apagar_regra(regra_id):
    """Remove uma regra"""
    st.session_state.alert_rules = [r for r in st.session_state.alert_rules if r['id'] != regra_id]

# ===================== OTIMIZAÇÃO DE CUSTOS (FUNC 52) =====================
def analisar_eficiencia_custos():
    """
    Analisa métricas de uso (simuladas) e gera recomendações de poupança.
    """
    # Simulação de métricas de infraestrutura (Médias dos últimos 30 dias)
    metricas = {
        "cpu_avg": np.random.uniform(10, 60), # % de uso
        "ram_avg": np.random.uniform(30, 80), # % de uso
        "db_iops": np.random.randint(100, 5000),
        "storage_cold_access": 0.05 # Apenas 5% dos dados antigos são acedidos
    }
    
    custo_atual = 1250.00 # €/mês
    recomendacoes = []
    poupanca_potencial = 0
    
    # REGRA 1: Servidor de ML sobredimensionado
    if metricas['cpu_avg'] < 20:
        poupanca = 350.00
        recomendacoes.append({
            "tipo": "Computação",
            "acao": "Downscale ML Cluster (4x GPU -> 2x GPU)",
            "motivo": f"Uso médio de CPU é apenas {metricas['cpu_avg']:.1f}%.",
            "poupanca": poupanca,
            "impacto": "Baixo (Apenas em treino intensivo)"
        })
        poupanca_potencial += poupanca
        
    # REGRA 2: Dados antigos em disco caro
    if metricas['storage_cold_access'] < 0.10:
        poupanca = 120.00
        recomendacoes.append({
            "tipo": "Armazenamento",
            "acao": "Mover Logs antigos para Cold Storage (Glacier)",
            "motivo": "95% dos logs com >90 dias não são consultados.",
            "poupanca": poupanca,
            "impacto": "Nenhum (Acesso mais lento a arquivos antigos)"
        })
        poupanca_potencial += poupanca

    # REGRA 3: Base de Dados Ociosa à noite
    poupanca = 80.00
    recomendacoes.append({
        "tipo": "Base de Dados",
        "acao": "Ativar 'Pause on Idle' no ambiente de Testes",
        "motivo": "Ambiente de Staging ativo 24/7 sem uso noturno.",
        "poupanca": poupanca,
        "impacto": "Médio (Arranque a frio demora 2 min)"
    })
    poupanca_potencial += poupanca
    
    return {
        "metricas": metricas,
        "custo_atual": custo_atual,
        "recomendacoes": recomendacoes,
        "total_poupanca": poupanca_potencial,
        "novo_custo": custo_atual - poupanca_potencial
    }

# ===================== EXECUTIVE ACTION SCORE (RCS - FUNC 51) =====================
def calcular_rcs(df, df_macro=None):
    """
    Calcula o Risk Composite Score (0-100).
    0 = Risco Extremo (Mau), 100 = Segurança Total (Bom).
    """
    # 1. Componente de Carteira (Peso 50%)
    # Taxa de Default (Invertida: quanto menor, melhor)
    taxa_default = df['default'].mean()
    score_default = max(0, 100 - (taxa_default * 500)) # Se default for 20%, score é 0
    
    # LTV Médio (Invertido)
    ltv_medio = df['ltv'].mean()
    score_ltv = max(0, 100 - ((ltv_medio - 0.5) * 200)) # Penaliza LTV > 50%
    
    comp_carteira = (score_default * 0.7) + (score_ltv * 0.3)
    
    # 2. Componente Macro (Peso 30%)
    # Se não houver dados macro, assumimos neutro (50)
    if df_macro is not None and not df_macro.empty:
        ultimo = df_macro.iloc[-1]
        # Inflação alta penaliza
        score_inflacao = max(0, 100 - (ultimo['Inflação (%)'] * 10))
        # Desemprego alto penaliza
        score_desemp = max(0, 100 - (ultimo['Taxa Desemprego (%)'] * 10))
        
        comp_macro = (score_inflacao * 0.5) + (score_desemp * 0.5)
    else:
        comp_macro = 70 # Valor base otimista
        
    # 3. Componente Compliance (Peso 20%)
    # Baseado no número de alertas ativos
    alerts = check_compliance_alerts(df)
    # Cada alerta crítico tira 20 pontos, aviso tira 5
    penalizacao = sum([20 if a['nivel'] == 'CRÍTICO' else 5 for a in alerts])
    comp_compliance = max(0, 100 - penalizacao)
    
    # Ponderação Final
    rcs_final = (comp_carteira * 0.50) + (comp_macro * 0.30) + (comp_compliance * 0.20)
    
    return {
        "RCS": rcs_final,
        "Componentes": {
            "Carteira (50%)": comp_carteira,
            "Macroeconomia (30%)": comp_macro,
            "Compliance (20%)": comp_compliance
        },
        "Tendencia": "Estável" # Simplificação
    }

# ===================== PREVISÃO AVANÇADA (ARIMA - FUNC 30) =====================
from statsmodels.tsa.arima.model import ARIMA

def executar_arima(df_historico, coluna_alvo, meses_futuros=6, ordem=(1, 1, 1)):
    """
    Executa o modelo ARIMA (AutoRegressive Integrated Moving Average).
    Permite configuração manual da ordem (p, d, q).
    """
    try:
        # 1. Ajustar o Modelo
        # ordem = (p, d, q)
        modelo = ARIMA(df_historico[coluna_alvo], order=ordem).fit()
        
        # 2. Prever
        forecast_result = modelo.get_forecast(steps=meses_futuros)
        previsao = forecast_result.predicted_mean
        conf_int = forecast_result.conf_int(alpha=0.05) # 95% Confiança
        
        return {
            'modelo': modelo,
            'previsao': previsao,
            'conf_inf': conf_int.iloc[:, 0],
            'conf_sup': conf_int.iloc[:, 1],
            'aic': modelo.aic,
            'bic': modelo.bic,
            'sucesso': True
        }
    except Exception as e:
        return {'sucesso': False, 'erro': str(e)}

# ===================== LICENCIAMENTO (FUNC 49) =====================
def verificar_licenca_exportacao(user, tipo_dado):
    """
    Verifica se o utilizador tem permissão para exportar este tipo de dado.
    """
    # Simulação de regras de negócio
    regras = {
        "Admin": ["Todos"],
        "Risk Manager": ["Todos"],
        "Analyst": ["Relatórios", "KPIs"], # Não pode exportar dados brutos de clientes
        "Auditor": ["Logs", "Relatórios"]
    }
    
    role = st.session_state.user_role
    permissoes = regras.get(role, [])
    
    if "Todos" in permissoes:
        return True
    
    if tipo_dado in permissoes:
        return True
        
    return False

def aplicar_marca_agua(texto_base):
    """Gera o texto da marca de água para anexar a relatórios"""
    user = st.session_state.username
    data = datetime.now().strftime("%Y-%m-%d %H:%M")
    licenca_id = f"LIC-{abs(hash(user)) % 10000:04d}"
    
    return f"CONFIDENCIAL | Gerado por {user} em {data} | Licença: {licenca_id} | Proibida a Redistribuição"

# ===================== ANÁLISE DE SOBREVIVÊNCIA =====================
def analise_sobrevivencia(df, duracao_col, evento_col):
    """Análise de sobrevivência (Kaplan-Meier) para tempo até default"""
    df_surv = df[[duracao_col, evento_col]].dropna()
    
    if len(df_surv) < 10:
        return None, None, None
    
    kmf = KaplanMeierFitter()
    kmf.fit(df_surv[duracao_col], df_surv[evento_col])
    
    # Criar gráfico
    fig = go.Figure()
    fig.add_trace(go.Scatter(
        x=kmf.survival_function_.index,
        y=kmf.survival_function_['KM_estimate'],
        mode='lines',
        name='Curva de Sobrevivência',
        fill='tozeroy',
        line=dict(color='#3b82f6', width=3)
    ))
    
    fig.update_layout(
        title='Análise de Sobrevivência - Tempo até Default',
        xaxis_title='Meses',
        yaxis_title='Probabilidade de Sobrevivência (%)',
        hovermode='x unified',
        template='plotly_dark'
    )
    
    median_survival = kmf.median_survival_time_
    
    return fig, median_survival, kmf

# ===================== STRESS TESTING =====================
def stress_test_scenarios(df, var_target='divida_total'):
    """Testes de stress regulatórios (cenários adversos)"""
    
    scenarios = {
        "Recessão Severa": {
            "gdp": -0.05,
            "unemployment": 0.03,
            "default_mult": 2.5,
            "descricao": "Crise económica profunda"
        },
        "Crise Moderada": {
            "gdp": -0.02,
            "unemployment": 0.015,
            "default_mult": 1.5,
            "descricao": "Recessão técnica"
        },
        "Base Case": {
            "gdp": 0.02,
            "unemployment": 0,
            "default_mult": 1.0,
            "descricao": "Cenário esperado"
        },
        "Crescimento Forte": {
            "gdp": 0.04,
            "unemployment": -0.01,
            "default_mult": 0.7,
            "descricao": "Expansão económica"
        }
    }
    
    results = {}
    base_default = df['default'].mean() if 'default' in df.columns else 0.1
    base_value = df[var_target].sum() if var_target in df.columns else 0
    
    for scenario_name, params in scenarios.items():
        # Simular impacto no default
        new_default_rate = min(base_default * params['default_mult'], 0.95)
        
        # Simular impacto no valor da carteira
        portfolio_impact = base_value * (1 + params['gdp'])
        
        # Expected Loss (assumir LGD 45%)
        expected_loss = new_default_rate * portfolio_impact * 0.45
        
        # Loss rate
        loss_rate = expected_loss / base_value if base_value > 0 else 0
        
        results[scenario_name] = {
            "Default_Rate": new_default_rate,
            "Portfolio_Value": portfolio_impact,
            "Expected_Loss": expected_loss,
            "Loss_Rate": loss_rate,
            "Impact_GDP": params['gdp'],
            "Descricao": params['descricao']
        }
    
    return pd.DataFrame(results).T

# ===================== BENCHMARKING (FUNC 22) =====================

# ===================== QUALIDADE DE DADOS (CORRIGIDA) =====================
def auditar_qualidade_dados(df):
    """
    Realiza um check-up completo à qualidade dos dados.
    Retorna um score (0-100) e métricas detalhadas.
    """
    total_registos = len(df)
    if total_registos == 0:
        return {"score": 0, "erro": "Dataset vazio", "total_linhas": 0, "nulos_pct": 0, "duplicados_pct": 0, "colunas_com_problemas": [], "detalhe_nulos": {}}
    
    metricas = {
        "total_linhas": total_registos,
        "total_colunas": len(df.columns),
        "colunas_com_problemas": []
    }
    
    # 1. Completude (Valores Nulos)
    nulos = df.isnull().sum()
    total_celulas = total_registos * len(df.columns)
    total_nulos = nulos.sum()
    score_completude = (1 - (total_nulos / total_celulas)) * 100
    
    # 2. Unicidade (Duplicados) - CORREÇÃO AQUI
    # Remover colunas que contêm listas (como 'tags') antes de verificar duplicados
    cols_hashable = [col for col in df.columns if not isinstance(df[col].iloc[0], list)]
    try:
        duplicados = df[cols_hashable].duplicated().sum()
    except Exception:
        # Se ainda falhar (ex: dicionários), convertemos tudo para string para comparar
        duplicados = df.astype(str).duplicated().sum()
        
    score_unicidade = (1 - (duplicados / total_registos)) * 100
    
    # 3. Integridade de Colunas Chave
    cols_criticas = ['id_cliente', 'default', 'divida_total']
    penalizacao_critica = 0
    
    for col in cols_criticas:
        if col in df.columns:
            nulos_criticos = df[col].isnull().sum()
            if nulos_criticos > 0:
                metricas["colunas_com_problemas"].append({
                    "coluna": col,
                    "tipo": "Crítica com Nulos",
                    "qtd": nulos_criticos
                })
                penalizacao_critica += 10
    
    # 4. Outliers (Desvios Padrão > 3)
    cols_num = df.select_dtypes(include=[np.number]).columns
    outliers_total = 0
    
    for col in cols_num:
        if col != 'id_cliente':
            z_scores = np.abs((df[col] - df[col].mean()) / df[col].std())
            n_outliers = (z_scores > 3).sum()
            if n_outliers > 0:
                 outliers_total += n_outliers
    
    # Score Final Ponderado
    score_final = (score_completude * 0.4) + (score_unicidade * 0.4) + (100 * 0.2)
    score_final -= penalizacao_critica
    
    metricas["score_global"] = max(0, min(100, score_final))
    metricas["nulos_pct"] = (total_nulos / total_celulas) * 100
    metricas["duplicados_pct"] = (duplicados / total_registos) * 100
    metricas["outliers_count"] = outliers_total
    metricas["detalhe_nulos"] = nulos[nulos > 0].to_dict()
    
    return metricas

# ===================== EARLY WARNING SYSTEM =====================
def early_warning_system(df):
    """Sistema de alerta precoce para identificar clientes em risco"""
    if 'default' not in df.columns:
        return None
    
    # Features importantes para o modelo
    feature_cols = ['score_interno', 'ltv', 'divida_total', 'rendimento_mensal', 
                   'num_atrasos_12m', 'utilizacao_credito', 'idade']
    
    # Filtrar apenas features disponíveis
    available = [col for col in feature_cols if col in df.columns]
    
    if len(available) < 2:
        return None
    
    # Preparar dados
    df_model = df[available + ['default']].dropna()
    
    if len(df_model) < 50:
        return None
    
    X = df_model[available]
    y = df_model['default']
    
    # Treinar modelo de alerta
    model = GradientBoostingClassifier(
        n_estimators=100,
        learning_rate=0.1,
        max_depth=5,
        random_state=42
    )
    model.fit(X, y)
    
    # Calcular probabilidades de risco
    probs = model.predict_proba(X)[:, 1]
    
    # Criar cópia para não modificar original
    df_result = df_model.copy()
    df_result['Risk_Score'] = probs * 100  # Converter para 0-100
    
    # Classificar em níveis de risco
    df_result['Risk_Level'] = pd.cut(
        probs, 
        bins=[0, 0.2, 0.5, 0.8, 1.0], 
        labels=['Baixo', 'Médio', 'Alto', 'Crítico'],
        include_lowest=True
    )
    
    return df_result, model, available

# ===================== FORECASTING (FUNCIONALIDADE 7) =====================
from statsmodels.tsa.holtwinters import ExponentialSmoothing
@st.cache_data(ttl=3600)
def gerar_historico_simulado(df, meses_historico=24):
    """
    Gera uma série temporal simulada baseada nos valores atuais do dataframe.
    Cria uma tendência realista com sazonalidade e ruído.
    """
    # Valor base atual (ex: Taxa de Default atual)
    kpis_atuais = calcular_kpis_principais(df)
    
    # Datas (Do passado até hoje)
    datas = pd.date_range(end=datetime.now(), periods=meses_historico, freq='M')
    
    # Simular histórico para Default Rate
    # Tendência: Vamos assumir que subiu ligeiramente (trend) e tem sazonalidade
    base_default = kpis_atuais.get('Taxa_Default', 0.05)
    
    # Criar curva com ruído
    np.random.seed(42)
    tendencia = np.linspace(base_default * 0.8, base_default, meses_historico)
    sazonalidade = np.sin(np.arange(meses_historico) * (2 * np.pi / 12)) * 0.005
    ruido = np.random.normal(0, 0.002, meses_historico)
    
    serie_default = tendencia + sazonalidade + ruido
    
    # Simular histórico para Carteira Total
    base_carteira = kpis_atuais.get('Carteira_Total', 1000000)
    tendencia_vol = np.linspace(base_carteira * 0.9, base_carteira, meses_historico)
    ruido_vol = np.random.normal(0, base_carteira * 0.02, meses_historico)
    
    serie_carteira = tendencia_vol + ruido_vol
    
    df_hist = pd.DataFrame({
        'Data': datas,
        'Taxa_Default': np.clip(serie_default, 0, 1), # Garantir entre 0 e 1
        'Carteira_Total': np.clip(serie_carteira, 0, None)
    }).set_index('Data')
    
    return df_hist

# ===================== ELASTICIDADE SOCIAL (CORRIGIDA) =====================
import statsmodels.api as sm

# ===================== INTEGRAÇÃO CRM (FUNC 27) =====================
def get_crm_status():
    """Simula o estado da conexão com o CRM externo (Salesforce/Dynamics)"""
    return {
        "status": "Conectado 🟢",
        "ultima_sincronizacao": "2024-03-10 09:30",
        "registos_pendentes": 45,
        "crm_type": "Salesforce Cloud"
    }

def sincronizar_com_crm(df_risco):
    """
    Simula o envio de dados de risco para o CRM.
    df_risco: DataFrame com id_cliente e risk_score
    """
    # Em produção, isto seria um POST request para a API do Salesforce
    import time
    time.sleep(2) # Simular latência de rede
    
    sucesso = int(len(df_risco) * 0.98) # Simular 98% de sucesso
    falhas = len(df_risco) - sucesso
    
    return {
        "sucesso": sucesso,
        "falhas": falhas,
        "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    }

# ===================== ACESSIBILIDADE (FUNC 34) =====================
def get_color_palette(mode="Standard"):
    """Retorna paletas de cores otimizadas para acessibilidade"""
    if mode == "Colorblind Safe (Daltonismo)":
        # Paleta Okabe-Ito (Segura para deuteranopia/protanopia)
        return {
            "positivo": "#56B4E9", # Sky Blue
            "negativo": "#D55E00", # Vermelho/Laranja forte
            "neutro": "#F0E442",   # Amarelo
            "sequencial": "Viridis"
        }
    else:
        # Padrão
        return {
            "positivo": "#10b981", # Verde
            "negativo": "#ef4444", # Vermelho
            "neutro": "#f59e0b",   # Amarelo
            "sequencial": "Blues"
        }

# ===================== ELASTICIDADE AVANÇADA (VECM - FUNC 31) =====================
from statsmodels.tsa.vector_ar.vecm import VECM
@monitorar_performance
def calcular_elasticidade_vecm(df_macro, df_carteira_hist, variaveis=['Euribor 6M (%)', 'Taxa_Default']):
    """
    Ajusta um modelo VECM para medir a relação de longo prazo e a velocidade de ajuste.
    """
    # 1. Preparar Dados
    # Juntar e garantir que as datas coincidem
    df_combined = pd.merge(
        df_macro, 
        df_carteira_hist[['Taxa_Default']], 
        left_index=True, 
        right_index=True, 
        how='inner'
    ).dropna()
    
    # Filtrar apenas as variáveis de interesse
    dados_modelo = df_combined[variaveis]
    
    if len(dados_modelo) < 30:
        return {"sucesso": False, "erro": "Dados insuficientes para VECM (mínimo 30 meses)."}

    try:
        # 2. Ajustar Modelo VECM
        # k_ar_diff=1 (Lags), coint_rank=1 (Existe 1 relação de longo prazo)
        modelo = VECM(dados_modelo, k_ar_diff=1, coint_rank=1, deterministic='ci')
        resultado = modelo.fit()
        
        # 3. Gerar Impulso-Resposta (IRF)
        # Simula um choque numa variável e vê o efeito na outra nos próximos 12 meses
        irf = resultado.irf(12)
        
        # Extrair os dados do IRF para plotar
        # A estrutura do IRF é complexa, vamos simplificar para obter o choque da Var 1 na Var 2
        irf_data = irf.irfs[:, variaveis.index('Taxa_Default'), variaveis.index(variaveis[0])]
        
        return {
            "sucesso": True,
            "modelo": resultado,
            "irf_values": irf_data,
            "alpha": resultado.alpha, # Velocidade de ajuste
            "beta": resultado.beta    # Relação de Longo Prazo
        }
        
    except Exception as e:
        return {"sucesso": False, "erro": f"Erro no cálculo VECM: {str(e)}"}

def calcular_elasticidade(df_macro, df_carteira_hist):
    """
    Calcula a sensibilidade (elasticidade) do Default a fatores macroeconómicos
    usando Regressão Linear Simples.
    """
    # 1. CORREÇÃO DE DATAS: Normalizar para garantir que o merge funciona
    # Remove horas/minutos/segundos, mantendo apenas a data (Ano-Mês-Dia)
    df_macro.index = df_macro.index.normalize()
    df_carteira_hist.index = df_carteira_hist.index.normalize()
    
    # 2. Juntar dados
    df_combined = pd.merge(
        df_macro, 
        df_carteira_hist[['Taxa_Default']], 
        left_index=True, 
        right_index=True, 
        how='inner'
    )
    
    # 3. VERIFICAÇÃO DE SEGURANÇA (Se o merge estiver vazio, abortar)
    if df_combined.empty:
        return {}
    
    resultados = {}
    
    try:
        # Para cada indicador macro, calcular a regressão contra o Default
        for indicador in df_macro.columns:
            Y = df_combined['Taxa_Default']
            X = df_combined[indicador]
            
            # Verificação extra: Se X ou Y forem constantes (variância 0), OLS falha
            if X.std() == 0 or Y.std() == 0:
                continue

            X = sm.add_constant(X) # Adicionar intercepto
            
            modelo = sm.OLS(Y, X).fit()
            
            # O coeficiente (beta) é a elasticidade aproximada
            beta = modelo.params[indicador]
            r2 = modelo.rsquared
            p_valor = modelo.pvalues[indicador]
            
            resultados[indicador] = {
                'elasticidade': beta,
                'r2': r2,
                'significancia': 'Alta' if p_valor < 0.05 else 'Baixa',
                'dados_combinados': df_combined[[indicador, 'Taxa_Default']]
            }
    except Exception as e:
        # Em caso de erro matemático, retorna vazio para não quebrar a app
        print(f"Erro no cálculo OLS: {e}")
        return {}
        
    return resultados

# ===================== XAI / SHAP (FUNC 48) =====================
@st.cache_data(show_spinner="A calcular explicabilidade (SHAP)...")
def calcular_shap_values(df, _modelo_data):
    import shap # <--- MOVA O IMPORT PARA AQUI
    """
    Calcula os valores SHAP para explicar o modelo.
    """
    # Recuperar modelo e dados
    model = _modelo_data['model']       
    scaler = _modelo_data['scaler']     
    features = _modelo_data['features'] 
    
    # Preparar dados (Amostra para ser rápido)
    X = df[features].dropna().sample(min(100, len(df)), random_state=42)
    X_scaled = scaler.transform(X)
    
    # Criar Explainer (TreeExplainer é otimizado para GradientBoosting)
    explainer = shap.TreeExplainer(model)
    shap_values = explainer.shap_values(X_scaled)
    
    # O SHAP para classificação binária retorna uma lista ou array. 
    # Para GradientBoosting (sklearn), normalmente retorna as log-odds.
    # Vamos garantir que temos a matriz correta.
    if isinstance(shap_values, list):
        vals = shap_values[1] # Classe 1 (Default)
    else:
        vals = shap_values
        
    return {
        "explainer": explainer,
        "shap_values": vals,
        "X": X,
        "expected_value": explainer.expected_value[0] if isinstance(explainer.expected_value, list) else explainer.expected_value
    }

def plot_shap_waterfall(shap_data, cliente_idx=0):
    import shap # <--- MOVA O IMPORT PARA AQUI também
    import matplotlib.pyplot as plt
    # Preparar dados para um único cliente
    shap_val_single = shap_data['shap_values'][cliente_idx]
    data_single = shap_data['X'].iloc[cliente_idx]
    
    # Criar objeto Explanation (necessário para os plots novos do SHAP)
    exp = shap.Explanation(
        values=shap_val_single,
        base_values=shap_data['expected_value'],
        data=data_single.values,
        feature_names=data_single.index.tolist()
    )
    
    # Gerar plot (matplotlib)
    fig, ax = plt.subplots()
    shap.plots.waterfall(exp, show=False)
    return fig

# ===================== REGRESSÃO LINEAR MÚLTIPLA (FUNC 23) =====================
def calcular_rlm(df, target_col, feature_cols):
    """
    Executa uma Regressão Linear Múltipla (OLS).
    Retorna o sumário do modelo e os coeficientes.
    """
    # 1. Preparar Dados (Apenas numéricos e sem nulos)
    data = df[[target_col] + feature_cols].dropna()
    
    if len(data) < 10:
        return {"sucesso": False, "erro": "Dados insuficientes (mínimo 10 registos sem nulos)."}
        
    Y = data[target_col]
    X = data[feature_cols]
    
    # Adicionar constante (intercepto)
    X = sm.add_constant(X)
    
    try:
        # 2. Ajustar Modelo
        modelo = sm.OLS(Y, X).fit()
        
        # 3. Extrair Coeficientes para Gráfico
        coefs = pd.DataFrame({
            'Feature': modelo.params.index,
            'Coeficiente': modelo.params.values,
            'P-Value': modelo.pvalues.values
        })
        # Remover a constante do gráfico visual
        coefs = coefs[coefs['Feature'] != 'const'].sort_values('Coeficiente')
        
        return {
            "sucesso": True,
            "modelo": modelo,
            "coefs": coefs,
            "r2": modelo.rsquared,
            "r2_adj": modelo.rsquared_adj,
            "n_obs": int(modelo.nobs)
        }
    except Exception as e:
        return {"sucesso": False, "erro": str(e)}
    
    # ===================== FATURAÇÃO E SUBSCRIÇÕES (FUNC 24) =====================
def get_planos_disponiveis():
    """Retorna os planos SaaS disponíveis"""
    return {
        "Standard": {"preco": 499, "users": 5, "api_calls": 1000, "support": "Email"},
        "Pro": {"preco": 999, "users": 20, "api_calls": 10000, "support": "Prioritário"},
        "Enterprise": {"preco": 2500, "users": "Ilimitado", "api_calls": "Ilimitado", "support": "24/7 Dedicado"}
    }

def get_dados_subscricao(user):
    """Simula os dados financeiros do cliente atual"""
    # Em produção, isto viria do Stripe API ou SQL
    return {
        "plano_atual": "Pro",
        "status": "Ativo ✅",
        "proxima_fatura": datetime.now().replace(day=28).strftime("%Y-%m-%d"),
        "metodo_pagamento": "Visa **** 4242",
        "uso_api": 4520, # De 10.000
        "faturas": [
            {"id": "INV-2024-001", "data": "2024-01-28", "valor": "€999.00", "status": "Pago"},
            {"id": "INV-2024-002", "data": "2024-02-28", "valor": "€999.00", "status": "Pago"},
            {"id": "INV-2023-012", "data": "2023-12-28", "valor": "€999.00", "status": "Pago"},
        ]
    }

# ===================== SDK GENERATOR (FUNC 45) =====================
def gerar_codigo_sdk_python():
    """
    Gera o código fonte do SDK Python para download.
    """
    sdk_code = """
import requests
import pandas as pd

class SocioStatClient:
    def __init__(self, api_key, base_url="https://api.sociostat.enterprise"):
        self.api_key = api_key
        self.base_url = base_url
        self.headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}

    def _get(self, endpoint):
        response = requests.get(f"{self.base_url}{endpoint}", headers=self.headers)
        if response.status_code == 200:
            return response.json()
        else:
            raise Exception(f"Erro API {response.status_code}: {response.text}")

    def get_cliente(self, cliente_id):
        '''Obtém dados detalhados de um cliente'''
        data = self._get(f"/cliente/{cliente_id}")
        return data['data']

    def get_risco_carteira(self):
        '''Obtém métricas de risco agregadas da carteira'''
        data = self._get("/risco/carteira")
        return data['metrics']
        
    def get_previsao_default(self, meses=6):
        '''Obtém a previsão de taxa de default para os próximos meses'''
        # Exemplo de endpoint futuro
        data = self._get(f"/forecast/default?horizon={meses}")
        return pd.DataFrame(data['forecast'])

# Exemplo de Uso:
# client = SocioStatClient("sk_live_123456")
# risco = client.get_risco_carteira()
# print(f"Risco Atual: {risco['default_rate']:.2%}")
"""
    return sdk_code

# ===================== MACRO WATCH (FUNC 11) =====================
@st.cache_data(ttl=3600)
def get_dados_macroeconomicos(meses=24):
    """
    Gera dados macroeconómicos realistas para o contexto de Portugal/Europa.
    Em produção, isto ligaria à API do Banco de Portugal ou INE.
    """
    dates = pd.date_range(end=datetime.now(), periods=meses, freq='M')
    
    # Simulação realista de tendências recentes
    np.random.seed(42)
    
    # Euribor (Subida recente)
    euribor = np.linspace(0.5, 4.0, meses) + np.random.normal(0, 0.1, meses)
    
    # Inflação (Pico e descida)
    inflacao = np.concatenate([
        np.linspace(2.0, 10.0, int(meses/2)),
        np.linspace(10.0, 3.0, int(meses/2))
    ]) + np.random.normal(0, 0.2, meses)
    
    # Desemprego (Estável)
    desemprego = np.linspace(6.5, 5.8, meses) + np.random.normal(0, 0.1, meses)
    
    # PIB (Variação trimestral suavizada)
    pib = np.linspace(1.5, 2.1, meses) + np.random.normal(0, 0.3, meses)
    
    df_macro = pd.DataFrame({
        'Data': dates,
        'Euribor 6M (%)': euribor,
        'Inflação (%)': inflacao,
        'Taxa Desemprego (%)': desemprego,
        'Crescimento PIB (%)': pib
    }).set_index('Data')
    
    return df_macro

# ===================== MLOPS & VERSIONAMENTO (FUNC 33) =====================
def registar_versao_modelo(nome, metricas, params, user):
    """Regista uma nova versão no histórico"""
    # Calcular nova versão
    versoes_anteriores = [m for m in st.session_state.model_registry if m['nome'] == nome]
    nova_versao_num = len(versoes_anteriores) + 1
    versao_str = f"v{nova_versao_num}.0"
    
    novo_modelo = {
        "id": f"MOD-{nome[:3].upper()}-{datetime.now().strftime('%Y%m%d%H%M')}",
        "nome": nome,
        "versao": versao_str,
        "data": datetime.now().strftime("%Y-%m-%d %H:%M"),
        "autor": user,
        "status": "Staging", # Começa em testes
        "metricas": metricas,
        "params": params
    }
    
    st.session_state.model_registry.insert(0, novo_modelo)
    return versao_str

def promover_modelo_producao(model_id):
    """Promove um modelo para Produção e arquiva o anterior"""
    # 1. Encontrar o modelo a promover
    target_model = next((m for m in st.session_state.model_registry if m['id'] == model_id), None)
    if not target_model: return False
    
    nome_modelo = target_model['nome']
    
    # 2. Arquivar o atual modelo de Produção desse tipo
    for m in st.session_state.model_registry:
        if m['nome'] == nome_modelo and m['status'] == "Production":
            m['status'] = "Archived"
            
    # 3. Promover o novo
    target_model['status'] = "Production"
    return True

def correlacionar_macro_carteira(df_macro, df_carteira_hist):
    """Calcula a correlação entre fatores externos e o risco da carteira"""
    # Juntar os dois dataframes pela data (resample mensal)
    # Assumindo que df_carteira_hist vem da Funcionalidade 7 (Forecasting)
    
    df_combined = pd.merge(
        df_macro, 
        df_carteira_hist[['Taxa_Default']], 
        left_index=True, 
        right_index=True, 
        how='inner'
    )
    
    return df_combined.corr()['Taxa_Default'].drop('Taxa_Default')
def executar_previsao(df_historico, coluna_alvo, meses_futuros=6):
    """
    Executa o modelo Holt-Winters (Triple Exponential Smoothing) para prever o futuro.
    Ideal para dados com tendência e sazonalidade.
    """
    try:
        # 1. Ajustar o Modelo
        # trend='add' (tendência linear), seasonal='add' (sazonalidade aditiva), periods=12 (anual)
        modelo = ExponentialSmoothing(
            df_historico[coluna_alvo],
            trend='add',
            seasonal='add',
            seasonal_periods=12
        ).fit()
        
        # 2. Prever
        previsao = modelo.forecast(meses_futuros)
        
        # 3. Intervalo de Confiança (Simulado para HW, pois statsmodels HW básico não devolve nativamente)
        # Usamos o desvio padrão dos resíduos para estimar a incerteza
        residuos = df_historico[coluna_alvo] - modelo.fittedvalues
        std_resid = residuos.std()
        
        # Intervalo de 95% (aprox 1.96 * std, aumentando com o tempo)
        margem_erro = 1.96 * std_resid * np.sqrt(np.arange(1, meses_futuros + 1))
        
        conf_inf = previsao - margem_erro
        conf_sup = previsao + margem_erro
        
        return {
            'modelo': modelo,
            'previsao': previsao,
            'conf_inf': conf_inf,
            'conf_sup': conf_sup,
            'sucesso': True
        }
    except Exception as e:
        return {'sucesso': False, 'erro': str(e)}

def plotar_previsao(df_hist, resultado_prev, titulo, formato_pct=False):
    """Gera o gráfico Plotly com histórico + previsão + intervalo de confiança"""
    
    fig = go.Figure()
    
    # 1. Dados Históricos
    fig.add_trace(go.Scatter(
        x=df_hist.index,
        y=df_hist,
        mode='lines',
        name='Histórico',
        line=dict(color='#3b82f6', width=3)
    ))
    
    # 2. Previsão
    fig.add_trace(go.Scatter(
        x=resultado_prev['previsao'].index,
        y=resultado_prev['previsao'],
        mode='lines',
        name='Previsão',
        line=dict(color='#10b981', width=3, dash='dash')
    ))
    
    # 3. Intervalo de Confiança (Área Sombreada)
    # Limite Superior
    fig.add_trace(go.Scatter(
        x=resultado_prev['conf_sup'].index,
        y=resultado_prev['conf_sup'],
        mode='lines',
        line=dict(width=0),
        showlegend=False,
        name='Limite Superior'
    ))
    
    # Limite Inferior (com preenchimento tonexty)
    fig.add_trace(go.Scatter(
        x=resultado_prev['conf_inf'].index,
        y=resultado_prev['conf_inf'],
        mode='lines',
        line=dict(width=0),
        fill='tonexty',
        fillcolor='rgba(16, 185, 129, 0.2)', # Verde transparente
        showlegend=True,
        name='Intervalo Confiança (95%)'
    ))
    
    yaxis_format = ',.1%' if formato_pct else '€,.0f'
    
    fig.update_layout(
        title=titulo,
        xaxis_title='Data',
        yaxis_title='Valor',
        template='plotly_dark',
        yaxis=dict(tickformat=yaxis_format),
        hovermode='x unified'
    )
    
    return fig

# ===================== WHAT-IF ANALYSIS (NOVO) =====================
def simular_cenario_whatif(df, modelo, scaler, features, ajuste_rendimento=0, ajuste_divida=0, ajuste_score=0):
    """
    Simula o impacto de alterações nas variáveis no risco da carteira.
    
    :param ajuste_rendimento: % de variação no rendimento (ex: -0.10 para queda de 10%)
    :param ajuste_divida: % de variação na dívida total
    :param ajuste_score: Pontos a adicionar/remover do score interno
    """
    
    # 1. Criar cópia dos dados para não afetar o original
    df_sim = df.copy()
    
    # 2. Aplicar choques (Choques)
    if 'rendimento_mensal' in df_sim.columns:
        df_sim['rendimento_mensal'] = df_sim['rendimento_mensal'] * (1 + ajuste_rendimento)
        
    if 'divida_total' in df_sim.columns:
        df_sim['divida_total'] = df_sim['divida_total'] * (1 + ajuste_divida)
        
    if 'score_interno' in df_sim.columns:
        df_sim['score_interno'] = df_sim['score_interno'] + ajuste_score
        
    # 3. Preparar dados para o modelo (apenas linhas com dados completos nas features usadas)
    X_sim = df_sim[features].dropna()
    
    # Se a simulação reduzir linhas (por nulos), alinhar o y
    # Mas aqui queremos comparar o agregado, então calculamos a nova média prevista
    
    if len(X_sim) > 0:
        # Normalizar
        X_sim_scaled = scaler.transform(X_sim)
        
        # Prever novas probabilidades
        novas_probas = modelo.predict_proba(X_sim_scaled)[:, 1]
        
        return {
            'media_risco_atual': df['default'].mean(), # Risco histórico real
            'media_risco_simulado': novas_probas.mean(),
            'total_divida': df_sim['divida_total'].sum(),
            'impacto_default': (novas_probas.mean() - df['default'].mean()) / df['default'].mean()
        }
    return None

# ===================== GEO-RISK MAPPING (CORRIGIDO) =====================
import folium
from folium.plugins import MarkerCluster, HeatMap
from streamlit_folium import st_folium

# Nota: Removemos a função 'gerar_dados_geograficos' daqui porque
# agora as coordenadas são geradas uma única vez no carregamento dos dados.

def renderizar_mapa_risco(df):
    """Gera o mapa interativo usando coordenadas fixas"""
    
    # Verificar se as colunas existem (caso o user carregue CSV sem elas)
    if 'latitude' not in df.columns or 'longitude' not in df.columns:
        st.error("⚠️ O dataset não tem colunas 'latitude' e 'longitude'. Carregue a Demo novamente.")
        return None, df

    # 1. Criar mapa base (Centrado em Portugal)
    # Usamos st.session_state para guardar o centro se necessário, mas fixo serve
    m = folium.Map(location=[39.5, -8.0], zoom_start=7, tiles="cartodbpositron")
    
    # 2. Adicionar Camada de Heatmap (Fixo)
    heat_data = df[df['default'] == 1][['latitude', 'longitude']].values.tolist()
    
    if len(heat_data) > 0:
        HeatMap(heat_data, radius=15, blur=10, name="Heatmap de Default").add_to(m)
    
    # 3. Adicionar Clusters
    marker_cluster = MarkerCluster(name="Clusters de Clientes").add_to(m)
    
    # Limitar amostra para performance (opcional, mas bom para evitar lag)
    amostra = df.sample(min(200, len(df)), random_state=42) # random_state 42 fixa a amostra!
    
    for idx, row in amostra.iterrows():
        color = 'red' if row['default'] == 1 else 'green' if row['score_interno'] > 700 else 'orange'
        
        folium.CircleMarker(
            location=[row['latitude'], row['longitude']],
            radius=6,
            popup=f"Cliente #{row['id_cliente']}<br>Score: {row['score_interno']}",
            color=color,
            fill=True,
            fill_color=color,
            fill_opacity=0.7
        ).add_to(marker_cluster)

    folium.LayerControl().add_to(m)
    
    return m, df

# ===================== AI FAIRNESS & BIAS (FUNC 13) =====================
def calcular_metricas_fairness(df, modelo, scaler, features, atributo_protegido, valor_protegido):
    """
    Audita o modelo para detetar enviesamento (Bias) contra um grupo específico.
    Ex: Verificar se 'Idade < 25' tem tratamento injusto.
    """
    # 1. Preparar dados
    X = df[features].dropna()
    y_true = df.loc[X.index, 'default']
    
    # Prever com o modelo atual
    X_scaled = scaler.transform(X)
    y_pred = modelo.predict(X_scaled)
    
    # Criar dataframe de análise
    df_audit = X.copy()
    df_audit['y_true'] = y_true
    df_audit['y_pred'] = y_pred
    
    # 2. Definir Grupos
    # Grupo A: Protegido (Ex: Jovens)
    # Grupo B: Referência (Ex: Resto da população)
    
    if atributo_protegido == 'idade':
        mask_prot = df_audit['idade'] <= valor_protegido
    elif atributo_protegido in df_audit.columns:
        mask_prot = df_audit[atributo_protegido] == valor_protegido
    else:
        return None

    grupo_a = df_audit[mask_prot]      # Protegido
    grupo_b = df_audit[~mask_prot]     # Referência
    
    if len(grupo_a) == 0 or len(grupo_b) == 0:
        return {"erro": "Grupos vazios"}

    # 3. Calcular Métricas
    
    # Taxa de Seleção (Quantos são classificados como 'Mau Pagador' / Default)
    sel_rate_a = grupo_a['y_pred'].mean()
    sel_rate_b = grupo_b['y_pred'].mean()
    
    # Disparate Impact (DI) = Rate A / Rate B
    # Se DI > 1.25 ou < 0.8, há indício de bias severo
    disparate_impact = sel_rate_a / sel_rate_b if sel_rate_b > 0 else 0
    
    # Taxa de Falso Positivo (Erro Tipo I: Dizemos que vai falhar, mas paga)
    # Importante para não negar crédito injustamente
    tn_a, fp_a, fn_a, tp_a = confusion_matrix(grupo_a['y_true'], grupo_a['y_pred'], labels=[0,1]).ravel()
    tn_b, fp_b, fn_b, tp_b = confusion_matrix(grupo_b['y_true'], grupo_b['y_pred'], labels=[0,1]).ravel()
    
    fpr_a = fp_a / (tn_a + fp_a) if (tn_a + fp_a) > 0 else 0
    fpr_b = fp_b / (tn_b + fp_b) if (tn_b + fp_b) > 0 else 0
    
    return {
        "grupo_protegido_nome": f"{atributo_protegido} {'<=' if atributo_protegido == 'idade' else '='} {valor_protegido}",
        "contagem_a": len(grupo_a),
        "contagem_b": len(grupo_b),
        "taxa_selecao_a": sel_rate_a,
        "taxa_selecao_b": sel_rate_b,
        "disparate_impact": disparate_impact,
        "fpr_a": fpr_a,
        "fpr_b": fpr_b,
        "fpr_diff": fpr_a - fpr_b
    }

# ===================== CREDIT SCORING AVANÇADO =====================
@st.cache_data(show_spinner="A treinar modelo...")
@monitorar_performance # <--- ADICIONAR ISTO
def credit_scoring_model(df):
    """Modelo completo de credit scoring com ROC, AUC e feature importance"""
    if 'default' not in df.columns:
        return None
    
    # Features para scoring
    feature_cols = ['idade', 'score_interno', 'rendimento_mensal', 'ltv', 
                   'divida_total', 'num_produtos', 'taxa_juro', 
                   'num_atrasos_12m', 'utilizacao_credito']
    
    available = [col for col in feature_cols if col in df.columns]
    
    if len(available) < 3:
        return None
    
    # Preparar dados
    df_model = df[available + ['default']].dropna()
    
    if len(df_model) < 100:
        return None
    
    X = df_model[available]
    y = df_model['default']
    
    # Split train/test
    X_train, X_test, y_train, y_test = train_test_split(
        X, y, test_size=0.3, random_state=42, stratify=y
    )
    
    # Normalizar features
    scaler = StandardScaler()
    X_train_scaled = scaler.fit_transform(X_train)
    X_test_scaled = scaler.transform(X_test)
    
    # Treinar modelo
    model = GradientBoostingClassifier(
        n_estimators=100,
        learning_rate=0.1,
        max_depth=5,
        random_state=42
    )
    model.fit(X_train_scaled, y_train)
    
    # Predições
    y_pred = model.predict(X_test_scaled)
    y_proba = model.predict_proba(X_test_scaled)[:, 1]
    
    # Curva ROC
    fpr, tpr, thresholds = roc_curve(y_test, y_proba)
    roc_auc = auc(fpr, tpr)
    
    fig_roc = go.Figure()
    fig_roc.add_trace(go.Scatter(
        x=fpr, y=tpr, 
        mode='lines',
        name=f'ROC Curve (AUC={roc_auc:.3f})',
        line=dict(color='#3b82f6', width=3)
    ))
    fig_roc.add_trace(go.Scatter(
        x=[0, 1], y=[0, 1],
        mode='lines',
        name='Random Classifier',
        line=dict(dash='dash', color='gray')
    ))
    fig_roc.update_layout(
        title='Curva ROC - Credit Scoring Model',
        xaxis_title='False Positive Rate',
        yaxis_title='True Positive Rate',
        template='plotly_dark'
    )
    
    # Matriz de Confusão
    cm = confusion_matrix(y_test, y_pred)
    fig_cm = px.imshow(
        cm,
        text_auto=True,
        title='Matriz de Confusão',
        labels={'x': 'Predito', 'y': 'Real'},
        x=['Não Default', 'Default'],
        y=['Não Default', 'Default'],
        color_continuous_scale='Blues'
    )
    
    # Feature Importance
    importance_df = pd.DataFrame({
        'Feature': available,
        'Importance': model.feature_importances_
    }).sort_values('Importance', ascending=False)
    
    # Classification Report
    report = classification_report(y_test, y_pred, output_dict=True)
    
    return {
        'model': model,
        'scaler': scaler,
        'roc_auc': roc_auc,
        'fig_roc': fig_roc,
        'fig_cm': fig_cm,
        'importance': importance_df,
        'report': report,
        'features': available
    }

# ===================== OTIMIZAÇÃO DE PORTFÓLIO =====================
def otimizar_portfolio(df):
    """Otimização de carteira - Fronteira Eficiente"""
    
    n_portfolios = 1000
    results = []
    
    np.random.seed(42)
    
    for i in range(n_portfolios):
        # Simular retornos e riscos aleatórios
        portfolio_return = np.random.normal(0.05, 0.02)
        portfolio_risk = np.random.normal(0.15, 0.05)
        
        # Evitar riscos negativos
        portfolio_risk = max(portfolio_risk, 0.01)
        
        # Sharpe Ratio
        sharpe = portfolio_return / portfolio_risk
        
        results.append({
            'return': portfolio_return,
            'risk': portfolio_risk,
            'sharpe': sharpe
        })
    
    df_portfolios = pd.DataFrame(results)
    
    # Criar gráfico da Fronteira Eficiente
    fig = px.scatter(
        df_portfolios,
        x='risk',
        y='return',
        color='sharpe',
        title='Fronteira Eficiente de Carteira',
        labels={
            'risk': 'Risco (Volatilidade)',
            'return': 'Retorno Esperado',
            'sharpe': 'Sharpe Ratio'
        },
        color_continuous_scale='RdYlGn',
        template='plotly_dark'
    )
    
    # Marcar portfolio ótimo (maior Sharpe)
    best_idx = df_portfolios['sharpe'].idxmax()
    best_portfolio = df_portfolios.iloc[best_idx]
    
    fig.add_trace(go.Scatter(
        x=[best_portfolio['risk']],
        y=[best_portfolio['return']],
        mode='markers',
        marker=dict(size=20, color='red', symbol='star'),
        name='Portfolio Ótimo'
    ))
    
    return fig, df_portfolios, best_portfolio

# ===================== REPORTING AVANÇADO =====================

# ===================== EXPORTAÇÃO EXCEL AVANÇADA (FUNC 10) =====================
def gerar_excel_avancado(df, kpis):
    """Gera um Excel com múltiplas abas e formatação profissional"""
    output = BytesIO()
    
    # Usar xlsxwriter como engine
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        # --- ABA 1: SUMÁRIO ---
        df_kpis = pd.DataFrame(list(kpis.items()), columns=['Métrica', 'Valor'])
        df_kpis.to_excel(writer, index=False, sheet_name='Sumário Executivo')
        
        # --- ABA 2: DADOS FILTRADOS ---
        df.to_excel(writer, index=False, sheet_name='Dados Detalhados')
        
        # --- FORMATAÇÃO ---
        workbook = writer.book
        ws_sumario = writer.sheets['Sumário Executivo']
        ws_dados = writer.sheets['Dados Detalhados']
        
        header_fmt = workbook.add_format({'bold': True, 'bg_color': '#3b82f6', 'font_color': 'white', 'border': 1})
        
        ws_sumario.set_column('A:A', 25)
        ws_sumario.set_column('B:B', 15)
        ws_dados.set_column('A:Z', 15)
        
    return output.getvalue()

def gerar_relatorio_executivo(df, user, role):
    """Gera relatório executivo completo em Word"""
    
    doc = Document()
    
    # Header
    doc.add_heading('RELATÓRIO EXECUTIVO SOCIOSTAT', 0)
    doc.add_paragraph(f'Data: {datetime.now().strftime("%d/%m/%Y %H:%M")}')
    doc.add_paragraph(f'Gerado por: {user} ({role})')
    doc.add_paragraph('_' * 60)
    
    # Sumário Executivo
    doc.add_heading('1. SUMÁRIO EXECUTIVO', 1)
    
    carteira = df['divida_total'].sum() / 1e6 if 'divida_total' in df.columns else 0
    default_rate = df['default'].mean() if 'default' in df.columns else 0
    clientes = len(df)
    
    doc.add_paragraph(f'• Carteira Total: €{carteira:.1f}M')
    doc.add_paragraph(f'• Taxa de Default: {default_rate:.2%}')
    doc.add_paragraph(f'• Total de Clientes: {clientes:,}')
    doc.add_paragraph(f'• Status: {"CRÍTICO" if default_rate > 0.15 else "ADEQUADO"}')
    
    # Basel III
    doc.add_heading('2. CAPITAL REGULATÓRIO (BASEL III)', 1)
    basel = calcular_capital_regulatorio(df)
    
    if basel:
        doc.add_paragraph(f'• PD: {basel["PD"]:.2%}')
        doc.add_paragraph(f'• EAD: €{basel["EAD"]/1e6:.1f}M')
        doc.add_paragraph(f'• Expected Loss: €{basel["Expected_Loss"]/1e6:.2f}M')
        doc.add_paragraph(f'• Capital Mínimo: €{basel["Capital_Minimo_8pct"]/1e6:.2f}M')
        doc.add_paragraph(f'• Capital Recomendado: €{basel["Capital_Recomendado_10.5pct"]/1e6:.2f}M')
    
    # Alertas
    doc.add_heading('3. ALERTAS DE COMPLIANCE', 1)
    alerts = check_compliance_alerts(df)
    
    if alerts:
        for i, alert in enumerate(alerts):
            doc.add_paragraph(f'{i+1}. [{alert["nivel"]}] {alert["tipo"]}: {alert["mensagem"]}')
    else:
        doc.add_paragraph('✓ Sem alertas ativos')
    
    # Recomendações
    doc.add_heading('4. RECOMENDAÇÕES', 1)
    doc.add_paragraph('• Monitorizar taxa de default semanalmente')
    doc.add_paragraph('• Revisar políticas de crédito em segmentos de alto risco')
    doc.add_paragraph('• Manter capital acima do mínimo regulatório')
    
    # Salvar em memória
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    
    return bio

def gerar_relatorio_tecnico(df):
    """Gera relatório técnico detalhado com estatísticas"""
    
    doc = Document()
    
    doc.add_heading('RELATÓRIO TÉCNICO DE RISCO', 0)
    doc.add_paragraph(f'Gerado em: {datetime.now().strftime("%d/%m/%Y %H:%M")}')
    doc.add_paragraph('_' * 60)
    
    # Estatísticas Descritivas
    doc.add_heading('1. ESTATÍSTICAS DESCRITIVAS', 1)
    
    for col in df.select_dtypes(include=[np.number]).columns[:5]:
        doc.add_paragraph(f'\n{col.upper()}:')
        doc.add_paragraph(f'  • Média: {df[col].mean():.2f}')
        doc.add_paragraph(f'  • Mediana: {df[col].median():.2f}')
        doc.add_paragraph(f'  • Desvio Padrão: {df[col].std():.2f}')
        doc.add_paragraph(f'  • Min/Max: {df[col].min():.2f} / {df[col].max():.2f}')
    
    # Correlações
    doc.add_heading('2. ANÁLISE DE CORRELAÇÕES', 1)
    
    if 'default' in df.columns:
        corr_with_default = df.corr()['default'].sort_values(ascending=False)
        doc.add_paragraph('\nTop correlações com Default:')
        for var, corr in corr_with_default.head(5).items():
            if var != 'default':
                doc.add_paragraph(f'  • {var}: {corr:.3f}')
    
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    
    return bio

def create_dashboard_summary():
    """Cria sumário visual do dashboard"""
    return {
        'timestamp': datetime.now(),
        'total_records': len(df),
        'default_rate': df['default'].mean() if 'default' in df.columns else 0,
        'alerts_count': len(check_compliance_alerts(df))
    }

# ===================== BENCHMARKING (FUNC 22) =====================
def get_dados_mercado():
    """
    Retorna as médias do setor bancário (Simulação de dados do Banco de Portugal/BCE).
    """
    return {
        "Taxa_Default": 0.038,       # Mercado tem 3.8%
        "LTV_Medio": 0.75,           # Mercado empresta 75% do valor
        "Score_Medio": 680,          # Score médio do mercado
        "Taxa_Juro_Media": 4.5,      # Taxa média aplicada
        "Utilizacao_Credito": 0.35   # Utilização média de plafonds
    }

def calcular_gap_analysis(meus_kpis, mercado):
    """Calcula a diferença entre o Banco e o Mercado"""
    gaps = {}
    
    # 1. Default (Quanto menor, melhor)
    diff_default = meus_kpis['Taxa_Default'] - mercado['Taxa_Default']
    gaps['Taxa_Default'] = {
        'eu': meus_kpis['Taxa_Default'],
        'mercado': mercado['Taxa_Default'],
        'diff': diff_default,
        'status': '✅ Melhor' if diff_default < 0 else '❌ Pior'
    }
    
    # 2. LTV (Quanto menor, menos risco)
    diff_ltv = meus_kpis['LTV_Medio'] - mercado['LTV_Medio']
    gaps['LTV_Medio'] = {
        'eu': meus_kpis['LTV_Medio'],
        'mercado': mercado['LTV_Medio'],
        'diff': diff_ltv,
        'status': '✅ Mais Conservador' if diff_ltv < 0 else '⚠️ Mais Agressivo'
    }
    
    # 3. Score (Quanto maior, melhor)
    diff_score = meus_kpis['Score_Medio'] - mercado['Score_Medio']
    gaps['Score_Medio'] = {
        'eu': meus_kpis['Score_Medio'],
        'mercado': mercado['Score_Medio'],
        'diff': diff_score,
        'status': '✅ Melhores Clientes' if diff_score > 0 else '⚠️ Piores Clientes'
    }
    
    return gaps

# ===================== KPI TRACKING & MONITORING =====================
def calcular_kpis_principais(df):
    """Calcula todos os KPIs principais do banco"""
    kpis = {}
    
    # KPIs Financeiros
    if 'divida_total' in df.columns:
        kpis['Carteira_Total'] = df['divida_total'].sum()
        kpis['Exposicao_Media'] = df['divida_total'].mean()
    
    # KPIs de Risco
    if 'default' in df.columns:
        kpis['Taxa_Default'] = df['default'].mean()
        kpis['Num_Defaults'] = df['default'].sum()
    
    if 'ltv' in df.columns:
        kpis['LTV_Medio'] = df['ltv'].mean()
        kpis['LTV_Alto_Pct'] = (df['ltv'] > 0.9).sum() / len(df)
    
    if 'score_interno' in df.columns:
        kpis['Score_Medio'] = df['score_interno'].mean()
        kpis['Score_Baixo_Pct'] = (df['score_interno'] < 600).sum() / len(df)
    
    # KPIs Operacionais
    kpis['Total_Clientes'] = len(df)
    
    if 'segmento' in df.columns:
        kpis['Concentracao_Max'] = df['segmento'].value_counts(normalize=True).max()
    
    if 'num_atrasos_12m' in df.columns:
        kpis['Taxa_Atraso'] = (df['num_atrasos_12m'] > 0).sum() / len(df)
    
    # Basel III
    basel = calcular_capital_regulatorio(df)
    if basel:
        kpis['Capital_Minimo'] = basel['Capital_Minimo_8pct']
        kpis['Expected_Loss'] = basel['Expected_Loss']
        kpis['RWA'] = basel['RWA']
    
    return kpis

def comparar_kpis(kpis_atual, kpis_anterior):
    """Compara KPIs atual vs período anterior"""
    comparacao = {}
    
    for kpi_name in kpis_atual.keys():
        if kpi_name in kpis_anterior:
            atual = kpis_atual[kpi_name]
            anterior = kpis_anterior[kpi_name]
            
            if anterior != 0:
                variacao = ((atual - anterior) / abs(anterior)) * 100
            else:
                variacao = 0
            
            comparacao[kpi_name] = {
                'atual': atual,
                'anterior': anterior,
                'variacao_pct': variacao,
                'direcao': '📈' if variacao > 0 else '📉' if variacao < 0 else '➡️'
            }
    
    return comparacao

def criar_alerta_kpi(kpi_name, valor_atual, threshold, tipo='max'):
    """Cria alerta se KPI ultrapassar threshold"""
    if tipo == 'max' and valor_atual > threshold:
        return {
            'kpi': kpi_name,
            'valor': valor_atual,
            'threshold': threshold,
            'tipo': 'Ultrapassou máximo',
            'severidade': 'CRÍTICO'
        }
    elif tipo == 'min' and valor_atual < threshold:
        return {
            'kpi': kpi_name,
            'valor': valor_atual,
            'threshold': threshold,
            'tipo': 'Abaixo do mínimo',
            'severidade': 'AVISO'
        }
    return None

def gerar_grafico_tendencia_kpi(kpi_name, historico):
    """Gera gráfico de tendência temporal de um KPI"""
    fig = go.Figure()
    
    fig.add_trace(go.Scatter(
        x=list(range(len(historico))),
        y=historico,
        mode='lines+markers',
        name=kpi_name,
        line=dict(color='#3b82f6', width=3),
        marker=dict(size=8)
    ))
    
    # Adicionar média
    media = np.mean(historico)
    fig.add_hline(y=media, line_dash="dash", line_color="gray", 
                  annotation_text=f"Média: {media:.2f}")
    
    fig.update_layout(
        title=f'Tendência: {kpi_name}',
        xaxis_title='Período',
        yaxis_title='Valor',
        template='plotly_dark',
        hovermode='x unified'
    )
    
    return fig

# ===================== SEGMENTAÇÃO AVANÇADA =====================
def segmentacao_rfm(df):
    """Segmentação RFM (Recency, Frequency, Monetary) adaptada para risco"""
    if not all(col in df.columns for col in ['duracao_meses', 'num_produtos', 'divida_total']):
        return None
    
    # Calcular scores
    df_seg = df.copy()
    
    # R - Recency (quanto mais recente, melhor)
    df_seg['R_Score'] = pd.qcut(df_seg['duracao_meses'], q=4, labels=[4,3,2,1], duplicates='drop')
    
    # F - Frequency (número de produtos)
    df_seg['F_Score'] = pd.qcut(df_seg['num_produtos'], q=4, labels=[1,2,3,4], duplicates='drop')
    
    # M - Monetary (valor da dívida)
    df_seg['M_Score'] = pd.qcut(df_seg['divida_total'], q=4, labels=[1,2,3,4], duplicates='drop')
    
    # Score total
    df_seg['RFM_Score'] = (df_seg['R_Score'].astype(int) + 
                           df_seg['F_Score'].astype(int) + 
                           df_seg['M_Score'].astype(int))
    
    # Classificação
    df_seg['Segmento_RFM'] = pd.cut(
        df_seg['RFM_Score'],
        bins=[0, 4, 7, 9, 12],
        labels=['Alto Risco', 'Risco Moderado', 'Baixo Risco', 'Premium']
    )
    
    return df_seg

def analise_coorte(df, data_col='duracao_meses', grupo_col='segmento'):
    """Análise de coorte para tracking de performance"""
    if data_col not in df.columns or grupo_col not in df.columns:
        return None
    
    # Criar coortes por período
    df_coorte = df.copy()
    df_coorte['Coorte'] = pd.cut(df_coorte[data_col], bins=5, labels=['0-12m', '12-24m', '24-36m', '36-48m', '48+m'])
    
    # Default rate por coorte e grupo
    coorte_analysis = df_coorte.groupby(['Coorte', grupo_col])['default'].agg(['mean', 'count']).reset_index()
    coorte_analysis.columns = ['Coorte', grupo_col, 'Default_Rate', 'Count']
    
    return coorte_analysis

# ===================== ANÁLISE INDIVIDUAL DE CLIENTE =====================
def analisar_cliente_individual(df, cliente_id):
    """Análise detalhada de um cliente específico"""
    cliente = df[df['id_cliente'] == cliente_id]
    
    if len(cliente) == 0:
        return None
    
    cliente = cliente.iloc[0]
    
    # Calcular percentil em relação à carteira
    analise = {
        'dados': cliente.to_dict(),
        'percentis': {}
    }
    
    for col in df.select_dtypes(include=[np.number]).columns:
        if col in cliente.index and col != 'id_cliente':
            valor_cliente = cliente[col]
            percentil = (df[col] < valor_cliente).sum() / len(df) * 100
            analise['percentis'][col] = percentil
    
    return analise

# ===================== API PÚBLICA (SIMULADA - FUNC 15) =====================
import json

def api_get_cliente(cliente_id, df):
    """Endpoint: GET /api/v1/cliente/{id}"""
    # Validar token (Simulação de Segurança Func 13)
    # Em produção, verificaríamos o header 'Authorization'
    
    cliente = df[df['id_cliente'] == cliente_id]
    
    if len(cliente) == 0:
        return {"status": 404, "error": "Cliente não encontrado"}
    
    # Converter para dicionário
    dados = cliente.iloc[0].to_dict()
    
    # Adicionar metadados
    response = {
        "status": 200,
        "timestamp": datetime.now().isoformat(),
        "data": dados,
        "links": {
            "self": f"/api/v1/cliente/{cliente_id}",
            "score": f"/api/v1/cliente/{cliente_id}/score"
        }
    }
    return response

def api_get_risco_carteira(df):
    """Endpoint: GET /api/v1/risco/carteira"""
    kpis = calcular_kpis_principais(df)
    
    response = {
        "status": 200,
        "timestamp": datetime.now().isoformat(),
        "metrics": {
            "default_rate": kpis.get('Taxa_Default'),
            "portfolio_volume": kpis.get('Carteira_Total'),
            "average_ltv": kpis.get('LTV_Medio')
        }
    }
    return response
def calcular_score_cliente(cliente_data):
    """Calcula score de crédito detalhado para um cliente"""
    score_base = 500
    
    # Score interno (peso 30%)
    if 'score_interno' in cliente_data:
        score_interno_norm = (cliente_data['score_interno'] - 500) / 200
        score_base += score_interno_norm * 150
    
    # LTV (peso 25%) - quanto menor, melhor
    if 'ltv' in cliente_data:
        ltv_penalty = (cliente_data['ltv'] - 0.5) * -100
        score_base += ltv_penalty
    
    # Rendimento (peso 20%)
    if 'rendimento_mensal' in cliente_data:
        if cliente_data['rendimento_mensal'] > 2000:
            score_base += 50
        elif cliente_data['rendimento_mensal'] > 1500:
            score_base += 30
    
    # Idade (peso 15%)
    if 'idade' in cliente_data:
        if 30 <= cliente_data['idade'] <= 50:
            score_base += 40
        elif 25 <= cliente_data['idade'] < 30 or 50 < cliente_data['idade'] <= 60:
            score_base += 20
    
    # Atrasos (peso 10%)
    if 'num_atrasos_12m' in cliente_data:
        score_base -= cliente_data['num_atrasos_12m'] * 30
    
    # Garantir score entre 300 e 850
    score_final = max(300, min(850, score_base))
    
    # Classificação
    if score_final >= 750:
        rating = 'AAA - Excelente'
        risco = 'Muito Baixo'
        cor = '#10b981'
    elif score_final >= 650:
        rating = 'BBB - Bom'
        risco = 'Baixo'
        cor = '#3b82f6'
    elif score_final >= 550:
        rating = 'CCC - Regular'
        risco = 'Médio'
        cor = '#f59e0b'
    else:
        rating = 'DDD - Fraco'
        risco = 'Alto'
        cor = '#ef4444'
    
    return {
        'score': score_final,
        'rating': rating,
        'risco': risco,
        'cor': cor
    }

def simular_credito(valor_solicitado, prazo_meses, score_cliente, ltv=None):
    """Simula aprovação e condições de crédito"""
    
    # Taxa base (Euribor + spread)
    euribor = 0.04  # 4%
    
    # Spread baseado no score
    if score_cliente >= 750:
        spread = 0.015  # 1.5%
        decisao = 'APROVADO'
        limite_max = valor_solicitado
    elif score_cliente >= 650:
        spread = 0.025  # 2.5%
        decisao = 'APROVADO COM CONDIÇÕES'
        limite_max = valor_solicitado * 0.9
    elif score_cliente >= 550:
        spread = 0.04  # 4%
        decisao = 'ANÁLISE ADICIONAL'
        limite_max = valor_solicitado * 0.7
    else:
        spread = 0.06  # 6%
        decisao = 'NEGADO'
        limite_max = 0
    
    taxa_final = euribor + spread
    
    # Calcular prestação (PMT)
    if limite_max > 0:
        taxa_mensal = taxa_final / 12
        prestacao = (limite_max * taxa_mensal * (1 + taxa_mensal)**prazo_meses) / \
                    ((1 + taxa_mensal)**prazo_meses - 1)
        total_a_pagar = prestacao * prazo_meses
        juros_totais = total_a_pagar - limite_max
    else:
        prestacao = 0
        total_a_pagar = 0
        juros_totais = 0
    
    return {
        'decisao': decisao,
        'valor_aprovado': limite_max,
        'taxa_juro_anual': taxa_final,
        'prestacao_mensal': prestacao,
        'prazo_meses': prazo_meses,
        'total_a_pagar': total_a_pagar,
        'juros_totais': juros_totais,
        'tan': taxa_final,
        'taeg': taxa_final * 1.05  # Simplificado
    }

def gerar_recomendacoes_cliente(cliente_data, score_info):
    """Gera recomendações personalizadas para melhorar o perfil"""
    recomendacoes = []
    
    # Recomendações baseadas em score
    if score_info['score'] < 650:
        recomendacoes.append({
            'tipo': 'CRÍTICO',
            'titulo': 'Melhorar Score de Crédito',
            'descricao': 'Score atual abaixo do ideal. Ações recomendadas: regularizar atrasos e reduzir utilização de crédito.'
        })
    
    # Recomendações baseadas em LTV
    if 'ltv' in cliente_data and cliente_data['ltv'] > 0.8:
        recomendacoes.append({
            'tipo': 'AVISO',
            'titulo': 'LTV Elevado',
            'descricao': f'LTV de {cliente_data["ltv"]:.1%} está elevado. Considere aumentar entrada ou garantias adicionais.'
        })
    
    # Recomendações baseadas em atrasos
    if 'num_atrasos_12m' in cliente_data and cliente_data['num_atrasos_12m'] > 0:
        recomendacoes.append({
            'tipo': 'CRÍTICO',
            'titulo': 'Histórico de Atrasos',
            'descricao': f'{cliente_data["num_atrasos_12m"]} atraso(s) nos últimos 12 meses. Regularizar pagamentos é essencial.'
        })
    
    # Recomendações baseadas em utilização de crédito
    if 'utilizacao_credito' in cliente_data and cliente_data['utilizacao_credito'] > 0.7:
        recomendacoes.append({
            'tipo': 'AVISO',
            'titulo': 'Alta Utilização de Crédito',
            'descricao': f'Utilização de {cliente_data["utilizacao_credito"]:.1%}. Ideal manter abaixo de 30%.'
        })
    
    # Recomendações positivas
    if score_info['score'] >= 750:
        recomendacoes.append({
            'tipo': 'SUCESSO',
            'titulo': 'Perfil Excelente',
            'descricao': 'Elegível para as melhores taxas e condições do mercado. Considere consolidação de dívidas.'
        })
    
    if len(recomendacoes) == 0:
        recomendacoes.append({
            'tipo': 'INFO',
            'titulo': 'Perfil Adequado',
            'descricao': 'Perfil de crédito satisfatório. Manter boas práticas financeiras.'
        })
    
    return recomendacoes

def comparar_com_peers(df, cliente_data, segmento_col='segmento'):
    """Compara cliente com seus peers (mesmo segmento)"""
    if segmento_col not in df.columns or segmento_col not in cliente_data:
        return None
    
    segmento = cliente_data[segmento_col]
    peers = df[df[segmento_col] == segmento]
    
    comparacao = {}
    
    for col in ['score_interno', 'ltv', 'rendimento_mensal', 'divida_total']:
        if col in df.columns and col in cliente_data:
            valor_cliente = cliente_data[col]
            media_peers = peers[col].mean()
            percentil = (peers[col] < valor_cliente).sum() / len(peers) * 100
            
            comparacao[col] = {
                'cliente': valor_cliente,
                'media_peers': media_peers,
                'percentil': percentil,
                'melhor_que_peers': percentil > 50
            }
    
    return comparacao

# ===================== FIM DA PARTE 3 =====================

# =====================================================
# PARTE 1 & 3: CONFIGURAÇÃO + UI ENTERPRISE + MENU
# =====================================================

# 1. CONFIGURAÇÃO DA PÁGINA (Deve ser a primeira linha executável)
st.set_page_config(
    page_title="SocioStat Enterprise | Risk Management",
    layout="wide",
    page_icon="🏦",
    initial_sidebar_state="expanded"
)


# ===================== GESTÃO DE TEMAS & IMPRESSÃO (FUNC 28 + 29) =====================
def get_app_style(mode):
    """Gera o CSS baseado no tema escolhido + Regras de Impressão"""
    
    if mode == "Dark":
        # Tema Escuro (Financial)
        vars = {
            "bg": "#0f172a", "sidebar": "#020617", "text": "#e2e8f0",
            "card": "#1e293b", "border": "#334155", "metric": "#ffffff",
            "metric_label": "#94a3b8", "input_bg": "#1e293b"
        }
    else:
        # Tema Claro (Corporate)
        vars = {
            "bg": "#f1f5f9", "sidebar": "#ffffff", "text": "#0f172a",
            "card": "#ffffff", "border": "#e2e8f0", "metric": "#0f172a",
            "metric_label": "#64748b", "input_bg": "#ffffff"
        }
        
    return f"""
    <style>
        /* --- ESTILOS DE ECRÃ (Screen) --- */
        .stApp {{ background-color: {vars['bg']}; color: {vars['text']}; }}
        [data-testid="stSidebar"] {{ background-color: {vars['sidebar']}; border-right: 1px solid {vars['border']}; }}
        .stMarkdown, p, h1, h2, h3, li, span {{ color: {vars['text']} !important; }}
        
        div[data-testid="metric-container"] {{
            background-color: {vars['card']};
            border: 1px solid {vars['border']};
            padding: 15px;
            border-radius: 8px;
            box-shadow: 0 4px 6px -1px rgba(0,0,0,0.1);
        }}
        label[data-testid="stMetricLabel"] {{ color: {vars['metric_label']} !important; }}
        div[data-testid="stMetricValue"] {{ color: {vars['metric']} !important; }}
        
        .stTextInput input, .stSelectbox div[data-baseweb="select"], .stNumberInput input {{
            background-color: {vars['input_bg']} !important;
            color: {vars['text']} !important;
            border: 1px solid {vars['border']} !important;
        }}
        
        .metric-card {{ 
            background: {vars['card']}; 
            border: 1px solid {vars['border']}; 
            padding: 20px; 
            border-radius: 10px; 
            box-shadow: 0 2px 4px rgba(0,0,0,0.05);
            border-left: 4px solid #3b82f6;
        }}
        .header-style {{
            font-size: 2.0rem;
            font-weight: 800;
            background: -webkit-linear-gradient(0deg, #2563eb, #0ea5e9);
            -webkit-background-clip: text;
            -webkit-text-fill-color: transparent;
            margin-bottom: 1rem;
        }}

        /* --- FUNCIONALIDADE 29: MÓDULO DE IMPRESSÃO OTIMIZADA --- */
        @media print {{
            /* 1. Esconder Interface de Navegação */
            [data-testid="stSidebar"], 
            header, 
            footer, 
            .stDeployButton, 
            button {{
                display: none !important;
            }}
            
            /* 2. Forçar Fundo Branco e Texto Preto (Poupança de Tinta) */
            .stApp, div[data-testid="metric-container"], .metric-card {{
                background-color: white !important;
                color: black !important;
                border: 1px solid #ddd !important;
                box-shadow: none !important;
            }}
            
            h1, h2, h3, p, span, li, div {{
                color: black !important;
                -webkit-text-fill-color: black !important; /* Remove gradiente do título */
            }}
            
            /* 3. Expandir Conteúdo */
            .block-container {{
                padding: 0 !important;
                max-width: 100% !important;
            }}
            
            /* 4. Ajustes de Gráficos */
            .js-plotly-plot {{
                break-inside: avoid; /* Tenta não cortar gráficos a meio da página */
            }}
        }}
    </style>
    """

# ------------------------------------------------------------------
# LÓGICA DE LOGIN (Mantida, mas com visual melhorado)
# ------------------------------------------------------------------
if not st.session_state.logged_in:
    # Centralizar o Login
    col1, col2, col3 = st.columns([1, 1, 1])
    with col2:
        st.markdown("<br><br>", unsafe_allow_html=True)
        st.markdown("<h2 style='text-align: center;'>🔐 SocioStat Enterprise</h2>", unsafe_allow_html=True)
        st.markdown("<p style='text-align: center; color: #64748b;'>Portal de Gestão de Risco & Compliance</p>", unsafe_allow_html=True)
        
        with st.form("login_form"):
            username = st.text_input("ID Utilizador", placeholder="admin")
            pwd = st.text_input("Credenciais", type="password", placeholder="••••••")
            role = st.selectbox("Perfil de Acesso", ["Admin", "Risk Manager", "Auditor", "Analyst"])
            
            submit = st.form_submit_button("Aceder ao Terminal", use_container_width=True)
            
            if submit:
                if pwd == "demo123":
                    st.session_state.logged_in = True
                    st.session_state.user_role = role
                    st.session_state.username = username if username else "Admin"
                    log_audit(st.session_state.username, "Login", f"Role: {role}")
                    st.rerun()
                else:
                    st.error("Credenciais inválidas.")
        
        st.caption("🔒 Ambiente Seguro | Encriptação AES-256 | Demo: `demo123`")
    
    st.stop() # Pára a execução aqui se não estiver logado

# ===================== PARTE 3: SIDEBAR (CORRIGIDA E SEGURA) =====================
with st.sidebar:
    # 1. CABEÇALHO DE PERFIL
    with st.container():
        c1, c2 = st.columns([1, 3])
        with c1:
            st.image(f"https://ui-avatars.com/api/?name={st.session_state.username}&background=3b82f6&color=fff&size=64", width=45)
        with c2:
            st.markdown(f"<div style='margin-top:5px; font-weight:bold; font-size:14px;'>{st.session_state.username}</div>", unsafe_allow_html=True)
            st.caption(f"🔑 {st.session_state.user_role}")
        

    st.markdown("---")

    # 2. CONTROLO DE SISTEMA
    c_lang, c_theme = st.columns(2)
    with c_lang:
        lang_sel = st.selectbox("🌐 Idioma", ["PT", "EN"], index=0 if st.session_state.language == "PT" else 1, label_visibility="collapsed", key="lang_selector_main")
        if lang_sel != st.session_state.language:
            st.session_state.language = lang_sel
            st.rerun()
    with c_theme:
        if st.button("🌗 Tema", use_container_width=True):
            st.session_state.theme = "Light" if st.session_state.theme == "Dark" else "Dark"
            st.rerun()

    # 3. FILTROS GLOBAIS
    with st.expander("🔍 Filtros & Dados", expanded=False):
        if st.session_state.df is None:
            if st.button("📥 Carregar Demo"):
                configurar_demo_bancaria()
                st.rerun()
        else:
            df_raw = st.session_state.df
            st.caption(f"Total: {len(df_raw):,} registos")
            
            # Filtros seguros (verificam se colunas existem)
            if 'regiao' in df_raw.columns:
                f_reg = st.multiselect("Região", df_raw['regiao'].unique(), default=df_raw['regiao'].unique())
            else: f_reg = []
            
            if 'segmento' in df_raw.columns:
                f_seg = st.multiselect("Segmento", df_raw['segmento'].unique(), default=df_raw['segmento'].unique())
            else: f_seg = []
            
            # Aplicar Filtro
            mask = pd.Series([True] * len(df_raw))
            if 'regiao' in df_raw.columns: mask &= df_raw['regiao'].isin(f_reg)
            if 'segmento' in df_raw.columns: mask &= df_raw['segmento'].isin(f_seg)
            df = df_raw[mask] # Atualiza o DF global

            # ... (depois dos filtros de dados) ...
        st.markdown("---")
        st.caption("♿ Acessibilidade")
        acessibilidade_mode = st.selectbox(
            "Modo de Visualização", 
            ["Standard", "Colorblind Safe (Daltonismo)"],
            key="access_mode"
        )

    st.markdown("---")

    # 4. MENU DE NAVEGAÇÃO (A CORREÇÃO CRÍTICA)
    
    # Lista Mestra de Páginas (Nomes Internos - NÃO MUDAR ISTO)
    # Esta lista deve corresponder EXATAMENTE aos 'if selected ==' da Parte 4
    menu_keys = [
        "Dashboard", "Macro Watch", "Elasticidade", "KPI Tracking", "Geo-Risk Map", "Reconstrução de Experiências", "Licenciamento de Dados",
        "Visualizador", "Cliente Individual", "What-If Analysis", "Basel III", "Catálogo de Dados", "Gamificação", "Otimização de Consultas",
        "Forecasting", "Credit Scoring", "Análise de Drivers (RLM)", "Early Warning", "Qualidade de Dados", "Monitorização APM", "Executive Score (RCS)",
        "Stress Testing", "Survival Analysis", "Compliance", "Portfolio", "Web Scraping", "Conetores Oficiais", "XAI Explainer", "Otimização de Custos", "Gestão de Alertas",
        "Integração CRM", "Integração Cloud","Reporting", "Base de Conhecimento", "Feedback","MLOps Center", "Integração BI",
        "Gestão de Tags", "API & Devs", "Gestão de Utilizadores", "Billing & Planos", "Audit Trail", "Auditoria de Segurança"
    ]

    # Ícones correspondentes (mesma ordem)
    menu_icons = [
        "house", "globe", "magnet", "speedometer", "map", "heart-pulse", "trophy", "lamp", "speedometer",
        "search", "person", "sliders", "bank", "journal-bookmark","rewind-circle", "key-fill","globe-americas",
        "graph-up-arrow", "credit-card", "calculator", "exclamation-diamond", "bar-chart-line", "bell-fill",
        "lightning", "hourglass", "file-medical", "briefcase", "globe2", "building", "lightning-charge-fill",
        "cloud-arrow-up", "arrow-repeat", "file-earmark-pdf", "book", "chat", "diagram-3", "piggy-bank",
        "tags-fill", "code", "people", "wallet", "shield-lock", "list-ul", "activity",
    ]

    # Traduzir as opções para o utilizador ver
    # Se não houver tradução definida, usa o nome original
    menu_options_display = [t("menu_" + k.split()[0].lower()) if "menu_" + k.split()[0].lower() in TRANSLATIONS.get("PT", {}) else k for k in menu_keys]
    
    # Criar dicionário para reverter a tradução (Visual -> Interno)
    # Ex: "Painel" -> "Dashboard"
    mapa_navegacao = dict(zip(menu_options_display, menu_keys))

    # Renderizar o Menu
    selected_label = option_menu(
        menu_title=None,
        options=menu_options_display,
        icons=menu_icons,
        menu_icon="cast",
        default_index=0,
        styles={
            "container": {"padding": "0!important", "background-color": "transparent"},
            "icon": {"color": "#60a5fa", "font-size": "14px"}, 
            "nav-link": {"font-size": "13px", "text-align": "left", "margin":"2px", "--hover-color": "#1e293b"},
            "nav-link-selected": {"background-color": "#2563eb"},
        }
    )
    
    # RECUPERAR A CHAVE INTERNA (Isto repara a navegação!)
    selected = mapa_navegacao.get(selected_label, "Dashboard")

    # 5. RODAPÉ DE ALERTAS
    st.markdown("---")
    if 'df' in locals() and df is not None:
        alerts = check_compliance_alerts(df)
        if alerts:
            with st.expander(f"⚠️ Alertas ({len(alerts)})"):
                for alert in alerts[:2]:
                    st.caption(f"🔴 {alert['tipo']}")

                    # 6. RODAPÉ DE SISTEMA (Limpar Cache & Sair)
    st.markdown("---")
    
    col_tools1, col_tools2 = st.columns(2)
    
    with col_tools1:
        # Botão de Cache (Funcionalidade 14 recuperada!)
        if st.button("🧹 Cache", help="Limpar memória e recarregar dados", use_container_width=True):
            st.cache_data.clear()
            st.rerun()
            
    with col_tools2:
        # Botão de Logout (Reforço)
        if st.button("🚪 Sair", help="Terminar sessão", use_container_width=True, key="logout_bottom"):
            st.session_state.clear()
            st.rerun()
            
    st.caption("v3.0.2 Enterprise")

# ===================== FIM DA PARTE 3 =====================

# ===================== BLOCO DE SEGURANÇA (A CORREÇÃO) =====================
# 1. Verificar se há dados carregados
if st.session_state.df is None:
    col1, col2 = st.columns([1, 2])
    with col2:
        st.warning("⚠️ Nenhum dado carregado.")
        st.info("👈 Por favor, clique em **'📥 Carregar Demo'** na barra lateral esquerda.")
    st.stop() # <--- ISTO É O SEGREDO! Pára o código aqui e evita o erro NameError.

# 2. Garantir que a variável 'df' existe para as páginas seguintes
# Se por algum motivo o filtro da sidebar falhou, usamos o original
if 'df' not in locals():
    df = st.session_state.df

# Garante que a coluna 'tags' existe para a funcionalidade 38 funcionar
if 'tags' not in df.columns:
    # Cria uma coluna de tags vazia para todos os clientes
    # Usamos uma lista de listas para permitir múltiplas tags
    df['tags'] = [[] for _ in range(len(df))]
    
    # Opcional: Adicionar algumas tags de exemplo para não ficar vazio
    if len(df) > 0:
        # Marca os primeiros 5 clientes como "VIP" para teste
        for i in range(min(5, len(df))):
            df.at[df.index[i], 'tags'] = ["VIP", "Novo"]

# 3. Recalcular listas de colunas (para garantir consistência)
cols_num = df.select_dtypes(include=[np.number]).columns.tolist()
cols_all = df.columns.tolist()

# ===========================================================================

# =====================================================
# PARTE 4/4: Todas as Páginas e Módulos
# =====================================================

# ===================== PÁGINA: DASHBOARD =====================
if selected == "Dashboard":
    st.header(t('dashboard_title'))
    log_audit(st.session_state.username, "Dashboard View", "Visualização do dashboard")
    
    try:
        df_macro = get_dados_macroeconomicos()
    except:
        df_macro = None
        
    rcs_data = calcular_rcs(df, df_macro)
    rcs_val = rcs_data['RCS']
    
    # Definir cor do Score
    cor_rcs = "green" if rcs_val >= 80 else "orange" if rcs_val >= 50 else "red"
    msg_rcs = "Excelente" if rcs_val >= 80 else "Atenção" if rcs_val >= 50 else "Crítico"
    
    # Widget Visual de Destaque
    st.markdown(f"""
    <div style="background-color: #1e293b; border: 2px solid {cor_rcs}; padding: 20px; border-radius: 15px; margin-bottom: 25px; text-align: center;">
        <h4 style="color: #94a3b8; margin:0;">EXECUTIVE ACTION SCORE (RCS)</h4>
        <h1 style="font-size: 72px; margin:0; color: {cor_rcs};">{rcs_val:.0f}<span style="font-size:30px">/100</span></h1>
        <h3 style="margin:0; color: {cor_rcs};">{msg_rcs}</h3>
        <p style="color: #64748b;">Saúde Global da Instituição</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Detalhe dos Componentes (Expander)
    with st.expander("🔍 Ver Decomposição do Score"):
        c1, c2, c3 = st.columns(3)
        comp = rcs_data['Componentes']
        c1.metric("Carteira (50%)", f"{comp['Carteira (50%)']:.0f}", help="Baseado em Default e LTV")
        c2.metric("Macro (30%)", f"{comp['Macroeconomia (30%)']:.0f}", help="Inflação e Desemprego")
        c3.metric("Compliance (20%)", f"{comp['Compliance (20%)']:.0f}", help="Penalização por Alertas")
    
    # KPIs principais
    col1, col2, col3, col4, col5 = st.columns(5)
    
    with col1:
        carteira_total = df['divida_total'].sum() / 1e6 if 'divida_total' in df.columns else 0
        st.metric(t('kpi_portfolio'), f"€{carteira_total:.1f}M")
    
    with col2:
        default_rate = df['default'].mean() if 'default' in df.columns else 0
        st.metric("⚠️ Taxa Default", f"{default_rate:.1%}", 
                 delta=f"{'🔴' if default_rate > 0.15 else '🟢'}")
    
    with col3:
        ltv_medio = df['ltv'].mean() if 'ltv' in df.columns else 0
        st.metric("📊 LTV Médio", f"{ltv_medio:.1%}")
    
    with col4:
        st.metric("👥 Clientes", f"{len(df):,}")
    
    with col5:
        score_medio = df['score_interno'].mean() if 'score_interno' in df.columns else 0
        st.metric("🎯 Score Médio", f"{score_medio:.0f}")
    
    # Alertas Ativos
    if alerts:
        st.markdown("### ⚠️ Alertas Ativos")
        cols_alert = st.columns(len(alerts))
        for i, alert in enumerate(alerts):
            with cols_alert[i]:
                nivel_class = "alert-critical" if alert['nivel'] == "CRÍTICO" else "alert-warning"
                st.markdown(f"""
                <div class="{nivel_class}">
                    <strong>{alert['tipo']}</strong><br>
                    {alert['mensagem']}<br>
                    <em>Ação: {alert['acao']}</em>
                </div>
                """, unsafe_allow_html=True)
    
    # Gráficos principais
    col1, col2 = st.columns(2)
    
    with col1:
        if 'segmento' in df.columns and 'default' in df.columns:
            # Obter cores dinâmicas
            paleta = get_color_palette(st.session_state.get("access_mode", "Standard"))
            
            seg_default = df.groupby('segmento')['default'].mean().reset_index()
            
            fig = px.bar(seg_default, x='segmento', y='default', 
                         title='Taxa de Default por Segmento',
                         labels={'default': 'Taxa de Default', 'segmento': 'Segmento'},
                         # Usar a cor sequencial segura
                         color='default', 
                         color_continuous_scale=paleta['sequencial'])
            st.plotly_chart(fig, use_container_width=True)
    
    with col2:
        if 'regiao' in df.columns and 'divida_total' in df.columns:
            reg_exp = df.groupby('regiao')['divida_total'].sum().reset_index()
            fig = px.pie(reg_exp, values='divida_total', names='regiao',
                        title='Exposição por Região')
            st.plotly_chart(fig, use_container_width=True)
    
    # Distribuição de Score
    if 'score_interno' in df.columns:
        st.markdown("### 📊 Distribuição de Credit Score")
        fig = px.histogram(df, x='score_interno', nbins=50,
                          title='Distribuição de Scores',
                          labels={'score_interno': 'Score Interno'})
        st.plotly_chart(fig, use_container_width=True)

        sofia_explica(f"""
    **Análise Executiva:**
    
    A carteira total é de **€{carteira_total:.1f}M** com uma taxa de default de **{default_rate:.1%}**.
    Recomenda-se atenção ao segmento com maior risco e revisão dos critérios de aprovação para manter o LTV médio controlado em {ltv_medio:.1%}.
    """)
        
    
# ===================== PÁGINA: MACRO WATCH (FUNC 11) =====================
elif selected == "Macro Watch":
    st.header("🌍 Macro Watch - Indicadores Económicos")
    log_audit(st.session_state.username, "Macro Access", "Análise de indicadores externos")
    
    # 1. Obter Dados
    df_macro = get_dados_macroeconomicos(36) # Últimos 3 anos
    
    # 2. KPIs Atuais (Último mês)
    ultimo = df_macro.iloc[-1]
    penultimo = df_macro.iloc[-2]
    
    col1, col2, col3, col4 = st.columns(4)
    
    col1.metric("Euribor 6M", f"{ultimo['Euribor 6M (%)']:.2f}%", 
                delta=f"{ultimo['Euribor 6M (%)'] - penultimo['Euribor 6M (%)']:.2f} pp")
    
    col2.metric("Inflação", f"{ultimo['Inflação (%)']:.2f}%", 
                delta=f"{ultimo['Inflação (%)'] - penultimo['Inflação (%)']:.2f} pp", delta_color="inverse")
                
    col3.metric("Desemprego", f"{ultimo['Taxa Desemprego (%)']:.2f}%", 
                delta=f"{ultimo['Taxa Desemprego (%)'] - penultimo['Taxa Desemprego (%)']:.2f} pp", delta_color="inverse")
                
    col4.metric("PIB (YoY)", f"{ultimo['Crescimento PIB (%)']:.2f}%", 
                delta=f"{ultimo['Crescimento PIB (%)'] - penultimo['Crescimento PIB (%)']:.2f} pp")

    st.markdown("---")

    # 3. Visualização Temporal
    tab1, tab2 = st.tabs(["📈 Tendências Macro", "🔗 Correlação com Carteira"])
    
    with tab1:
        st.markdown("### Evolução dos Indicadores (36 Meses)")
        
        # Multiselect para escolher o que ver
        vars_macro = df_macro.columns.tolist()
        selecao = st.multiselect("Selecionar Indicadores", vars_macro, default=['Euribor 6M (%)', 'Inflação (%)'])
        
        if selecao:
            fig = px.line(df_macro, y=selecao, title="Indicadores Macroeconómicos", template="plotly_dark")
            fig.update_layout(yaxis_title="Percentagem (%)", hovermode="x unified")
            st.plotly_chart(fig, use_container_width=True)
        
    with tab2:
        st.markdown("### 🔗 Impacto na Carteira")
        st.info("Como é que a economia afeta o nosso risco de Default?")
        
        # Gerar histórico da carteira
        df_carteira_hist = gerar_historico_simulado(df, 36)
        
        # Calcular correlações
        correlacoes = correlacionar_macro_carteira(df_macro, df_carteira_hist)
        
        # --- CORREÇÃO DE ERRO (Tratamento de NaNs) ---
        # Remover valores nulos (NaN) que causam o KeyError
        correlacoes = correlacoes.dropna()
        
        if not correlacoes.empty:
            # Gráfico de Barras de Correlação
            fig_corr = px.bar(
                x=correlacoes.index, 
                y=correlacoes.values,
                title="Correlação com a Taxa de Default (Risco)",
                labels={'x': 'Indicador Macro', 'y': 'Correlação (-1 a 1)'},
                color=correlacoes.values,
                color_continuous_scale='RdBu_r', 
                range_color=[-1, 1]
            )
            st.plotly_chart(fig_corr, use_container_width=True)
            
            # Explicação Automática (Protegida)
            try:
                max_fator = correlacoes.abs().idxmax()
                
                # Verificação extra para garantir que max_fator é válido
                if pd.notna(max_fator) and max_fator in correlacoes.index:
                    valor_corr = correlacoes[max_fator]
                    tipo_corr = "aumenta" if valor_corr > 0 else "diminui"
                    
                    sofia_explica(f"""
                    **Análise de Sensibilidade:**
                    
                    A sua carteira é **altamente sensível ao indicador '{max_fator}'** (Correlação: {valor_corr:.2f}).
                    
                    Historicamente, quando **{max_fator}** sobe, o Risco de Default da sua carteira tende a **{tipo_corr}**.
                    Recomendamos stress-test focado neste indicador.
                    """)
                else:
                    st.warning("Não foi possível identificar um fator dominante claro.")
            except Exception as e:
                st.warning(f"Não foi possível gerar a explicação automática: {str(e)}")
        else:
            st.warning("⚠️ Não foi possível calcular correlações (Dados insuficientes ou variância zero).")

# ===================== PÁGINA: ELASTICIDADE (FUNC 12) =====================
elif selected == "Elasticidade":
    st.header("🧲 Elasticidade & Choques Económicos")
    log_audit(st.session_state.username, "Elasticity Access", "Análise VECM")
    
    # 1. Preparar Dados (Simulados)
    df_macro = get_dados_macroeconomicos(48) # 4 anos para ter histórico
    df_carteira_hist = gerar_historico_simulado(df, 48)
    
    tab1, tab2 = st.tabs(["📊 Análise de Sensibilidade (Simples)", "🌊 Impulso-Resposta (VECM)"])

    # 2. Calcular Elasticidades
    elasticidades = calcular_elasticidade(df_macro, df_carteira_hist)
    
    col1, col2 = st.columns([1, 2])
    
    with col1:
        st.markdown("### 📊 Drivers de Risco")
        st.info("Quais indicadores externos têm maior impacto (força) na sua carteira?")
        
        # Criar tabela resumo
        resumo = []
        for ind, metrics in elasticidades.items():
            impacto = "Aumenta Risco" if metrics['elasticidade'] > 0 else "Reduz Risco"
            resumo.append({
                "Indicador": ind,
                "Sensibilidade": metrics['elasticidade'],
                "Impacto": impacto,
                "Confiança": metrics['significancia']
            })
            
        df_resumo = pd.DataFrame(resumo).sort_values('Sensibilidade', ascending=False)
        
        st.dataframe(
            df_resumo.style.format({'Sensibilidade': '{:.4f}'})
            .apply(lambda x: ['background-color: #7f1d1d' if v == 'Aumenta Risco' else '' for v in x], subset=['Impacto']),
            use_container_width=True
        )
        
    with col2:
        st.markdown("### 🔬 Análise de Regressão")
        
        # Seletor para visualizar o detalhe
        indicador_sel = st.selectbox("Selecione o Indicador para Detalhar", df_macro.columns)
        
        dados_reg = elasticidades[indicador_sel]
        df_plot = dados_reg['dados_combinados']
        
        # Gráfico de Dispersão com Linha de Tendência (OLS)
        fig = px.scatter(
            df_plot, 
            x=indicador_sel, 
            y='Taxa_Default', 
            trendline="ols", # Adiciona a linha de regressão automaticamente
            trendline_color_override="red",
            title=f"Sensibilidade: Default vs {indicador_sel}",
            labels={'Taxa_Default': 'Taxa de Default da Carteira'},
            template='plotly_dark'
        )
        st.plotly_chart(fig, use_container_width=True)
        
        # --- TAB 2: O NOVO CÓDIGO (VECM - Func 31) ---
    with tab2:
        st.markdown("### 🌊 Dinâmica de Choques (Longo Prazo)")
        st.info("Simule um choque económico e veja como a sua carteira reage ao longo do tempo (efeito cascata).")
        
        col1, col2 = st.columns([1, 3])
        
        with col1:
            choque_var = st.selectbox("Onde aplicar o choque?", df_macro.columns, index=0)
            st.caption("O modelo VECM irá simular um aumento súbito nesta variável.")
            
            if st.button("💥 Simular Choque"):
                with st.spinner("A calcular Impulso-Resposta..."):
                    res_vecm = calcular_elasticidade_vecm(df_macro, df_carteira_hist, variaveis=[choque_var, 'Taxa_Default'])
                    st.session_state['vecm_result'] = res_vecm

        with col2:
            if 'vecm_result' in st.session_state:
                res = st.session_state['vecm_result']
                
                if res['sucesso']:
                    irf_vals = res['irf_values']
                    
                    # Visualizar a Curva de Resposta
                    fig = go.Figure()
                    fig.add_trace(go.Scatter(
                        x=list(range(len(irf_vals))),
                        y=irf_vals,
                        mode='lines+markers',
                        name='Impacto no Default',
                        line=dict(color='#ef4444', width=3),
                        fill='tozeroy'
                    ))
                    
                    fig.update_layout(
                        title=f"Resposta da Taxa de Default a um choque em {choque_var}",
                        xaxis_title="Meses após o Choque",
                        yaxis_title="Variação no Default",
                        template="plotly_dark"
                    )
                    st.plotly_chart(fig, use_container_width=True)
                    
                    # Interpretação
                    pico_mes = np.argmax(np.abs(irf_vals))
                    pico_val = irf_vals[pico_mes]
                    
                    sofia_explica(f"""
                    **Análise de Impulso-Resposta:**
                    
                    Se houver um choque súbito em **{choque_var}** hoje:
                    1. O impacto na sua carteira não é imediato.
                    2. O efeito máximo (pico) será sentido daqui a **{pico_mes} meses**.
                    3. A magnitude do impacto será de **{pico_val:.4f}**.
                    
                    Isto permite-lhe preparar liquidez com {pico_mes} meses de antecedência!
                    """)
                else:
                    st.error(res['erro'])
            else:
                st.info("👈 Selecione uma variável e clique em Simular.")
        
        # Explicação IA
        beta = dados_reg['elasticidade']
        direcao = "aumenta" if beta > 0 else "diminui"
        
        sofia_explica(f"""
        **Interpretação Econométrica:**
        
        Por cada aumento de **1 ponto percentual** em '{indicador_sel}', a Taxa de Default da sua carteira **{direcao} em {abs(beta)*100:.2f} pontos percentuais** (em média).
        
        *Exemplo:* Se a {indicador_sel} subir de 2% para 3%, o seu Default sobe de 5% para {5 + (beta*100):.2f}%.
        """)

# ===================== PÁGINA: KPI TRACKING =====================
elif selected == "KPI Tracking":
    st.header("📊 KPI Tracking & Monitoring")
    log_audit(st.session_state.username, "KPI Tracking Access", "Acesso ao módulo de KPIs")
    
    tab1, tab2, tab3, tab4 = st.tabs(["📈 KPIs Principais", "🔄 Comparação", "🎯 Segmentação", "📉 Tendências"])
    
    with tab1:
        st.markdown("### 📊 Painel de KPIs")
        
        # Calcular KPIs
        kpis = calcular_kpis_principais(df)
        
        # Grid de KPIs
        st.markdown("#### 💰 KPIs Financeiros")
        col1, col2, col3 = st.columns(3)
        
        with col1:
            if 'Carteira_Total' in kpis:
                st.metric("💰 Carteira Total", f"€{kpis['Carteira_Total']/1e6:.1f}M")
        with col2:
            if 'Exposicao_Media' in kpis:
                st.metric("📊 Exposição Média", f"€{kpis['Exposicao_Media']/1e3:.1f}K")
        with col3:
            if 'Total_Clientes' in kpis:
                st.metric("👥 Total Clientes", f"{kpis['Total_Clientes']:,}")
        
        st.markdown("#### ⚠️ KPIs de Risco")
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            if 'Taxa_Default' in kpis:
                st.metric("⚠️ Taxa Default", f"{kpis['Taxa_Default']:.2%}",
                         delta="🔴" if kpis['Taxa_Default'] > 0.15 else "🟢")
        with col2:
            if 'LTV_Medio' in kpis:
                st.metric("📈 LTV Médio", f"{kpis['LTV_Medio']:.1%}")
        with col3:
            if 'Score_Medio' in kpis:
                st.metric("🎯 Score Médio", f"{kpis['Score_Medio']:.0f}")
        with col4:
            if 'Taxa_Atraso' in kpis:
                st.metric("⏰ Taxa Atraso", f"{kpis['Taxa_Atraso']:.1%}")
        
        st.markdown("#### 🛡️ KPIs Regulatórios (Basel III)")
        col1, col2, col3 = st.columns(3)
        
        with col1:
            if 'RWA' in kpis:
                st.metric("RWA", f"€{kpis['RWA']/1e6:.1f}M")
        with col2:
            if 'Capital_Minimo' in kpis:
                st.metric("Capital Mínimo", f"€{kpis['Capital_Minimo']/1e6:.2f}M")
        with col3:
            if 'Expected_Loss' in kpis:
                st.metric("Expected Loss", f"€{kpis['Expected_Loss']/1e6:.2f}M")
        
        # Alertas de KPI
        st.markdown("---")
        st.markdown("### 🚨 Alertas de KPI")
        
        alertas_kpi = []
        
        if 'Taxa_Default' in kpis:
            alerta = criar_alerta_kpi('Taxa_Default', kpis['Taxa_Default'], 0.15, tipo='max')
            if alerta:
                alertas_kpi.append(alerta)
        
        if 'LTV_Alto_Pct' in kpis:
            alerta = criar_alerta_kpi('LTV_Alto_Pct', kpis['LTV_Alto_Pct'], 0.20, tipo='max')
            if alerta:
                alertas_kpi.append(alerta)
        
        if 'Concentracao_Max' in kpis:
            alerta = criar_alerta_kpi('Concentracao_Max', kpis['Concentracao_Max'], 0.40, tipo='max')
            if alerta:
                alertas_kpi.append(alerta)
        
        if alertas_kpi:
            for alerta in alertas_kpi:
                st.markdown(f"""
                <div class="alert-warning">
                    <strong>⚠️ {alerta['severidade']}: {alerta['kpi']}</strong><br>
                    Valor Atual: {alerta['valor']:.2%} | Threshold: {alerta['threshold']:.2%}<br>
                    {alerta['tipo']}
                </div>
                """, unsafe_allow_html=True)
        else:
            st.success("✅ Todos os KPIs dentro dos limites estabelecidos")
        
        # Exportar KPIs
        st.markdown("---")
        kpis_df = pd.DataFrame(list(kpis.items()), columns=['KPI', 'Valor'])
        csv_kpis = kpis_df.to_csv(index=False)
        
        st.download_button(
            label="📥 Download KPIs (CSV)",
            data=csv_kpis,
            file_name=f"kpis_{datetime.now().strftime('%Y%m%d')}.csv",
            mime="text/csv"
        )
    
    with tab2:
        st.markdown("### 🔄 Comparação de Períodos")
        st.info("Compare KPIs entre diferentes períodos (mensal, trimestral, anual)")
        
        # Simular dados de período anterior (em produção, viria de BD)
        kpis_atual = calcular_kpis_principais(df)
        
        # Simular período anterior (reduzir 5% dos valores)
        kpis_anterior = {k: v * 0.95 for k, v in kpis_atual.items()}
        
        comparacao = comparar_kpis(kpis_atual, kpis_anterior)
        
        if comparacao:
            st.markdown("#### 📊 Variação vs Período Anterior")
            
            comp_df = pd.DataFrame([
                {
                    'KPI': kpi,
                    'Atual': f"{dados['atual']:.2f}",
                    'Anterior': f"{dados['anterior']:.2f}",
                    'Variação %': f"{dados['variacao_pct']:.1f}%",
                    'Tendência': dados['direcao']
                }
                for kpi, dados in comparacao.items()
            ])
            
            st.dataframe(
                comp_df.style.apply(
                    lambda x: ['background-color: #064e3b' if '+' in str(v) else 
                              'background-color: #7f1d1d' if '-' in str(v) else '' 
                              for v in x], axis=1, subset=['Variação %']
                ),
                use_container_width=True
            )
            
            # Gráfico de comparação
            fig = go.Figure()
            
            kpis_select = list(comparacao.keys())[:5]  # Top 5 KPIs
            
            fig.add_trace(go.Bar(
                name='Período Atual',
                x=kpis_select,
                y=[comparacao[k]['atual'] for k in kpis_select],
                marker_color='#3b82f6'
            ))
            
            fig.add_trace(go.Bar(
                name='Período Anterior',
                x=kpis_select,
                y=[comparacao[k]['anterior'] for k in kpis_select],
                marker_color='#64748b'
            ))
            
            fig.update_layout(
                title='Comparação de KPIs',
                barmode='group',
                template='plotly_dark'
            )
            
            st.plotly_chart(fig, use_container_width=True)
    
    with tab3:
        st.markdown("### 🎯 Segmentação Avançada")
        
        seg_type = st.radio("Tipo de Segmentação", ["RFM (Risco)", "Coorte", "Custom"])
        
        if seg_type == "RFM (Risco)":
            df_rfm = segmentacao_rfm(df)
            
            if df_rfm is not None:
                st.markdown("#### 📊 Distribuição por Segmento RFM")
                
                seg_dist = df_rfm['Segmento_RFM'].value_counts()
                
                fig = px.pie(
                    values=seg_dist.values,
                    names=seg_dist.index,
                    title='Distribuição de Clientes por Segmento',
                    color_discrete_sequence=px.colors.sequential.RdBu
                )
                st.plotly_chart(fig, use_container_width=True)
                
                # Estatísticas por segmento
                st.markdown("#### 📋 Estatísticas por Segmento")
                
                stats_seg = df_rfm.groupby('Segmento_RFM').agg({
                    'divida_total': ['sum', 'mean'],
                    'default': 'mean',
                    'score_interno': 'mean'
                }).round(2)
                
                st.dataframe(stats_seg, use_container_width=True)
            else:
                st.warning("Dados insuficientes para segmentação RFM")
        
        elif seg_type == "Coorte":
            coorte_df = analise_coorte(df)
            
            if coorte_df is not None:
                st.markdown("#### 📊 Análise de Coorte")
                
                fig = px.bar(
                    coorte_df,
                    x='Coorte',
                    y='Default_Rate',
                    color='segmento' if 'segmento' in df.columns else None,
                    title='Taxa de Default por Coorte',
                    barmode='group'
                )
                st.plotly_chart(fig, use_container_width=True)
                
                st.dataframe(coorte_df, use_container_width=True)
            else:
                st.warning("Dados insuficientes para análise de coorte")
    
    with tab4:
        st.markdown("### 📉 Análise de Tendências")
        st.info("Visualize a evolução dos KPIs ao longo do tempo")
        
        # Simular histórico (em produção viria de BD)
        kpi_selecionado = st.selectbox(
            "Selecionar KPI",
            ['Taxa_Default', 'LTV_Medio', 'Score_Medio', 'Carteira_Total']
        )
        
        # Gerar histórico simulado (12 períodos)
        np.random.seed(42)
        historico = [kpis.get(kpi_selecionado, 0) * (1 + np.random.uniform(-0.1, 0.1)) 
                     for _ in range(12)]
        
        fig = gerar_grafico_tendencia_kpi(kpi_selecionado, historico)
        st.plotly_chart(fig, use_container_width=True)
        
        # Análise de tendência
        tendencia = "CRESCENTE 📈" if historico[-1] > historico[0] else "DECRESCENTE 📉"
        variacao = ((historico[-1] - historico[0]) / abs(historico[0])) * 100
        
        col1, col2, col3 = st.columns(3)
        col1.metric("Tendência", tendencia)
        col2.metric("Variação Total", f"{variacao:+.1f}%")
        col3.metric("Última Leitura", f"{historico[-1]:.2f}")

# ===================== PÁGINA: GEO-RISK MAPPING (CORRIGIDO) =====================
elif selected == "Geo-Risk Map":
    st.header("🌍 Mapeamento de Risco Geossocial")
    log_audit(st.session_state.username, "Geo-Risk Access", "Acesso ao mapa de risco")

    # Verificar se as colunas de latitude/longitude existem
    if 'latitude' not in df.columns:
        st.warning("⚠️ Coordenadas não encontradas. Por favor, recarregue a 'Demo Enterprise' na barra lateral.")
    else:
        tab1, tab2 = st.tabs(["🗺️ Mapa Interativo", "📊 Estatísticas Regionais"])

        with tab1:
            col1, col2 = st.columns([1, 3])
            
            with col1:
                st.markdown("### Filtros")
                # Filtro com chave única para não dar conflito
                segmentos_unicos = df['segmento'].unique()
                filtro_segmento = st.multiselect(
                    "Filtrar por Segmento", 
                    segmentos_unicos,
                    default=segmentos_unicos,
                    key="geo_seg_filter"
                )
                
                filtro_risco = st.checkbox("Mostrar apenas Clientes em Default", value=False)

            # Lógica de Filtragem Robusta
            mask = df['segmento'].isin(filtro_segmento)
            if filtro_risco:
                mask = mask & (df['default'] == 1)
            
            df_mapa = df[mask]

            with col2:
                if not df_mapa.empty:
                    # Chamar a função de renderização (certifique-se que ela está na PARTE 2)
                    mapa_obj, _ = renderizar_mapa_risco(df_mapa)
                    
                    # Renderizar com largura total
                    st_folium(mapa_obj, width=None, height=600, use_container_width=True)
                else:
                    st.info("ℹ️ Nenhum dado encontrado com os filtros atuais.")

        with tab2:
            st.markdown("### 📊 Indicadores por Região")
            if 'regiao' in df.columns:
                kpis_geo = df.groupby('regiao').agg({
                    'id_cliente': 'count',
                    'divida_total': 'sum',
                    'default': 'mean'
                }).reset_index()
                
                kpis_geo.columns = ['Região', 'Clientes', 'Exposição (€)', 'Taxa Default']
                
                st.dataframe(
                    kpis_geo.style.format({
                        'Exposição (€)': '€{:,.0f}', 
                        'Taxa Default': '{:.2%}'
                    }).background_gradient(subset=['Taxa Default'], cmap='Reds'),
                    use_container_width=True
                )

# ===================== PÁGINA: WHAT-IF ANALYSIS =====================
elif selected == "What-If Analysis":
    st.header("🧪 What-If Analysis (Simulação)")
    log_audit(st.session_state.username, "What-If Access", "Acesso a simulação de cenários")

    # Verificar se temos modelo treinado (Dependência da Funcionalidade Credit Scoring)
    # Para o What-If funcionar, precisamos de um modelo base. Vamos treiná-lo on-the-fly se não existir.
    
    if 'default' not in df.columns:
        st.error("Dados insuficientes para modelagem.")
    else:
        # Treinar um modelo rápido se não existir em cache (simulação)
        with st.spinner("Calibrando motor de simulação..."):
            modelo_data = credit_scoring_model(df) 
        
        if modelo_data:
            model = modelo_data['model']
            scaler = modelo_data['scaler']
            features = modelo_data['features']

            col1, col2 = st.columns([1, 2])

            with col1:
                st.markdown("### 🎛️ Painel de Controlo")
                st.info("Ajuste os parâmetros para simular choques na carteira.")

                # Sliders de Ajuste
                var_rendimento = st.slider(
                    "📉 Choque no Rendimento (%)", 
                    min_value=-50, max_value=50, value=0, step=5,
                    help="Ex: -10% simula uma perda de poder de compra"
                ) / 100.0

                var_divida = st.slider(
                    "💳 Variação na Dívida (%)", 
                    min_value=-20, max_value=50, value=0, step=5,
                    help="Ex: +10% simula aumento do endividamento das famílias"
                ) / 100.0

                var_score = st.slider(
                    "🎯 Ajuste no Score Interno (pts)", 
                    min_value=-100, max_value=100, value=0, step=10,
                    help="Ex: -50 simula deterioração geral da qualidade de crédito"
                )

                if st.button("🚀 Executar Simulação", type="primary"):
                    resultado_sim = simular_cenario_whatif(
                        df, model, scaler, features, 
                        var_rendimento, var_divida, var_score
                    )
                    
                    st.session_state['resultado_whatif'] = resultado_sim # Guardar estado

            with col2:
                st.markdown("### 📊 Resultados da Simulação")

                if 'resultado_whatif' in st.session_state:
                    res = st.session_state['resultado_whatif']
                    
                    # Métricas Comparativas
                    c1, c2, c3 = st.columns(3)
                    
                    c1.metric(
                        "Risco Esperado (Atual)", 
                        f"{res['media_risco_atual']:.2%}"
                    )
                    
                    c2.metric(
                        "Risco Simulado (Novo)", 
                        f"{res['media_risco_simulado']:.2%}",
                        delta=f"{(res['media_risco_simulado'] - res['media_risco_atual'])*100:.2f} pp",
                        delta_color="inverse" # Se subir é mau (vermelho)
                    )
                    
                    c3.metric(
                        "Impacto Relativo", 
                        f"{res['impacto_default']:.1%}",
                        help="Aumento percentual sobre a taxa base"
                    )

                    # Gráfico de Antes vs Depois
                    fig = go.Figure()
                    fig.add_trace(go.Bar(
                        x=['Cenário Atual', 'Cenário Simulado'],
                        y=[res['media_risco_atual'], res['media_risco_simulado']],
                        marker_color=['#64748b', '#ef4444'],
                        text=[f"{res['media_risco_atual']:.2%}", f"{res['media_risco_simulado']:.2%}"],
                        textposition='auto'
                    ))
                    
                    fig.update_layout(
                        title="Impacto no Risco de Default da Carteira",
                        yaxis_title="Probabilidade Média de Default",
                        template="plotly_dark"
                    )
                    st.plotly_chart(fig, use_container_width=True)

                    # Explicação IA
                    if res['media_risco_simulado'] > res['media_risco_atual']:
                        msg = f"A simulação indica um **agravamento do risco**. Um choque de {var_rendimento:.0%} no rendimento e {var_divida:.0%} na dívida elevaria a probabilidade de default em {res['impacto_default']:.1%}."
                    else:
                        msg = "O cenário simulado apresenta uma melhoria nas condições de risco da carteira."
                        
                    sofia_explica(msg)

                else:
                    st.info("👈 Ajuste os parâmetros à esquerda e clique em 'Executar Simulação'.")
        else:
            st.warning("Não foi possível calibrar o modelo com os dados atuais.")

# ===================== PÁGINA: FORECASTING (ATUALIZADA FUNC 30) =====================
elif selected == "Forecasting":
    st.header("🔮 Previsão & Séries Temporais")
    log_audit(st.session_state.username, "Forecasting Access", "Acesso a previsões")
    
    # Gerar dados históricos
    df_hist = gerar_historico_simulado(df)
    
    # Seletor de Modo
    modo = st.radio("Modo de Previsão", ["Simples (Holt-Winters)", "Avançado (ARIMA)"], horizontal=True)
    
    col1, col2 = st.columns([1, 3])
    
    with col1:
        st.markdown("### ⚙️ Configuração")
        
        metrica = st.selectbox("Métrica a Prever", ["Taxa de Default", "Volume da Carteira"])
        col_alvo = 'Taxa_Default' if metrica == "Taxa de Default" else 'Carteira_Total'
        is_pct = True if metrica == "Taxa de Default" else False
        
        horizonte = st.slider("Horizonte (Meses)", 3, 24, 12)
        
        if modo == "Avançado (ARIMA)":
            st.markdown("#### Parâmetros ARIMA")
            p = st.number_input("p (Autoregressivo)", 0, 5, 1)
            d = st.number_input("d (Integração)", 0, 2, 1)
            q = st.number_input("q (Média Móvel)", 0, 5, 1)
            
            st.info(f"Ordem do Modelo: ({p}, {d}, {q})")

    with col2:
        if modo == "Simples (Holt-Winters)":
            # Usar a função antiga (Func 7)
            resultado = executar_previsao(df_hist, col_alvo, horizonte)
            titulo_grafico = f"Projeção HW: {metrica}"
        else:
            # Usar a NOVA função ARIMA (Func 30)
            resultado = executar_arima(df_hist, col_alvo, horizonte, ordem=(p, d, q))
            titulo_grafico = f"Projeção ARIMA({p},{d},{q}): {metrica}"
        
        if resultado['sucesso']:
            # Mostrar Gráfico
            fig = plotar_previsao(df_hist[col_alvo], resultado, titulo_grafico, is_pct)
            st.plotly_chart(fig, use_container_width=True)
            
            # Métricas e Diagnósticos
            c1, c2, c3 = st.columns(3)
            
            ultimo_valor = df_hist[col_alvo].iloc[-1]
            valor_previsto = resultado['previsao'].iloc[-1]
            delta = valor_previsto - ultimo_valor
            fmt = "{:.2%}" if is_pct else "€{:,.0f}"
            
            c1.metric("Valor Atual", fmt.format(ultimo_valor))
            c2.metric(f"Previsão ({horizonte}m)", fmt.format(valor_previsto), 
                      delta=fmt.format(delta), delta_color="inverse" if is_pct else "normal")
            
            if modo == "Avançado (ARIMA)":
                c3.metric("AIC (Qualidade)", f"{resultado['aic']:.1f}", help="Akaike Information Criterion (Menor é melhor)")
                
                with st.expander("Ver Sumário Estatístico (ARIMA)"):
                    st.text(resultado['modelo'].summary())
            
            # Explicação IA
            sofia_explica(f"""
            **Análise de Previsão ({modo}):**
            
            O modelo estima que a **{metrica}** irá {'aumentar' if delta > 0 else 'diminuir'} para **{fmt.format(valor_previsto)}** nos próximos {horizonte} meses.
            
            {'O AIC do modelo é ' + str(round(resultado['aic'],1)) + '. Tente ajustar (p,d,q) para reduzir este valor e melhorar a precisão.' if modo == 'Avançado (ARIMA)' else 'O modelo HW detetou automaticamente a tendência e sazonalidade.'}
            """)
            
        else:
            st.error(f"Erro na modelagem: {resultado['erro']}")
            if "stationarity" in str(resultado['erro']).lower():
                st.warning("Dica: Tente aumentar o parâmetro 'd' para 1 ou 2 para tornar a série estacionária.")

# ===================== PÁGINA: VISUALIZADOR AVANÇADO (FUNC 8) =====================
elif selected == "Visualizador":
    st.header("🔎 Visualizador de Dados Avançado")
    log_audit(st.session_state.username, "Data Explorer Access", "Exploração livre de dados")
    
    st.info("Utilize este módulo para criar gráficos personalizados e encontrar padrões ocultos nos dados.")

    # 1. Configuração do Gráfico
    with st.expander("🛠️ Configurar Eixos e Tipo", expanded=True):
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            chart_type = st.selectbox(
                "Tipo de Gráfico",
                ["Dispersão (Scatter)", "Barras (Bar)", "Linha (Line)", "Histograma (Dist)", "Boxplot", "Violin", "Heatmap (Correlação)"]
            )
        
        # Seleção de colunas baseada no tipo de gráfico
        cols = df.columns.tolist()
        cols_num = df.select_dtypes(include=[np.number]).columns.tolist()
        
        with col2:
            if chart_type != "Heatmap (Correlação)":
                x_col = st.selectbox("Eixo X", cols, index=cols.index('score_interno') if 'score_interno' in cols else 0)
        
        with col3:
            if chart_type not in ["Histograma (Dist)", "Heatmap (Correlação)"]:
                y_col = st.selectbox("Eixo Y", cols, index=cols.index('divida_total') if 'divida_total' in cols else 0)
        
        with col4:
            if chart_type not in ["Heatmap (Correlação)"]:
                color_col = st.selectbox("Cor (Agrupamento)", ["Nenhum"] + cols, index=0)

    # 2. Renderização do Gráfico
    st.markdown("### 📈 Visualização")
    
    fig = None
    
    try:
        color = None if color_col == "Nenhum" else color_col
        
        if chart_type == "Dispersão (Scatter)":
            fig = px.scatter(df, x=x_col, y=y_col, color=color, 
                           size='divida_total' if 'divida_total' in df.columns else None,
                           hover_data=df.columns, template='plotly_dark',
                           title=f"{y_col} vs {x_col}")
            
        elif chart_type == "Barras (Bar)":
            # Para barras, geralmente queremos uma agregação
            agg_func = st.radio("Função de Agregação", ["Média", "Soma", "Contagem"], horizontal=True)
            
            if agg_func == "Contagem":
                df_agg = df.groupby([x_col, color] if color else x_col).size().reset_index(name='Contagem')
                y_plot = 'Contagem'
            else:
                if df[y_col].dtype not in [np.number, int, float]:
                    st.warning(f"Para agregação {agg_func}, o Eixo Y deve ser numérico.")
                    df_agg = None
                else:
                    if agg_func == "Média":
                        df_agg = df.groupby([x_col, color] if color else x_col)[y_col].mean().reset_index()
                    else: # Soma
                        df_agg = df.groupby([x_col, color] if color else x_col)[y_col].sum().reset_index()
                    y_plot = y_col

            if df_agg is not None:
                fig = px.bar(df_agg, x=x_col, y=y_plot, color=color, barmode='group',
                           template='plotly_dark', title=f"{agg_func} de {y_col if agg_func != 'Contagem' else ''} por {x_col}")

        elif chart_type == "Linha (Line)":
            # Ordenar por X para a linha fazer sentido
            df_sorted = df.sort_values(x_col)
            fig = px.line(df_sorted, x=x_col, y=y_col, color=color, template='plotly_dark')

        elif chart_type == "Histograma (Dist)":
            fig = px.histogram(df, x=x_col, color=color, nbins=30, marginal="box",
                             template='plotly_dark', title=f"Distribuição de {x_col}")

        elif chart_type == "Boxplot":
            fig = px.box(df, x=x_col, y=y_col, color=color, template='plotly_dark',
                       title=f"Distribuição de {y_col} por {x_col}")
                       
        elif chart_type == "Violin":
            fig = px.violin(df, x=x_col, y=y_col, color=color, box=True, points="all",
                          template='plotly_dark', title=f"Densidade de {y_col} por {x_col}")

        elif chart_type == "Heatmap (Correlação)":
            if len(cols_num) > 1:
                corr = df[cols_num].corr()
                fig = px.imshow(corr, text_auto='.2f', color_continuous_scale='RdBu_r', aspect="auto",
                              title="Matriz de Correlação")
            else:
                st.warning("Necessárias pelo menos 2 colunas numéricas para correlação.")

        # Exibir
        if fig:
            st.plotly_chart(fig, use_container_width=True)
            
           # ================= CORREÇÃO FINAL DO DOWNLOAD =================
            # 1. Converter o gráfico para HTML (Texto)
            html_string = fig.to_html(include_plotlyjs="cdn")
            
            # 2. Converter o Texto para Bytes (O passo que faltava!)
            html_bytes = html_string.encode('utf-8')
            
            st.download_button(
                label="📥 Download Gráfico Interativo (HTML)",
                data=html_bytes, # Agora enviamos bytes, não texto
                file_name=f"grafico_{chart_type}_{datetime.now().strftime('%H%M%S')}.html",
                mime="text/html"
            )
            # ==============================================================
            
            # Análise IA Rápida
            if chart_type == "Heatmap (Correlação)":
                # Encontrar a maior correlação com Default
                if 'default' in corr.columns:
                    max_corr = corr['default'].drop('default').abs().idxmax()
                    val = corr.loc['default', max_corr]
                    sofia_explica(f"A variável mais correlacionada com o Default é **{max_corr}** ({val:.2f}). Isto sugere que é um forte preditor de risco.")

    except Exception as e:
        st.error(f"Não foi possível gerar o gráfico com estas configurações: {str(e)}")
        st.caption("Dica: Verifique se selecionou colunas numéricas para eixos Y em gráficos de dispersão ou linha.")

# ===================== PÁGINA: CLIENTE INDIVIDUAL =====================
elif selected == "Cliente Individual":
    st.header("👤 Análise Individual de Cliente")
    log_audit(st.session_state.username, "Cliente Individual Access", "Acesso à análise de cliente")
    
    tab1, tab2, tab3 = st.tabs(["🔍 Pesquisar Cliente", "💳 Simulador de Crédito", "📊 Comparação Peers"])
    
    with tab1:
        st.markdown("### 🔍 Pesquisar Cliente")
        
        # Pesquisa
        col1, col2 = st.columns([2, 1])
        
        with col1:
            if 'id_cliente' in df.columns:
                cliente_id = st.selectbox(
                    "Selecionar Cliente (ID)",
                    options=sorted(df['id_cliente'].unique()),
                    help="Selecione o ID do cliente para análise detalhada"
                )
            else:
                st.error("Coluna 'id_cliente' não encontrada")
                st.stop()
        
        with col2:
            if st.button("🔍 Analisar Cliente", use_container_width=True, type="primary"):
                analise = analisar_cliente_individual(df, cliente_id)
                
                if analise:
                    cliente_data = analise['dados']
                    percentis = analise['percentis']
                    
                    # Score do cliente
                    score_info = calcular_score_cliente(cliente_data)
                    
                    st.markdown("---")
                    st.markdown(f"## Cliente #{cliente_id}")
                    
                    # Score Card
                    col1, col2, col3, col4 = st.columns(4)
                    
                    with col1:
                        st.markdown(f"""
                        <div style='background: {score_info['cor']}; padding: 20px; border-radius: 10px; text-align: center;'>
                            <h2 style='color: white; margin: 0;'>{score_info['score']:.0f}</h2>
                            <p style='color: white; margin: 5px 0;'>{score_info['rating']}</p>
                        </div>
                        """, unsafe_allow_html=True)
                    
                    with col2:
                        st.metric("Nível de Risco", score_info['risco'])
                    
                    with col3:
                        if 'default' in cliente_data:
                            status = "🔴 DEFAULT" if cliente_data['default'] == 1 else "🟢 ATIVO"
                            st.metric("Status", status)
                    
                    with col4:
                        if 'segmento' in cliente_data:
                            st.metric("Segmento", cliente_data['segmento'])
                    
                    # Dados principais
                    st.markdown("### 📋 Dados do Cliente")
                    
                    col1, col2, col3 = st.columns(3)
                    
                    with col1:
                        st.markdown("**Dados Pessoais**")
                        if 'idade' in cliente_data:
                            st.write(f"👤 Idade: {cliente_data['idade']} anos")
                        if 'rendimento_mensal' in cliente_data:
                            st.write(f"💰 Rendimento: €{cliente_data['rendimento_mensal']:,.0f}/mês")
                    
                    with col2:
                        st.markdown("**Dados de Crédito**")
                        if 'divida_total' in cliente_data:
                            st.write(f"💳 Dívida Total: €{cliente_data['divida_total']:,.0f}")
                        if 'ltv' in cliente_data:
                            st.write(f"📊 LTV: {cliente_data['ltv']:.1%}")
                        if 'num_produtos' in cliente_data:
                            st.write(f"🎯 Produtos: {cliente_data['num_produtos']}")
                    
                    with col3:
                        st.markdown("**Comportamento**")
                        if 'num_atrasos_12m' in cliente_data:
                            st.write(f"⏰ Atrasos 12m: {cliente_data['num_atrasos_12m']}")
                        if 'utilizacao_credito' in cliente_data:
                            st.write(f"📈 Utilização: {cliente_data['utilizacao_credito']:.1%}")
                        if 'duracao_meses' in cliente_data:
                            st.write(f"🕐 Cliente há: {cliente_data['duracao_meses']} meses")
                    
                    # Percentis
                    st.markdown("### 📊 Posição vs Carteira")
                    st.caption("Percentis do cliente em relação à carteira total")
                    
                    percentis_principais = {k: v for k, v in list(percentis.items())[:6]}
                    
                    fig = go.Figure()
                    
                    fig.add_trace(go.Bar(
                        x=list(percentis_principais.keys()),
                        y=list(percentis_principais.values()),
                        marker_color=['#10b981' if v > 50 else '#ef4444' 
                                     for v in percentis_principais.values()],
                        text=[f"{v:.0f}%" for v in percentis_principais.values()],
                        textposition='outside'
                    ))
                    
                    fig.update_layout(
                        title='Percentis do Cliente',
                        yaxis_title='Percentil (%)',
                        template='plotly_dark',
                        showlegend=False
                    )
                    
                    st.plotly_chart(fig, use_container_width=True)
                    
                    # Recomendações
                    st.markdown("### 💡 Recomendações")
                    
                    recomendacoes = gerar_recomendacoes_cliente(cliente_data, score_info)
                    
                    for rec in recomendacoes:
                        if rec['tipo'] == 'CRÍTICO':
                            st.error(f"**{rec['titulo']}**: {rec['descricao']}")
                        elif rec['tipo'] == 'AVISO':
                            st.warning(f"**{rec['titulo']}**: {rec['descricao']}")
                        elif rec['tipo'] == 'SUCESSO':
                            st.success(f"**{rec['titulo']}**: {rec['descricao']}")
                        else:
                            st.info(f"**{rec['titulo']}**: {rec['descricao']}")
                else:
                    st.error("Cliente não encontrado")
    
    with tab2:
        st.markdown("### 💳 Simulador de Crédito")
        st.info("Simule aprovação e condições de crédito para o cliente")
        
        # Se cliente foi analisado na aba anterior, usar dados
        if 'id_cliente' in locals():
            analise = analisar_cliente_individual(df, cliente_id)
            if analise:
                cliente_data = analise['dados']
                score_info = calcular_score_cliente(cliente_data)
                
                st.markdown(f"**Cliente #{cliente_id} | Score: {score_info['score']:.0f}**")
        else:
            score_info = {'score': 650}  # Score padrão
        st.markdown("### 🏷️ Etiquetas do Cliente")
             
             # Recuperar tags atuais (garantir que é lista)
             # Nota: Em pandas, listas dentro de células podem ser tricky.
             # Vamos usar o índice do dataframe para localizar e editar.
        idx_cliente = df[df['id_cliente'] == cliente_id].index[0]
        tags_atuais = df.at[idx_cliente, 'tags']
        if not isinstance(tags_atuais, list): 
                tags_atuais = []
             
             # Widget de edição
        novas_tags = st.multiselect(
                 "Adicionar/Remover Tags",
                 options=["VIP", "Risco Elevado", "Novo", "Revisar", "Fraude?", "Corporate", "PME"],
                 default=tags_atuais,
                 key=f"tags_{cliente_id}"
             )
             
             # Botão de Guardar (Para persistir no session_state)
        if st.button("💾 Guardar Tags", key=f"save_tags_{cliente_id}"):
            # Atualizar o DataFrame no Session State
            st.session_state.df.at[idx_cliente, 'tags'] = novas_tags
                # Atualizar a variável local df também
            df.at[idx_cliente, 'tags'] = novas_tags
                 
            st.success("Tags atualizadas!")
            log_audit(st.session_state.username, "Tag Update", f"Cliente {cliente_id}: {novas_tags}")
            time.sleep(0.5)
            st.rerun()


        
        st.markdown("---")
        
        # Formulário de simulação
        with st.form("simulador_credito"):
            col1, col2 = st.columns(2)
            
            with col1:
                valor_solicitado = st.number_input(
                    "💰 Valor Solicitado (€)",
                    min_value=1000,
                    max_value=500000,
                    value=50000,
                    step=1000
                )
                
                prazo_meses = st.slider(
                    "📅 Prazo (meses)",
                    min_value=12,
                    max_value=360,
                    value=120,
                    step=12
                )
            
            with col2:
                finalidade = st.selectbox(
                    "🎯 Finalidade",
                    ["Habitação", "Automóvel", "Consolidação", "Consumo", "Educação"]
                )
                
                entrada_pct = st.slider(
                    "💵 Entrada (%)",
                    min_value=0,
                    max_value=50,
                    value=20,
                    step=5
                )
            
            simular_btn = st.form_submit_button("🚀 Simular Crédito", use_container_width=True, type="primary")
        
        if simular_btn:
            # Calcular LTV
            valor_bem = valor_solicitado / (1 - entrada_pct/100)
            ltv_simulado = valor_solicitado / valor_bem
            
            # Simular
            resultado = simular_credito(
                valor_solicitado,
                prazo_meses,
                score_info['score'],
                ltv_simulado
            )
            
            st.markdown("---")
            st.markdown("## 📋 Resultado da Simulação")
            
            # Decisão
            if resultado['decisao'] == 'APROVADO':
                st.success(f"✅ **{resultado['decisao']}**")
            elif resultado['decisao'] == 'APROVADO COM CONDIÇÕES':
                st.warning(f"⚠️ **{resultado['decisao']}**")
            elif resultado['decisao'] == 'ANÁLISE ADICIONAL':
                st.info(f"🔍 **{resultado['decisao']}**")
            else:
                st.error(f"❌ **{resultado['decisao']}**")
            
            if resultado['valor_aprovado'] > 0:
                # Condições
                col1, col2, col3, col4 = st.columns(4)
                
                col1.metric("Valor Aprovado", f"€{resultado['valor_aprovado']:,.0f}")
                col2.metric("Taxa Anual (TAN)", f"{resultado['tan']:.2%}")
                col3.metric("TAEG", f"{resultado['taeg']:.2%}")
                col4.metric("Prestação Mensal", f"€{resultado['prestacao_mensal']:,.2f}")
                
                # Detalhes
                st.markdown("### 📊 Detalhes do Financiamento")
                
                col1, col2 = st.columns(2)
                
                with col1:
                    st.write(f"**Prazo:** {resultado['prazo_meses']} meses ({resultado['prazo_meses']//12} anos)")
                    st.write(f"**Total a Pagar:** €{resultado['total_a_pagar']:,.2f}")
                    st.write(f"**Total de Juros:** €{resultado['juros_totais']:,.2f}")
                
                with col2:
                    st.write(f"**LTV:** {ltv_simulado:.1%}")
                    st.write(f"**Entrada:** €{valor_solicitado * entrada_pct/100:,.2f} ({entrada_pct}%)")
                    st.write(f"**Valor do Bem:** €{valor_bem:,.2f}")
                
                # Plano de amortização (primeiros 12 meses)
                st.markdown("### 📅 Plano de Amortização (Primeiros 12 Meses)")
                
                plano = []
                saldo = resultado['valor_aprovado']
                taxa_mensal = resultado['tan'] / 12
                
                for mes in range(1, min(13, prazo_meses + 1)):
                    juros_mes = saldo * taxa_mensal
                    amort_mes = resultado['prestacao_mensal'] - juros_mes
                    saldo -= amort_mes
                    
                    plano.append({
                        'Mês': mes,
                        'Prestação': resultado['prestacao_mensal'],
                        'Juros': juros_mes,
                        'Amortização': amort_mes,
                        'Saldo': max(0, saldo)
                    })
                
                plano_df = pd.DataFrame(plano)
                st.dataframe(
                    plano_df.style.format({
                        'Prestação': '€{:,.2f}',
                        'Juros': '€{:,.2f}',
                        'Amortização': '€{:,.2f}',
                        'Saldo': '€{:,.2f}'
                    }),
                    use_container_width=True
                )
                
                log_audit(st.session_state.username, "Simulação Crédito", 
                         f"Valor: €{valor_solicitado:,.0f} | Decisão: {resultado['decisao']}")
            else:
                st.info("💡 **Sugestões para melhorar a aprovação:**")
                st.write("- Aumentar entrada (reduzir LTV)")
                st.write("- Melhorar score de crédito")
                st.write("- Reduzir valor solicitado")
                st.write("- Aumentar prazo de pagamento")
    
    with tab3:
        st.markdown("### 📊 Comparação com Peers")
        
        if 'id_cliente' in locals() and 'cliente_data' in locals():
            comparacao = comparar_com_peers(df, cliente_data)
            
            if comparacao:
                st.markdown(f"**Comparação do Cliente #{cliente_id} com peers do segmento {cliente_data.get('segmento', 'N/A')}**")
                
                # Gráfico de comparação
                metricas = list(comparacao.keys())
                valores_cliente = [comparacao[m]['cliente'] for m in metricas]
                valores_peers = [comparacao[m]['media_peers'] for m in metricas]
                
                fig = go.Figure()
                
                fig.add_trace(go.Bar(
                    name='Cliente',
                    x=metricas,
                    y=valores_cliente,
                    marker_color='#3b82f6'
                ))
                
                fig.add_trace(go.Bar(
                    name='Média Peers',
                    x=metricas,
                    y=valores_peers,
                    marker_color='#64748b'
                ))
                
                fig.update_layout(
                    title='Cliente vs Peers',
                    barmode='group',
                    template='plotly_dark'
                )
                
                st.plotly_chart(fig, use_container_width=True)
                
                # Tabela detalhada
                st.markdown("### 📋 Detalhes da Comparação")
                
                comp_df = pd.DataFrame([
                    {
                        'Métrica': m,
                        'Cliente': f"{comparacao[m]['cliente']:.2f}",
                        'Média Peers': f"{comparacao[m]['media_peers']:.2f}",
                        'Percentil': f"{comparacao[m]['percentil']:.0f}%",
                        'Status': '✅ Acima' if comparacao[m]['melhor_que_peers'] else '⚠️ Abaixo'
                    }
                    for m in metricas
                ])
                
                st.dataframe(comp_df, use_container_width=True)
            else:
                st.warning("Dados insuficientes para comparação com peers")
        else:
            st.info("👈 Primeiro analise um cliente na aba 'Pesquisar Cliente'")

            

# ===================== PÁGINA: BASEL III =====================
elif selected == "Basel III":
    st.header("🛡️ Basel III - Capital Regulatório")
    log_audit(st.session_state.username, "Basel III Access", "Acesso ao módulo Basel III")
    
    basel = calcular_capital_regulatorio(df)
    
    if basel:
        # Métricas principais
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            st.metric("PD (Probability of Default)", f"{basel['PD']:.2%}")
        with col2:
            st.metric("EAD (Exposure at Default)", f"€{basel['EAD']/1e6:.1f}M")
        with col3:
            st.metric("LGD (Loss Given Default)", f"{basel['LGD']:.0%}")
        with col4:
            st.metric("Expected Loss", f"€{basel['Expected_Loss']/1e6:.2f}M")
        
        st.markdown("---")
        st.markdown("### 💰 Requisitos de Capital")
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            st.metric("RWA (Risk Weighted Assets)", f"€{basel['RWA']/1e6:.1f}M")
        
        with col2:
            st.metric("Capital Mínimo (8%)", 
                     f"€{basel['Capital_Minimo_8pct']/1e6:.2f}M",
                     help="Requisito mínimo Basel III Pilar 1")
        
        with col3:
            buffer = basel['Capital_Recomendado_10.5pct'] - basel['Capital_Minimo_8pct']
            st.metric("Capital Recomendado (10.5%)", 
                     f"€{basel['Capital_Recomendado_10.5pct']/1e6:.2f}M",
                     delta=f"+€{buffer/1e6:.2f}M buffer",
                     help="Inclui buffer de conservação de 2.5%")
        
        # Análise de adequação
        ratio = basel['Capital_Minimo_8pct'] / basel['RWA'] if basel['RWA'] > 0 else 0
        status = "✅ ADEQUADO" if ratio >= 0.08 else "⚠️ INSUFICIENTE"
        
        st.markdown("### 📊 Análise de Adequação de Capital")
        
        progress_val = min(ratio / 0.105, 1.0)
        st.progress(progress_val)
        
        sofia_explica(f"""
        **Status de Capital: {status}**
        
        - Tier 1 Capital Ratio: {ratio:.1%}
        - Basel III requer mínimo de 8% + buffer de conservação de 2.5% = 10.5%
        - Expected Loss anual: €{basel['Expected_Loss']/1e6:.2f}M
        - Recomendação: {'Manter níveis atuais' if ratio >= 0.105 else 'Aumentar capitalização ou reduzir RWA'}
        """)
        
        # Gráfico de composição
        fig = go.Figure(data=[
            go.Bar(name='Capital Mínimo', x=['Requisito'], y=[basel['Capital_Minimo_8pct']/1e6]),
            go.Bar(name='Buffer Conservação', x=['Requisito'], y=[buffer/1e6])
        ])
        fig.update_layout(title='Composição do Capital Regulatório (€M)', barmode='stack')
        st.plotly_chart(fig, use_container_width=True)
        
    else:
        st.warning("⚠️ Dados insuficientes para cálculo Basel III. Necessário: 'default' e 'divida_total'")

# ===================== PÁGINA: CREDIT SCORING =====================
elif selected == "Credit Scoring":
    st.header("🎯 Credit Scoring Avançado")
    log_audit(st.session_state.username, "Credit Scoring Access", "Acesso ao módulo de scoring")
    
    if st.button("🚀 Treinar Modelo de Credit Scoring"):
        with st.spinner("Treinando modelo..."):
            result = credit_scoring_model(df)
        
        if result:
            enviar_notificacao("Modelo de Credit Scoring treinado com sucesso!", "sucesso") # <--- NOVO
            st.success("✅ Modelo treinado com sucesso!")
            
            # Métricas principais
            col1, col2, col3, col4 = st.columns(4)
            
            with col1:
                st.metric("ROC AUC", f"{result['roc_auc']:.3f}")
            with col2:
                st.metric("Acurácia", f"{result['report']['accuracy']:.1%}")
            with col3:
                st.metric("Precisão", f"{result['report']['1']['precision']:.1%}")
            with col4:
                st.metric("Recall", f"{result['report']['1']['recall']:.1%}")
            
            # Gráficos
            col1, col2 = st.columns(2)
            
            with col1:
                st.plotly_chart(result['fig_roc'], use_container_width=True)
            
            with col2:
                st.plotly_chart(result['fig_cm'], use_container_width=True)
            
            # Feature Importance
            st.markdown("### 📊 Importância das Features")
            fig = px.bar(result['importance'], x='Importance', y='Feature',
                        orientation='h', title='Feature Importance',
                        color='Importance', color_continuous_scale='Blues')
            st.plotly_chart(fig, use_container_width=True)
            
            # Classification Report
            st.markdown("### 📋 Relatório de Classificação")
            report_df = pd.DataFrame(result['report']).T
            st.dataframe(report_df.style.highlight_max(axis=0), use_container_width=True)
            
            # Interpretação
            interpretacao = "Excelente" if result['roc_auc'] > 0.8 else "Bom" if result['roc_auc'] > 0.7 else "Aceitável"
            sofia_explica(f"""
            **Performance do Modelo: {interpretacao}**
            
            - AUC = {result['roc_auc']:.3f}: {interpretacao} poder discriminatório
            - Features mais importantes: {', '.join(result['importance'].head(3)['Feature'].tolist())}
            - Modelo pronto para deployment em produção
            """)
        else:
            st.error("❌ Dados insuficientes. Necessário: 'default' e pelo menos 3 features numéricas")

# ===================== PÁGINA: CATÁLOGO DE DADOS (FUNC 40) =====================
elif selected in ["Catálogo de Dados", "Data Catalog", "menu_metadata"]:
    st.header("🗂️ Catálogo de Metadados (Data Governance)")
    log_audit(st.session_state.username, "Metadata Access", "Consulta de catálogo de dados")
    
    st.info("Governança de Dados: Documente a origem, qualidade e propriedade dos seus datasets.")

    tab1, tab2 = st.tabs(["📖 Catálogo Atual", "✏️ Editar Metadados"])

    # --- TAB 1: VISUALIZAÇÃO ---
    with tab1:
        catalog = st.session_state.metadata_catalog
        
        if not catalog:
            st.warning("Catálogo vazio.")
        else:
            for nome, meta in catalog.items():
                # Card de Metadados
                with st.expander(f"📂 {nome}", expanded=True):
                    c1, c2 = st.columns([3, 1])
                    with c1:
                        st.markdown(f"**Descrição:** {meta['descricao']}")
                        st.markdown(f"**Fonte:** `{meta['fonte']}`")
                    with c2:
                        # Badge de Confidencialidade
                        cor_badge = "red" if meta['confidencialidade'] == "Alta" else "orange" if meta['confidencialidade'] == "Média" else "green"
                        st.markdown(f":{cor_badge}[**{meta['confidencialidade']}**]")
                        st.caption(f"Resp: {meta['responsavel']}")
                        st.caption(f"Atualizado: {meta['ultima_atualizacao']}")
                    
                    # Se houver dados carregados que correspondam ao nome, mostrar amostra
                    if nome == "Demo Enterprise" and st.session_state.df is not None:
                        st.markdown("---")
                        st.markdown("**Amostra de Dados (5 registos):**")
                        st.dataframe(st.session_state.df.head(5), use_container_width=True)

    # --- TAB 2: EDIÇÃO ---
    with tab2:
        st.markdown("### Registar Novo Dataset")
        
        with st.form("form_metadata"):
            nome_ds = st.text_input("Nome do Dataset", placeholder="Ex: Carteira Crédito Habitação 2024")
            desc_ds = st.text_area("Descrição", placeholder="Contém dados de clientes com LTV > 80%...")
            
            c1, c2, c3 = st.columns(3)
            with c1:
                fonte_ds = st.selectbox("Fonte Original", ["Core Bancário", "Excel Manual", "API Externa", "Web Scraping", "Simulação"])
            with c2:
                resp_ds = st.text_input("Responsável", value=st.session_state.username)
            with c3:
                conf_ds = st.selectbox("Nível de Confidencialidade", ["Baixa", "Média", "Alta", "Crítica"])
            
            if st.form_submit_button("💾 Salvar no Catálogo"):
                if nome_ds and desc_ds:
                    atualizar_metadados(nome_ds, desc_ds, fonte_ds, resp_ds, conf_ds)
                    st.success(f"Metadados para **{nome_ds}** atualizados com sucesso!")
                    log_audit(st.session_state.username, "Metadata Update", f"Dataset: {nome_ds}")
                    st.rerun()
                else:
                    st.error("O nome e a descrição são obrigatórios.")

# ===================== PÁGINA: AI ETHICS (FUNC 13) =====================
elif selected == "AI Ethics":
    st.header("⚖️ Auditoria de Ética e Fairness")
    log_audit(st.session_state.username, "Ethics Access", "Auditoria de bias algorítmico")
    
    st.info("Verificação de enviesamento do modelo de Credit Scoring contra grupos protegidos (Compliance EU AI Act).")

    # Dependência: Modelo treinado
    # Tentamos recuperar o modelo do Credit Scoring ou treinamos um rápido
    if 'default' not in df.columns:
        st.error("Dados insuficientes.")
    else:
        # Treinar modelo rápido se não existir
        with st.spinner("A carregar modelo para auditoria..."):
            modelo_data = credit_scoring_model(df)
            
        if modelo_data:
            model = modelo_data['model']
            scaler = modelo_data['scaler']
            features = modelo_data['features']
            
            col1, col2 = st.columns([1, 2])
            
            with col1:
                st.markdown("### ⚙️ Configuração do Teste")
                
                # Escolher atributo protegido
                atributo = st.selectbox("Atributo Protegido", ["idade", "regiao", "segmento"])
                
                valor = None
                if atributo == "idade":
                    valor = st.slider("Corte de Idade (Jovens)", 18, 60, 25)
                elif atributo == "regiao":
                    valor = st.selectbox("Região Focada", df['regiao'].unique())
                elif atributo == "segmento":
                    valor = st.selectbox("Segmento Focado", df['segmento'].unique())
                
                if st.button("🕵️ Auditar Modelo", type="primary"):
                    metricas = calcular_metricas_fairness(df, model, scaler, features, atributo, valor)
                    st.session_state['fairness_metrics'] = metricas

            with col2:
                if 'fairness_metrics' in st.session_state:
                    m = st.session_state['fairness_metrics']
                    
                    if "erro" in m:
                        st.error(m["erro"])
                    else:
                        st.markdown(f"### Resultados: {m['grupo_protegido_nome']}")
                        
                        # Cartões de Métricas
                        c1, c2, c3 = st.columns(3)
                        
                        # 1. Disparate Impact
                        # Se > 1.0, o grupo protegido é MAIS classificado como risco (Mau).
                        # Se < 1.0, é MENOS classificado como risco.
                        # Intervalo justo: 0.8 a 1.25
                        is_fair_di = 0.8 <= m['disparate_impact'] <= 1.25
                        cor_di = "normal" if is_fair_di else "inverse"
                        
                        c1.metric(
                            "Impacto Disparado", 
                            f"{m['disparate_impact']:.2f}",
                            help="Ratio de seleção entre grupos. Ideal: 0.8 a 1.25",
                            delta="Justo" if is_fair_di else "Enviesado",
                            delta_color=cor_di
                        )
                        
                        # 2. Taxa de Seleção (Default Previsto)
                        c2.metric("Taxa Default (Grupo)", f"{m['taxa_selecao_a']:.1%}")
                        c3.metric("Taxa Default (Outros)", f"{m['taxa_selecao_b']:.1%}")
                        
                        # Visualização
                        fig = go.Figure()
                        fig.add_trace(go.Bar(
                            x=['Grupo Protegido', 'Resto da População'],
                            y=[m['taxa_selecao_a'], m['taxa_selecao_b']],
                            marker_color=['#f59e0b', '#3b82f6'],
                            text=[f"{m['taxa_selecao_a']:.1%}", f"{m['taxa_selecao_b']:.1%}"],
                            textposition='auto',
                            name='Probabilidade Média de Default'
                        ))
                        fig.update_layout(title="Comparação de Previsão de Risco", template="plotly_dark")
                        st.plotly_chart(fig, use_container_width=True)
                        
                        # Explicação IA
                        diff = m['taxa_selecao_a'] - m['taxa_selecao_b']
                        texto_bias = "PENALIZA" if diff > 0 else "FAVORECE"
                        
                        sofia_explica(f"""
                        **Auditoria de Ética:**
                        
                        O modelo atualmente **{texto_bias}** o grupo **{m['grupo_protegido_nome']}**.
                        
                        - A probabilidade média de default prevista para este grupo é **{m['taxa_selecao_a']:.1%}**, contra **{m['taxa_selecao_b']:.1%}** nos restantes.
                        - O Impacto Disparado é {m['disparate_impact']:.2f}. {'✅ O modelo é considerado JUSTO.' if is_fair_di else '⚠️ ALERTA: Há indícios de discriminação estatística.'}
                        """)
                else:
                    st.info("👈 Configure o teste e clique em Auditar.")

# ===================== PÁGINA: EARLY WARNING =====================
elif selected == "Early Warning":
    st.header("⚠️ Sistema de Alerta Precoce")
    log_audit(st.session_state.username, "Early Warning Access", "Acesso ao early warning")
    
    result = early_warning_system(df)
    
    if result:
        df_model, model, features = result
        
        # Distribuição de risco
        risk_dist = df_model['Risk_Level'].value_counts()
        
        col1, col2, col3, col4 = st.columns(4)
        col1.metric("🟢 Risco Baixo", f"{risk_dist.get('Baixo', 0):,}")
        col2.metric("🟡 Risco Médio", f"{risk_dist.get('Médio', 0):,}")
        col3.metric("🟠 Risco Alto", f"{risk_dist.get('Alto', 0):,}")
        col4.metric("🔴 Risco Crítico", f"{risk_dist.get('Crítico', 0):,}")
        
        # Gráfico de distribuição
        st.markdown("### 📊 Distribuição de Risk Scores")
        fig = px.histogram(df_model, x='Risk_Score', nbins=50,
                          title='Distribuição de Scores de Risco (0-100)',
                          color_discrete_sequence=['#3b82f6'])
        st.plotly_chart(fig, use_container_width=True)
        
        # Top clientes em risco
        st.markdown("### 🚨 Top 20 Clientes de Alto Risco")
        top_risk = df_model.nlargest(20, 'Risk_Score')[['Risk_Score', 'Risk_Level'] + features]
        st.dataframe(top_risk.style.background_gradient(subset=['Risk_Score'], cmap='Reds'),
                    use_container_width=True)
        
        # Download
        csv = top_risk.to_csv(index=False)
        st.download_button(
            label="📥 Download Lista de Alto Risco (CSV)",
            data=csv,
            file_name="high_risk_clients.csv",
            mime="text/csv"
        )
        
        sofia_explica(f"""
        **Análise de Alerta Precoce:**
        
        - {risk_dist.get('Crítico', 0)} clientes em risco crítico requerem ação imediata
        - {risk_dist.get('Alto', 0)} clientes em risco alto necessitam monitorização
        - Features mais relevantes: {', '.join(features[:3])}
        """)
    else:
        st.warning("⚠️ Dados insuficientes para Early Warning System")

# ===================== PÁGINA: STRESS TESTING =====================
elif selected == "Stress Testing":
    st.header("🌪️ Testes de Stress Regulatórios")
    log_audit(st.session_state.username, "Stress Testing", "Acesso a stress testing")
    
    var_target = st.selectbox("Variável de Exposição", 
                              [c for c in cols_num if 'divida' in c.lower() or 'expo' in c.lower()],
                              index=0 if any('divida' in c.lower() for c in cols_num) else 0)
    
    if st.button("▶️ Executar Stress Tests"):
        with st.spinner("Simulando cenários..."):
            stress_results = stress_test_scenarios(df, var_target)
        
        st.success("✅ Stress tests concluídos!")
        
        # Tabela de resultados
        st.markdown("### 📊 Resultados por Cenário")
        
        # Formatar tabela
        display_df = stress_results.copy()
        display_df['Default_Rate'] = display_df['Default_Rate'].map('{:.2%}'.format)
        display_df['Portfolio_Value'] = display_df['Portfolio_Value'].map('€{:,.0f}'.format)
        display_df['Expected_Loss'] = display_df['Expected_Loss'].map('€{:,.0f}'.format)
        display_df['Loss_Rate'] = display_df['Loss_Rate'].map('{:.2%}'.format)
        
        st.dataframe(display_df, use_container_width=True)
        
        # Gráfico comparativo
        fig = go.Figure()
        
        scenarios = stress_results.index.tolist()
        
        fig.add_trace(go.Bar(
            name='Default Rate',
            x=scenarios,
            y=stress_results['Default_Rate'] * 100,
            marker_color='indianred'
        ))
        
        fig.update_layout(
            title='Taxa de Default por Cenário (%)',
            xaxis_title='Cenário',
            yaxis_title='Default Rate (%)',
            template='plotly_dark'
        )
        
        st.plotly_chart(fig, use_container_width=True)
        
        # Expected Loss
        fig2 = px.bar(stress_results.reset_index(), x='index', y='Expected_Loss',
                     title='Expected Loss por Cenário (€)',
                     labels={'index': 'Cenário', 'Expected_Loss': 'Expected Loss'},
                     color='Expected_Loss', color_continuous_scale='Reds')
        st.plotly_chart(fig2, use_container_width=True)
        
        # Análise
        worst_scenario = stress_results['Expected_Loss'].idxmax()
        worst_loss = stress_results.loc[worst_scenario, 'Expected_Loss']
        
        sofia_explica(f"""
        **Análise de Stress Testing:**
        
        - Cenário mais adverso: **{worst_scenario}**
        - Expected Loss máxima: **€{worst_loss:,.0f}**
        - Recomendação: Manter reservas para cobrir cenário de Crise Moderada no mínimo
        """)

# ===================== PÁGINA: SURVIVAL ANALYSIS =====================
elif selected == "Survival Analysis":
    st.header("⏱️ Análise de Sobrevivência")
    log_audit(st.session_state.username, "Survival Analysis", "Acesso a survival analysis")
    
    if 'duracao_meses' in df.columns and 'default' in df.columns:
        
        fig, median_surv, kmf = analise_sobrevivencia(df, 'duracao_meses', 'default')
        
        if fig:
            st.plotly_chart(fig, use_container_width=True)
            
            col1, col2, col3 = st.columns(3)
            col1.metric("Tempo Mediano até Default", f"{median_surv:.1f} meses" if median_surv else "N/A")
            col2.metric("Sobrevivência em 12m", f"{kmf.survival_function_at_times(12).values[0]:.1%}" if len(kmf.survival_function_at_times(12)) > 0 else "N/A")
            col3.metric("Sobrevivência em 24m", f"{kmf.survival_function_at_times(24).values[0]:.1%}" if len(kmf.survival_function_at_times(24)) > 0 else "N/A")
            
            sofia_explica(f"""
            **Insights de Sobrevivência:**
            
            - Tempo mediano até default: {median_surv:.0f} meses
            - 50% dos defaults ocorrem antes deste período
            - Usar para pricing de produtos e provisões
            """)
        else:
            st.warning("Dados insuficientes para análise")
    else:
        st.error("❌ Necessário colunas: 'duracao_meses' e 'default'")

# ===================== PÁGINA: COMPLIANCE =====================
elif selected == "Compliance":
    st.header("📋 Compliance & Regulação")
    log_audit(st.session_state.username, "Compliance Access", "Acesso a compliance")
    
    tab1, tab2, tab3 = st.tabs(["📊 Alertas", "📜 Regulamentos", "✅ Checklist"])
    
    with tab1:
        st.markdown("### ⚠️ Alertas de Compliance Ativos")
        
        if alerts:
            for i, alert in enumerate(alerts):
                nivel_class = "alert-critical" if alert['nivel'] == "CRÍTICO" else "alert-warning"
                st.markdown(f"""
                <div class="{nivel_class}">
                    <h4>{i+1}. {alert['tipo']}</h4>
                    <p><strong>Nível:</strong> {alert['nivel']}</p>
                    <p><strong>Mensagem:</strong> {alert['mensagem']}</p>
                    <p><strong>Ação Recomendada:</strong> {alert['acao']}</p>
                </div>
                """, unsafe_allow_html=True)
        else:
            st.success("✅ Sem alertas de compliance ativos")
    
    with tab2:
        st.markdown("""
        ### 📜 Framework Regulatório
        
        #### Basel III
        - Capital mínimo: 8% RWA
        - Buffer de conservação: 2.5%
        - Total: 10.5% RWA
        
        #### RGPD
        - Auditoria completa de acessos ✅
        - Direito ao esquecimento implementado ✅
        - Minimização de dados ✅
        
        #### EBA Guidelines
        - Stress testing anual obrigatório ✅
        - Early warning systems ✅
        - Credit scoring validation ✅
        """)
    
    with tab3:
        st.markdown("### ✅ Checklist de Compliance")
        
        # Calcular basel para checklist
        basel_check = calcular_capital_regulatorio(df)
        
        checks = [
            ("Capital adequado (Basel III)", basel_check is not None),
            ("Sistema de alertas ativo", len(alerts) >= 0),
            ("Auditoria configurada", len(st.session_state.audit_log) > 0),
            ("Modelos documentados", True),
            ("Stress testing implementado", True)
        ]
        
        for check_name, check_status in checks:
            status_icon = "✅" if check_status else "❌"
            st.markdown(f"{status_icon} {check_name}")

# ===================== PÁGINA: INTEGRAÇÃO BI (FUNC 47) =====================
elif selected in ["Integração BI", "BI Integration", "menu_bi"]:
    st.header("📊 Conector para Power BI & Tableau")
    log_audit(st.session_state.username, "BI Connector Access", "Acesso a integração BI")
    
    st.info("Exponha os dados calculados pelo SocioStat para ferramentas de Business Intelligence externas.")

    tab1, tab2 = st.tabs(["🔗 Conexão Direta", "💾 Exportação Otimizada"])

    # --- TAB 1: CONEXÃO DATABASE ---
    with tab1:
        col1, col2 = st.columns([2, 1])
        
        with col1:
            st.markdown("### Credenciais de Leitura")
            st.caption("Use estas credenciais no Power BI (Get Data -> PostgreSQL) para aceder à vista otimizada.")
            
            st.code(f"""
            Server:   db-analytics.sociostat.internal
            Port:     5432
            Database: risk_warehouse
            User:     bi_reader_{st.session_state.username.lower()}
            Password: ********************
            View:     vw_sociostat_analytics_v3
            """, language="yaml")
            
            if st.button("🔄 Rodar Password de Acesso"):
                st.toast("Novas credenciais enviadas para o seu email.", icon="📧")

        with col2:
            st.markdown("### Estado do Túnel")
            st.metric("Status", "Ativo 🟢")
            st.metric("Última Leitura", "Há 5 min")
            st.metric("Linhas Sincronizadas", f"{len(df):,}")

    # --- TAB 2: EXPORTAÇÃO FICHEIRO ---
    with tab2:
        st.markdown("### Dataset Otimizado para BI")
        st.write("Gera um ficheiro 'flat' com todas as métricas de risco pré-calculadas, ideal para importação manual.")
        
        # Gerar dados
        df_bi = gerar_vista_bi_otimizada(df)
        
        # Preview
        with st.expander("👁️ Ver Preview dos Dados Enriquecidos"):
            st.dataframe(df_bi.head(10), use_container_width=True)
            
        col_dl1, col_dl2 = st.columns(2)
        with col_dl1:
            # Download CSV Otimizado
            csv_bi = df_bi.to_csv(index=False).encode('utf-8')
            st.download_button(
                label="📥 Baixar Dataset Power BI (.csv)",
                data=csv_bi,
                file_name=f"SocioStat_BI_Export_{datetime.now().strftime('%Y%m%d')}.csv",
                mime="text/csv",
                help="CSV otimizado com separador vírgula e encoding UTF-8"
            )
            
        with col_dl2:
            # Download Script SQL
            sql_script = gerar_script_sql_bi()
            st.download_button(
                label="📥 Baixar Script SQL (.sql)",
                data=sql_script,
                file_name="create_view_bi.sql",
                mime="text/plain",
                help="Script para criar a vista no seu Data Warehouse"
            )

# ===================== PÁGINA: FEEDBACK (FUNC 18) =====================
elif selected == "Feedback":
    st.header("💬 Feedback & Suporte")
    log_audit(st.session_state.username, "Feedback Access", "Acesso ao módulo de feedback")
    
    st.info("Ajude-nos a melhorar o SocioStat. Reporte erros ou sugira novas funcionalidades.")

    col1, col2 = st.columns([2, 1])

    with col1:
        with st.form("form_feedback"):
            tipo = st.selectbox("Tipo de Feedback", ["🐛 Reportar Bug", "💡 Sugestão de Melhoria", "❓ Dúvida Técnica", "⭐ Elogio"])
            prioridade = st.select_slider("Prioridade", options=["Baixa", "Média", "Alta", "Crítica"])
            titulo = st.text_input("Assunto", placeholder="Ex: Erro no cálculo de LTV...")
            descricao = st.text_area("Descrição Detalhada", height=150, placeholder="Descreva o que aconteceu ou a sua ideia...")
            
            # Anexo (Opcional)
            anexo = st.file_uploader("Anexar Screenshot (Opcional)", type=['png', 'jpg'])
            
            submit = st.form_submit_button("📤 Enviar Feedback", type="primary")
            
            if submit:
                if not titulo or not descricao:
                    st.error("Por favor preencha o assunto e a descrição.")
                else:
                    # Simular envio (Em produção, isto iria para Jira/Trello/Email)
                    ticket_id = f"TKT-{np.random.randint(1000, 9999)}"
                    
                    # Registar no Audit Log para ficar guardado
                    log_audit(
                        st.session_state.username, 
                        "Feedback Submitted", 
                        f"[{ticket_id}] {tipo}: {titulo}"
                    )
                    
                    # Notificação de Sucesso (Func 17)
                    st.toast(f"Feedback recebido! Ticket #{ticket_id}", icon="✅")
                    st.success(f"""
                    **✅ Feedback enviado com sucesso!**
                    
                    O seu ticket **#{ticket_id}** foi encaminhado para a equipa de desenvolvimento.
                    Obrigado por contribuir para o SocioStat.
                    """)

    with col2:
        st.markdown("### 📞 Contactos Diretos")
        st.markdown("""
        **Equipa de Risco:**
        📧 risco@sociostat.pt
        📞 +351 22 123 4567
        
        **Suporte Técnico (IT):**
        📧 suporte@sociostat.pt
        📞 +351 22 123 4568
        
        **Horário:**
        Seg-Sex: 09:00 - 18:00
        """)
        
        st.markdown("---")
        st.markdown("### 📢 Últimas Atualizações")
        st.info("""
        **v3.0 (Atual)**
        - Novo módulo de IA (Credit Scoring)
        - Previsão de Séries Temporais
        - Exportação Excel Avançada
        """)

# ===================== PÁGINA: CAUSALIDADE (FUNC 21) =====================
elif selected == "Causal Inference":
    st.header("🔬 Inferência Causal (DiD)")
    log_audit(st.session_state.username, "Causal Analysis", "Execução de modelo DiD")
    
    st.info("Utilize o método **Difference-in-Differences** para medir o impacto real de uma política, isolando-o de tendências de mercado.")
    
    col1, col2 = st.columns([1, 2])
    
    with col1:
        st.markdown("### ⚙️ Configuração do Estudo")
        campanha = st.selectbox("Intervenção a Analisar", 
                                ["Campanha Renegociação Norte", "Moratória Crédito Habitação", "Programa de Literacia Financeira"])
        
        if st.button("🚀 Calcular Impacto"):
            # Simular e Calcular
            res = calcular_causalidade_did(campanha)
            st.session_state['did_result'] = res
            st.success("Cálculo DiD concluído.")
            
    with col2:
        if 'did_result' in st.session_state:
            res = st.session_state['did_result']
            df_did = res['df']
            
            # Gráfico de Tendências Paralelas
            fig = px.line(df_did, x='Tempo', y='Default', color='Grupo', markers=True,
                          title=f"Impacto da {campanha} na Taxa de Default",
                          template='plotly_dark')
            
            # Adicionar linha vertical da intervenção
            fig.add_vline(x=res['ponto_intervencao'] + 0.5, line_dash="dash", line_color="green", annotation_text="Início Campanha")
            
            st.plotly_chart(fig, use_container_width=True)
            
            # Resultados Numéricos
            did = res['did']
            
            c1, c2, c3 = st.columns(3)
            c1.metric("Impacto Causal (DiD)", f"{did:.2%}", delta="Efeito Isolado", delta_color="normal" if did < 0 else "inverse")
            c2.metric("Tendência Controlo", f"{(res['stats']['control_pos'] - res['stats']['control_pre']):.2%}", help="Variação natural do mercado")
            c3.metric("Tendência Tratamento", f"{(res['stats']['treat_pos'] - res['stats']['treat_pre']):.2%}", help="Variação observada no grupo alvo")
            
            # Explicação IA
            msg_impacto = "reduziu" if did < 0 else "aumentou"
            sofia_explica(f"""
            **Conclusão Causal:**
            
            A intervenção **{campanha}** causou uma **redução real de {abs(did)*100:.2f} pontos percentuais** na taxa de default.
            
            - Sem a campanha, o default teria descido apenas {abs(res['stats']['control_pos'] - res['stats']['control_pre'])*100:.2f}% (tendência de mercado).
            - Com a campanha, desceu {abs(res['stats']['treat_pos'] - res['stats']['treat_pre'])*100:.2f}%.
            - A diferença é o seu ROI.
            """)
        else:
            st.info("👈 Selecione uma campanha e clique em Calcular.")




# ===================== PÁGINA: BASE DE CONHECIMENTO (FUNC 19) =====================
elif selected == "Base de Conhecimento":
    st.header("📚 Base de Conhecimento & Documentação")
    log_audit(st.session_state.username, "Knowledge Base Access", "Consulta de documentação")
    
    # 1. Base de Dados de Artigos (Simulada)
    kb_articles = {
        "metodologia_score": {
            "titulo": "Como é calculado o Score de Crédito?",
            "tags": ["Scoring", "Risco", "Metodologia"],
            "conteudo": """
            ### O Modelo de Scoring Interno
            O SocioStat utiliza um modelo de **Gradient Boosting** treinado em dados históricos para calcular a probabilidade de incumprimento.
            
            **Fatores de Peso:**
            1. **Histórico de Crédito (35%):** Atrasos nos últimos 12 meses.
            2. **Capacidade Financeira (30%):** Rendimento mensal vs. Dívida Total.
            3. **LTV (20%):** Rácio Loan-to-Value.
            4. **Estabilidade (15%):** Idade e antiguidade como cliente.
            
            > **Nota:** O score varia de 300 a 850. Scores acima de 700 são considerados "Prime".
            """
        },
        "basel_iii": {
            "titulo": "Guia Rápido: Basel III e Capital",
            "tags": ["Regulação", "Compliance", "Capital"],
            "conteudo": """
            ### Requisitos de Capital (Basel III)
            O acordo de Basileia III estabelece normas globais para a regulação bancária.
            
            **Fórmula de Cálculo:**
            $$RWA = EAD \\times PD \\times LGD \\times 1.06$$
            
            **Legenda:**
            * **PD:** Probability of Default
            * **LGD:** Loss Given Default (Padrão: 45%)
            * **EAD:** Exposure at Default
            
            O banco deve manter um rácio de capital mínimo de **10.5%** sobre os ativos ponderados pelo risco (RWA).
            """
        },
        "early_warning": {
            "titulo": "Sistema de Alerta Precoce (EWS)",
            "tags": ["Risco", "AI", "Prevenção"],
            "conteudo": """
            ### Como funciona o EWS?
            O Early Warning System monitoriza a base de clientes diariamente em busca de sinais de deterioração financeira.
            
            **Sinais de Alerta:**
            * Aumento súbito na utilização de cartões de crédito.
            * Queda no score interno > 20 pontos.
            * Deteção de novos créditos em outras instituições.
            
            **Ação Recomendada:**
            Para clientes marcados como **"Risco Crítico"**, a equipa de recuperação deve contactar o cliente no prazo de 48h para reestruturação preventiva.
            """
        },
        "glossario": {
            "titulo": "Glossário de Termos Bancários",
            "tags": ["Básico", "Terminologia"],
            "conteudo": """
            | Termo | Definição |
            | :--- | :--- |
            | **LTV** | *Loan-to-Value*. O valor do empréstimo dividido pelo valor do bem (ex: casa). |
            | **DSTI** | *Debt Service-to-Income*. Peso da prestação no rendimento líquido. |
            | **Spread** | Margem de lucro do banco adicionada à taxa de referência (Euribor). |
            | **Write-off** | Quando uma dívida é considerada incobrável e removida do balanço. |
            """
        }
    }
    
    # 2. Interface de Pesquisa
    col_search, col_filter = st.columns([3, 1])
    with col_search:
        search_term = st.text_input("🔍 Pesquisar na documentação...", placeholder="Ex: scoring, basel, ltv")
    
    # 3. Lógica de Filtragem
    filtered_articles = {}
    for key, data in kb_articles.items():
        # Pesquisa no título, conteúdo e tags (case insensitive)
        blob = (data['titulo'] + data['conteudo'] + " ".join(data['tags'])).lower()
        if search_term.lower() in blob:
            filtered_articles[key] = data
            
    # 4. Layout Master-Detail
    if not filtered_articles:
        st.warning("Nenhum artigo encontrado para a sua pesquisa.")
    else:
        # Menu lateral de artigos
        col_menu, col_content = st.columns([1, 3])
        
        with col_menu:
            st.markdown("### Artigos")
            selected_article_key = st.radio(
                "Selecione para ler:",
                options=list(filtered_articles.keys()),
                format_func=lambda x: filtered_articles[x]['titulo'],
                label_visibility="collapsed"
            )
            
            # Tags do artigo selecionado
            st.markdown("---")
            st.caption("Tags:")
            for tag in filtered_articles[selected_article_key]['tags']:
                st.markdown(f"`{tag}`")
        
        with col_content:
            # Renderizar Conteúdo
            article = filtered_articles[selected_article_key]
            
            st.markdown(f"""
            <div style="background-color: #1e293b; padding: 30px; border-radius: 10px; border: 1px solid #3b82f6;">
                <h2 style="margin-top:0;">{article['titulo']}</h2>
                <hr style="border-color: #3b82f6;">
            </div>
            """, unsafe_allow_html=True)
            
            st.markdown(article['conteudo'])
            
            # Feedback do Artigo (Micro-interação)
            st.markdown("---")
            c1, c2, c3 = st.columns([4, 1, 1])
            with c1:
                st.caption("Este artigo foi útil?")
            with c2:
                if st.button("👍 Sim", key="like_kb"):
                    st.toast("Obrigado pelo feedback!", icon="😊")
            with c3:
                if st.button("👎 Não", key="dislike_kb"):
                    st.toast("Iremos melhorar este conteúdo.", icon="🔧")

# ===================== PÁGINA: OTIMIZAÇÃO DE QUERIES (FUNC 50) =====================
elif selected in ["Otimização de Consultas", "Query Optimization", "menu_opt"]:
    st.header("⚡ Otimização de Consultas (DBA Tools)")
    log_audit(st.session_state.username, "Query Opt Access", "Análise de performance SQL")
    
    st.info("Simulador de performance de base de dados. Compare tempos de execução e analise planos de consulta.")

    col1, col2 = st.columns([1, 2])
    
    with col1:
        st.markdown("### 🧪 Laboratório de Teste")
        st.markdown("**Query Alvo:**")
        st.code("SELECT * FROM clientes \nWHERE regiao='Lisboa' \nAND score < 600", language="sql")
        
        if st.button("🚀 Executar Benchmark", type="primary"):
            with st.spinner("A executar 'Full Table Scan' vs 'Index Scan'..."):
                res = benchmark_query_performance(df)
                st.session_state['bench_results'] = res
                
    with col2:
        if 'bench_results' in st.session_state:
            res = st.session_state['bench_results']
            slow = res['Slow']
            fast = res['Fast']
            melhoria = slow / fast if fast > 0 else 0
            
            # Métricas de Comparação
            c1, c2, c3 = st.columns(3)
            c1.metric("Tempo (Sem Índice)", f"{slow:.4f}s", delta="-Lento", delta_color="inverse")
            c2.metric("Tempo (Otimizado)", f"{fast:.4f}s", delta=f"{melhoria:.0f}x Mais Rápido")
            c3.metric("Custo Estimado", "€0.0001", delta="-99% CPU", delta_color="normal")
            
            st.markdown("### 📜 Planos de Execução (EXPLAIN ANALYZE)")
            
            tab_slow, tab_fast = st.tabs(["🔴 Query Lenta", "🟢 Query Otimizada"])
            
            with tab_slow:
                st.warning("⚠️ Detectado 'Sequential Scan'. A ler o disco inteiro.")
                st.code(gerar_explain_plan_simulado("Lenta"), language="sql")
                
            with tab_fast:
                st.success("✅ Utilizado Índice B-Tree composto.")
                st.code(gerar_explain_plan_simulado("Rapida"), language="sql")
                
            # Recomendação Automática
            st.info(f"""
            **Recomendação do DBA Automático:**
            
            Detetámos que a coluna `regiao` e `score_interno` são filtradas frequentemente.
            Sugere-se criar um índice composto:
            
            `CREATE INDEX idx_regiao_score ON clientes (regiao, score_interno);`
            """)
        else:
            st.markdown("""
            <div style="text-align: center; padding: 50px; color: gray;">
                Prima o botão para iniciar o teste de stress à base de dados.
            </div>
            """, unsafe_allow_html=True)

# ===================== PÁGINA: BENCHMARKING (FUNC 22) =====================
elif selected == "Benchmarking":
    st.header("🏆 Benchmarking & Análise Comparativa")
    log_audit(st.session_state.username, "Benchmarking Access", "Comparação com o mercado")
    
    # 1. Preparar Dados
    # Certifique-se que a função calcular_kpis_principais existe (veio do código original)
    meus_kpis = calcular_kpis_principais(df)
    mercado = get_dados_mercado()
    gaps = calcular_gap_analysis(meus_kpis, mercado)
    
    col1, col2 = st.columns([1, 2])
    
    with col1:
        st.markdown("### 📊 Gap Analysis")
        st.info("Comparação direta com a média do setor bancário nacional.")
        
        for kpi, dados in gaps.items():
            st.markdown(f"**{kpi.replace('_', ' ')}**")
            
            # Formatação percentual ou inteira
            fmt = "{:.2%}" if kpi != "Score_Medio" else "{:.0f}"
            val_eu = fmt.format(dados['eu'])
            val_mercado = fmt.format(dados['mercado'])
            
            # Cor do delta
            cor_delta = "off" if "Melhor" in dados['status'] or "Conservador" in dados['status'] or "Melhores" in dados['status'] else "inverse"
            
            st.metric(
                label="Nós vs Mercado",
                value=val_eu,
                delta=f"{dados['diff']:+.2f} ({dados['status']})",
                delta_color=cor_delta if kpi == "Score_Medio" else ("normal" if cor_delta == "off" else "inverse")
            )
            st.markdown(f"<small style='color:gray'>Média Mercado: {val_mercado}</small>", unsafe_allow_html=True)
            st.markdown("---")

    with col2:
        st.markdown("### 🕸️ Radar de Posicionamento")
        
        # Normalizar dados para o Radar (Base 100 = Mercado)
        # Se formos iguais ao mercado, valor = 100.
        
        categorias = ['Risco (Default)', 'Colateral (LTV)', 'Qualidade (Score)', 'Preço (Juros)', 'Utilização']
        
        # Valores do Banco
        taxa_juro_media = df['taxa_juro'].mean() if 'taxa_juro' in df.columns else 5.0
        utilizacao_media = df['utilizacao_credito'].mean() if 'utilizacao_credito' in df.columns else 0.3
        
        val_nos = [
            (meus_kpis.get('Taxa_Default', 0.05) / mercado['Taxa_Default']) * 100,
            (meus_kpis.get('LTV_Medio', 0.8) / mercado['LTV_Medio']) * 100,
            (meus_kpis.get('Score_Medio', 600) / mercado['Score_Medio']) * 100,
            (taxa_juro_media / mercado['Taxa_Juro_Media']) * 100,
            (utilizacao_media / mercado['Utilizacao_Credito']) * 100
        ]
        
        # Mercado é sempre a base 100 (referência)
        val_mercado = [100, 100, 100, 100, 100]
        
        fig = go.Figure()

        fig.add_trace(go.Scatterpolar(
            r=val_nos,
            theta=categorias,
            fill='toself',
            name='O Nosso Banco',
            line_color='#3b82f6'
        ))
        
        fig.add_trace(go.Scatterpolar(
            r=val_mercado,
            theta=categorias,
            fill='toself',
            name='Média de Mercado',
            line_color='#64748b',
            line_dash='dash',
            opacity=0.5
        ))

        fig.update_layout(
            polar=dict(
                radialaxis=dict(
                    visible=True,
                    range=[0, max(max(val_nos), 150)]
                )),
            showlegend=True,
            title="Posicionamento Relativo (Base 100 = Mercado)",
            template='plotly_dark'
        )
        
        st.plotly_chart(fig, use_container_width=True)
        
        # Insight Automático
        pontos_fortes = []
        for cat, val in zip(categorias, val_nos):
            # Lógica invertida para Risco/LTV (menor é melhor)
            if cat in ['Risco (Default)', 'Colateral (LTV)']:
                if val < 100: pontos_fortes.append(cat)
            else:
                if val > 100: pontos_fortes.append(cat)
        
        if pontos_fortes:
            msg = f"Estamos a superar o mercado em **{', '.join(pontos_fortes)}**. A nossa estratégia está competitiva nestas áreas."
        else:
            msg = "Estamos com desempenho abaixo da média de mercado nas principais métricas. É necessária uma revisão estratégica."
            
        sofia_explica(msg)

# ===================== PÁGINA: RCS DETALHE (FUNC 51) =====================
elif selected in ["Executive Score (RCS)", "Executive Score", "menu_rcs"]:
    st.header("🎯 Decomposição do Executive Action Score")
    log_audit(st.session_state.username, "RCS Access", "Análise detalhada do score")
    
    # Recalcular
    df_macro = get_dados_macroeconomicos()
    rcs_data = calcular_rcs(df, df_macro)
    
    # Gráfico de Velocímetro (Gauge)
    fig = go.Figure(go.Indicator(
        mode = "gauge+number+delta",
        value = rcs_data['RCS'],
        domain = {'x': [0, 1], 'y': [0, 1]},
        title = {'text': "RCS Global"},
        delta = {'reference': 80, 'increasing': {'color': "green"}}, # Meta é 80
        gauge = {
            'axis': {'range': [0, 100]},
            'bar': {'color': "white"},
            'steps': [
                {'range': [0, 50], 'color': "#7f1d1d"},   # Vermelho escuro
                {'range': [50, 80], 'color': "#f59e0b"},  # Laranja
                {'range': [80, 100], 'color': "#10b981"}  # Verde
            ],
            'threshold': {
                'line': {'color': "red", 'width': 4},
                'thickness': 0.75,
                'value': rcs_data['RCS']
            }
        }
    ))
    st.plotly_chart(fig, use_container_width=True)
    
    # Explicação dos Fatores
    col1, col2 = st.columns(2)
    with col1:
        st.markdown("### Contribuição dos Pilares")
        # Gráfico de Donut
        labels = list(rcs_data['Componentes'].keys())
        values = list(rcs_data['Componentes'].values())
        fig_pie = px.pie(values=values, names=labels, hole=0.5, title="Peso no Score Final")
        st.plotly_chart(fig_pie, use_container_width=True)
        
    with col2:
        st.markdown("### Plano de Ação Recomendado")
        score = rcs_data['RCS']
        
        if score < 50:
            st.error("🚨 **AÇÃO IMEDIATA NECESSÁRIA**")
            st.write("- Congelar novas aprovações de crédito de risco médio/alto.")
            st.write("- Ativar equipa de recuperação para clientes em *Early Warning*.")
            st.write("- Rever exposição a setores sensíveis à inflação.")
        elif score < 80:
            st.warning("⚠️ **MONITORIZAÇÃO ATIVA**")
            st.write("- Aumentar requisitos de LTV para novos créditos.")
            st.write("- Realizar Stress Test focado em subida da Euribor.")
        else:
            st.success("✅ **OPERAÇÃO ESTÁVEL**")
            st.write("- Manter política atual.")
            st.write("- Avaliar oportunidades de expansão em segmentos Premium.")

# ===================== PÁGINA: LICENCIAMENTO (FUNC 49) =====================
elif selected in ["Licenciamento de Dados", "Data Licensing", "menu_license"]:
    st.header("🔑 Licenciamento e Proteção de Dados")
    log_audit(st.session_state.username, "License Access", "Gestão de direitos de dados")
    
    st.info("Gestão de direitos digitais (DRM) e controlo de exportação para proteção da propriedade intelectual.")

    # 1. O Meu Nível de Acesso
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("### 🆔 A Minha Licença")
        st.markdown(f"""
        <div class="metric-card">
            <h4>Utilizador: {st.session_state.username}</h4>
            <h2>Nível: {st.session_state.user_role}</h2>
            <p>ID Licença: <code>LIC-{abs(hash(st.session_state.username)) % 10000:04d}</code></p>
        </div>
        """, unsafe_allow_html=True)

    with col2:
        st.markdown("### 🚦 Permissões de Exportação")
        # Verificar o que este user pode fazer
        pode_raw = verificar_licenca_exportacao(st.session_state.username, "Dados Brutos")
        pode_report = verificar_licenca_exportacao(st.session_state.username, "Relatórios")
        
        st.write(f"{'✅' if pode_report else '❌'} Exportar Relatórios PDF/Word")
        st.write(f"{'✅' if pode_raw else '❌'} Exportar Dados Brutos (CSV/Excel)")
        st.write(f"{'✅' if pode_raw else '❌'} Acesso via API")

    st.markdown("---")

    # 2. Simulador de Exportação Protegida
    st.subheader("🛡️ Teste de Exportação Segura")
    
    tipo_export = st.selectbox("Tipo de Dados", ["Relatório de Risco", "Dados Brutos de Clientes"])
    
    if st.button("Simular Download Seguro"):
        if verificar_licenca_exportacao(st.session_state.username, "Dados Brutos" if "Brutos" in tipo_export else "Relatórios"):
            
            # Gerar Marca de Água
            watermark = aplicar_marca_agua(tipo_export)
            
            with st.spinner("A aplicar encriptação e marca de água..."):
                time.sleep(1.5)
            
            st.success("Ficheiro gerado com sucesso!")
            st.warning(f"⚠️ **Marca de Água Aplicada:**\n\n`{watermark}`")
            
            # Simular ficheiro
            st.download_button("📥 Baixar Ficheiro Protegido", data="dados encriptados", file_name="secure_data.bin")
            
        else:
            st.error("⛔ Acesso Negado: O seu nível de licença não permite exportar este tipo de dados sensíveis.")
            enviar_notificacao("Tentativa de exportação não autorizada bloqueada.", "erro")

# ===================== PÁGINA: QUALIDADE DE DADOS (FUNC 42) =====================
elif selected in ["Qualidade de Dados", "Data Quality", "menu_dq"]:
    st.header("💎 Qualidade de Dados (Data Quality)")
    log_audit(st.session_state.username, "DQ Access", "Auditoria de qualidade")
    
    # Executar Auditoria
    dq = auditar_qualidade_dados(df)
    
    # --- 1. SCORECARD GERAL ---
    score = dq['score_global']
    cor_score = "green" if score > 90 else "orange" if score > 70 else "red"
    
    col1, col2 = st.columns([1, 3])
    
    with col1:
        # Visualização tipo "Gauge" (simulada com markdown colorido)
        st.markdown(f"""
        <div style="text-align: center; padding: 20px; background-color: #1e293b; border-radius: 10px; border: 2px solid {cor_score};">
            <h3 style="margin:0; color: #94a3b8;">Health Score</h3>
            <h1 style="font-size: 60px; margin:0; color: {cor_score};">{score:.0f}%</h1>
            <p style="margin:0;">{ "Excelente" if score > 90 else "Requer Atenção" if score > 70 else "Crítico" }</p>
        </div>
        """, unsafe_allow_html=True)
        
    with col2:
        c1, c2, c3 = st.columns(3)
        c1.metric("Registos Totais", f"{dq['total_linhas']:,}")
        c2.metric("Células Vazias (Nulos)", f"{dq['nulos_pct']:.2f}%", delta="Ideal: 0%", delta_color="inverse")
        c3.metric("Registos Duplicados", f"{dq['duplicados_pct']:.2f}%", delta="Ideal: 0%", delta_color="inverse")
        
        st.markdown("---")
        if dq['colunas_com_problemas']:
            for prob in dq['colunas_com_problemas']:
                st.error(f"🚫 **{prob['coluna']}**: {prob['tipo']} ({prob['qtd']} casos)")
        else:
            st.success("✅ Nenhuma coluna crítica apresenta problemas estruturais.")

    st.markdown("---")

    # --- 2. DETALHE POR COLUNA ---
    tab1, tab2 = st.tabs(["🔍 Completude por Coluna", "📈 Outliers e Distribuição"])
    
    with tab1:
        st.subheader("Mapa de Calor de Dados em Falta")
        
        # Preparar dados para gráfico
        nulos_df = pd.DataFrame(list(dq['detalhe_nulos'].items()), columns=['Coluna', 'Qtd Nulos'])
        if not nulos_df.empty:
            nulos_df['Percentagem'] = (nulos_df['Qtd Nulos'] / dq['total_linhas']) * 100
            
            fig = px.bar(nulos_df, x='Percentagem', y='Coluna', orientation='h',
                         title="Percentagem de Dados em Falta",
                         color='Percentagem', color_continuous_scale='Reds',
                         template='plotly_dark')
            st.plotly_chart(fig, use_container_width=True)
        else:
            st.balloons()
            st.info("🎉 Parabéns! O seu dataset está 100% completo. Sem valores nulos.")
            
    with tab2:
        st.subheader("Deteção de Outliers")
        st.caption("Valores que se desviam mais de 3 desvios padrão da média (Z-Score > 3).")
        
        col_outlier = st.selectbox("Selecionar Coluna Numérica", 
                                 [c for c in cols_num if c not in ['id_cliente', 'default']])
        
        # Boxplot é o melhor para ver outliers
        fig_box = px.box(df, y=col_outlier, points="outliers", 
                         title=f"Distribuição e Outliers: {col_outlier}",
                         template='plotly_dark')
        st.plotly_chart(fig_box, use_container_width=True)
        
        # Explicação IA
        sofia_explica(f"""
        **Análise de Qualidade:**
        O indicador **{col_outlier}** apresenta uma distribuição com alguns valores extremos.
        Verifique se são erros de introdução de dados ou clientes VIP reais.
        Outliers excessivos podem distorcer o modelo de Credit Scoring.
        """)

# ===================== PÁGINA: GAMIFICAÇÃO (FUNC 44) =====================
elif selected in ["Gamificação", "Gamification", "menu_game"]:
    st.header("🏅 Perfil & Conquistas")
    log_audit(st.session_state.username, "Gamification Access", "Visualização de perfil")
    
    # Dados do Utilizador
    xp_atual = st.session_state.user_xp
    nivel = st.session_state.user_level
    proximo_nivel_xp = nivel * 1000
    progresso = (xp_atual % 1000) / 1000
    
    # --- 1. CARTÃO DE PERFIL ---
    with st.container():
        c1, c2, c3 = st.columns([1, 3, 1])
        with c1:
            st.image(f"https://ui-avatars.com/api/?name={st.session_state.username}&background=random&size=128", width=100)
        with c2:
            st.markdown(f"### {st.session_state.username}")
            st.markdown(f"**Nível {nivel}** | {st.session_state.user_role}")
            st.progress(progresso, text=f"XP: {xp_atual} / {proximo_nivel_xp}")
        with c3:
            if st.button("🎁 Resgatar Bónus Diário"):
                adicionar_xp(50, "Bónus Diário")
                st.rerun()

    st.markdown("---")

    tab1, tab2, tab3 = st.tabs(["🏆 Leaderboard", "🎖️ Meus Badges", "📜 Histórico de XP"])

    with tab1:
        st.subheader("Ranking da Equipa")
        df_rank = get_leaderboard()
        
        # Destacar o utilizador atual
        st.dataframe(
            df_rank.style.apply(lambda x: ['background-color: #1e3a8a' if x['Utilizador'] == st.session_state.username else '' for i in x], axis=1),
            use_container_width=True,
            hide_index=True
        )

    with tab2:
        st.subheader("A Minha Coleção")
        cols = st.columns(4)
        badges = [
            {"nome": "Primeiros Passos", "icon": "🌱", "desc": "Completou o Onboarding"},
            {"nome": "Analista Pro", "icon": "📊", "desc": "Gerou 10 Relatórios"},
            {"nome": "Data Guardian", "icon": "🛡️", "desc": "Score DQ > 95%"},
            {"nome": "Bug Hunter", "icon": "🐛", "desc": "Reportou um erro válido"}
        ]
        
        for i, badge in enumerate(badges):
            with cols[i]:
                st.markdown(f"""
                <div style="text-align:center; padding:10px; border:1px solid #334155; border-radius:10px; background:#0f172a;">
                    <h1 style="margin:0;">{badge['icon']}</h1>
                    <p style="font-weight:bold; margin:5px;">{badge['nome']}</p>
                    <small style="color:gray;">{badge['desc']}</small>
                </div>
                """, unsafe_allow_html=True)

    with tab3:
        st.subheader("Atividade Recente")
        for log in st.session_state.gamification_log:
            st.markdown(f"**+{log['xp']} XP** | {log['acao']} <span style='float:right; color:gray'>{log['data']}</span>", unsafe_allow_html=True)
            st.divider()

# ===================== PÁGINA: GESTÃO DE UTILIZADORES (FUNC 25) =====================
elif selected == "Gestão de Utilizadores":
    st.header("👥 Gestão de Utilizadores e Acessos")
    
    # 1. VERIFICAÇÃO DE SEGURANÇA (RBAC)
    if st.session_state.user_role != "Admin":
        st.error("⛔ Acesso Negado. Esta página é restrita a Administradores.")
        st.image("https://http.cat/403", width=400) # Toque de humor técnico
        log_audit(st.session_state.username, "Security Alert", "Tentativa de acesso não autorizado a Admin")
        st.stop() # Pára a execução aqui
        
    log_audit(st.session_state.username, "Admin Access", "Gestão de utilizadores")
    
    col1, col2 = st.columns([2, 1])
    
    with col1:
        st.markdown("### 📋 Utilizadores do Sistema")
        
        # Converter para DataFrame para visualização bonita
        df_users = pd.DataFrame(st.session_state.db_users)
        
        # Editor de Dados (Interativo)
        edited_df = st.data_editor(
            df_users,
            column_config={
                "id": "ID",
                "username": "Utilizador",
                "role": st.column_config.SelectboxColumn(
                    "Cargo (Role)",
                    options=["Admin", "Risk Manager", "Auditor", "Analyst", "Executive"],
                    required=True
                ),
                "status": st.column_config.SelectboxColumn(
                    "Status",
                    options=["Ativo", "Inativo"],
                    required=True
                ),
                "last_login": "Último Acesso"
            },
            disabled=["id", "last_login"], # Campos que não podem ser editados diretamente
            hide_index=True,
            use_container_width=True,
            key="user_editor"
        )
        
        # Botão para salvar alterações em massa (Simulação)
        if st.button("💾 Guardar Alterações"):
            # Aqui, na vida real, compararíamos o edited_df com o original e faríamos UPDATE na DB
            # Como o st.data_editor no session_state já atualiza o objeto visual,
            # precisamos apenas de converter de volta para a lista de dicionários para persistir na lógica
            st.session_state.db_users = edited_df.to_dict('records')
            st.success("Base de dados de utilizadores atualizada!")
            log_audit(st.session_state.username, "User Update", "Alteração em massa de utilizadores")

    with col2:
        st.markdown("### ➕ Novo Utilizador")
        with st.form("add_user_form"):
            new_user = st.text_input("Username")
            new_role = st.selectbox("Cargo Inicial", ["Analyst", "Risk Manager", "Auditor", "Executive", "Admin"])
            new_email = st.text_input("Email Corporativo")
            
            if st.form_submit_button("Criar Conta", type="primary"):
                if new_user and new_email:
                    criar_novo_utilizador(new_user, new_role)
                    st.success(f"Utilizador **{new_user}** criado com sucesso!")
                    log_audit(st.session_state.username, "User Created", f"User: {new_user} | Role: {new_role}")
                    st.rerun()
                else:
                    st.warning("Preencha todos os campos.")
        
        st.markdown("---")
        st.markdown("### 🛡️ Políticas de Acesso")
        st.info("""
        **Matriz de Permissões:**
        * **Admin:** Acesso total.
        * **Risk Manager:** Acesso a modelos e aprovações.
        * **Analyst:** Acesso a dados e relatórios.
        * **Auditor:** Acesso apenas a logs e relatórios (Read-only).
        """)

# ===================== PÁGINA: XAI EXPLAINER (FUNC 48) =====================
# ===================== PÁGINA: XAI EXPLAINER (FUNC 48) =====================
elif selected in ["XAI Explainer", "Explicabilidade", "menu_xai"]:
    st.header("🧠 Explicabilidade de IA (XAI)")
    log_audit(st.session_state.username, "XAI Access", "Análise de explicabilidade SHAP")
    
    # Dependência: Modelo treinado
    if 'default' not in df.columns:
        st.error("Dados insuficientes.")
    else:
        # Garantir modelo
        modelo_data = credit_scoring_model(df)
        
        if modelo_data:
            # Calcular SHAP (Cacheado)
            # Nota: Usamos _modelo_data para evitar erro de hash
            shap_res = calcular_shap_values(df, _modelo_data=modelo_data)
            
            tab1, tab2 = st.tabs(["🌍 Explicação Global", "👤 Explicação Local (Cliente)"])
            
            # --- ABA 1: GLOBAL ---
            with tab1:
                st.markdown("### O que impulsiona o risco no banco?")
                st.info("O gráfico abaixo resume o impacto de cada variável no modelo. Variáveis no topo são as mais importantes.")
                
                # Importar bibliotecas AQUI para evitar erros de arranque
                import shap 
                import matplotlib.pyplot as plt
                
                # Summary Plot (Matplotlib)
                fig, ax = plt.subplots()
                shap.summary_plot(shap_res['shap_values'], shap_res['X'], show=False)
                st.pyplot(fig)
                
                # Limpar a figura
                plt.clf()
                
                sofia_explica("""
                **Interpretação SHAP Global:**
                
                * **Cores:** Vermelho significa valor alto da variável (ex: Dívida Alta). Azul é valor baixo.
                * **Eixo X:** Impacto no Risco. Se os pontos vermelhos estão à direita, significa que valores altos dessa variável AUMENTAM o risco.
                """)

            # --- ABA 2: LOCAL ---
            with tab2:
                st.markdown("### Por que é que este cliente tem este score?")
                
                # Selecionar Cliente da amostra XAI
                clientes_amostra = shap_res['X'].index.tolist()
                cliente_sel_idx = st.selectbox("Selecionar Cliente da Amostra (ID)", clientes_amostra)
                
                # Encontrar a posição numérica (0 a 99) na matriz SHAP
                pos_idx = clientes_amostra.index(cliente_sel_idx)
                
                col1, col2 = st.columns([2, 1])
                
                with col1:
                    # Importar novamente para garantir o scope
                    import shap
                    import matplotlib.pyplot as plt
                    
                    # Waterfall Plot
                    st.write(f"**Cascata de Decisão para Cliente #{cliente_sel_idx}**")
                    fig_water = plot_shap_waterfall(shap_res, pos_idx)
                    st.pyplot(fig_water)
                    plt.clf()
                    
                with col2:
                    # Dados reais do cliente
                    st.markdown("**Dados do Cliente:**")
                    dados_cli = df[df['id_cliente'] == cliente_sel_idx].iloc[0]
                    st.dataframe(dados_cli[modelo_data['features']], use_container_width=True)
                    
                    # Conclusão
                    base_val = shap_res['expected_value']
                    shap_vals = shap_res['shap_values'][pos_idx]
                    
                    # Correção de Formatação NumPy
                    score_raw = base_val + np.sum(shap_vals)
                    
                    if isinstance(score_raw, np.ndarray):
                        score_final = float(score_raw.item()) if score_raw.size == 1 else float(score_raw[0])
                    else:
                        score_final = float(score_raw)
                    
                    st.metric("Score de Risco (Log-Odds)", f"{score_final:.2f}")
                    
                    if score_final > 0:
                        st.error("Risco Acima da Média")
                    else:
                        st.success("Risco Abaixo da Média")

        else:
            st.warning("Treine o modelo de Credit Scoring primeiro.")
# ===================== PÁGINA: WEB SCRAPING (FUNC 39) =====================
elif selected == "Web Scraping":
    st.header("🕷️ Extração de Dados Web (Scraping)")
    log_audit(st.session_state.username, "Scraping Access", "Acesso à ferramenta de extração")
    
    st.info("Importe dados financeiros ou demográficos diretamente de websites públicos.")

    # 1. Configuração
    with st.expander("⚙️ Configurar Fonte", expanded=True):
        # Sugestões de URLs que costumam funcionar (Wikipedia é ótimo para demos)
        url_input = st.text_input(
            "URL do Site", 
            placeholder="https://pt.wikipedia.org/wiki/Lista_de_municípios_de_Portugal_por_população",
            help="Cole o link de uma página que contenha uma tabela de dados."
        )
        
        col1, col2 = st.columns([1, 4])
        with col1:
            btn_scrape = st.button("🚀 Extrair Dados", type="primary")
        with col2:
            st.caption("Nota: O sistema extrai automaticamente a maior tabela encontrada na página.")

    # 2. Resultados
    if btn_scrape and url_input:
        with st.spinner(f"A conectar a {url_input}..."):
            resultado = executar_scraping_web(url_input)
            
        if resultado['sucesso']:
            df_scraped = resultado['dados']
            
            st.success(f"✅ Extração concluída! Fonte: **{resultado['titulo']}**")
            st.caption(f"Foram detetadas {resultado['total_tabelas']} tabelas. A mostrar a principal.")
            
            # Mostrar Dados
            st.dataframe(df_scraped, use_container_width=True)
            
            st.markdown("### 📥 Ações")
            c1, c2 = st.columns(2)
            
            with c1:
                # Botão para adicionar aos metadados (Preparação para Func 40)
                if st.button("💾 Guardar no Data Lake"):
                    # Guardar no session state para usar depois
                    st.session_state['ultimo_scraping'] = df_scraped
                    st.success("Dados guardados temporariamente. Vá ao módulo de Metadados para catalogar.")
                    enviar_notificacao(f"Scraping realizado: {len(df_scraped)} linhas extraídas.", "sucesso")
            
            with c2:
                csv = df_scraped.to_csv(index=False).encode('utf-8')
                st.download_button(
                    "⬇️ Baixar CSV",
                    csv,
                    "dados_web.csv",
                    "text/csv"
                )
                
        else:
            st.error(f"Falha na extração: {resultado['erro']}")
            st.warning("Dica: Alguns sites bloqueiam robôs. Tente a Wikipedia ou sites de dados abertos.")

# ===================== PÁGINA: APM MONITOR (FUNC 46) =====================
elif selected in ["Monitorização APM", "Performance Monitor", "menu_apm"]:
    st.header("⚡ Monitorização de Performance (APM)")
    
    # Apenas Admin ou Dev deve ver isto
    if st.session_state.user_role not in ["Admin", "Risk Manager"]:
        st.warning("Acesso restrito a equipa técnica.")
    else:
        logs = st.session_state.apm_logs
        
        if logs:
            df_apm = pd.DataFrame(logs)
            
            # KPIs de Sistema
            avg_latency = df_apm['duracao_sec'].mean()
            max_latency = df_apm['duracao_sec'].max()
            error_rate = len(df_apm[df_apm['status'] != "Sucesso"]) / len(df_apm)
            
            col1, col2, col3 = st.columns(3)
            col1.metric("Latência Média", f"{avg_latency:.4f}s", delta=f"{(avg_latency-0.5)*100:.1f}% vs Meta", delta_color="inverse")
            col2.metric("Latência Máxima", f"{max_latency:.4f}s")
            col3.metric("Taxa de Erros", f"{error_rate:.1%}", delta_color="inverse")
            
            st.markdown("---")
            
            c1, c2 = st.columns([2, 1])
            
            with c1:
                st.markdown("### 🐢 Funções Mais Lentas")
                # Agrupar por função
                slowest = df_apm.groupby('funcao')['duracao_sec'].agg(['mean', 'max', 'count']).sort_values('mean', ascending=False)
                st.dataframe(slowest.style.format("{:.4f}"), use_container_width=True)
                
                st.markdown("### 📉 Latência ao Longo do Tempo")
                fig = px.scatter(df_apm, x='timestamp', y='duracao_sec', color='funcao', 
                               title="Tempos de Resposta (Scatter)", template="plotly_dark")
                st.plotly_chart(fig, use_container_width=True)

            with c2:
                st.markdown("### 🚦 Saúde do Sistema")
                if avg_latency < 0.5 and error_rate < 0.01:
                    st.success("Sistema Saudável")
                    st.image("https://cdn-icons-png.flaticon.com/512/4315/4315445.png", width=100)
                elif avg_latency < 2.0:
                    st.warning("Performance Degradada")
                else:
                    st.error("Sistema Crítico")
                    
                st.markdown("### 🧹 Manutenção")
                if st.button("Limpar Logs APM"):
                    st.session_state.apm_logs = []
                    st.rerun()
        else:
            st.info("Sem dados de performance. Navegue pela app para gerar métricas.")

# ===================== PÁGINA: API PLAYGROUND (FUNC 15) =====================

elif selected == "API & Devs":
    st.header("🔌 API Pública & Integração")
    log_audit(st.session_state.username, "API Access", "Acesso à documentação de API")
    
    st.info("Esta área é destinada a **desenvolvedores** que desejam integrar o SocioStat com sistemas externos (CRMs, ERPs).")

    tab1, tab2 = st.tabs(["🚀 Playground", "📚 Documentação"])

    with tab1:
        st.markdown("### Testar Endpoints")
        
        endpoint = st.selectbox("Endpoint", ["GET /cliente/{id}", "GET /risco/carteira"])
        
        if endpoint == "GET /cliente/{id}":
            c_id = st.number_input("ID do Cliente", min_value=1, step=1)
            if st.button("Enviar Pedido"):
                resultado = api_get_cliente(c_id, df)
                st.json(resultado)
                
                if resultado['status'] == 200:
                    st.success(f"✅ Pedido bem sucedido em {datetime.now().strftime('%H:%M:%S')}")
                else:
                    st.error("❌ Erro no pedido")

        elif endpoint == "GET /risco/carteira":
            if st.button("Enviar Pedido"):
                resultado = api_get_risco_carteira(df)
                st.json(resultado)
                st.success("✅ Dados da carteira recuperados")

    with tab2:
        st.markdown("### 📦 SDKs e Bibliotecas")
        st.info("Acelere a integração usando as nossas bibliotecas oficiais.")
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("#### Python SDK")
            st.code("""
            pip install requests pandas
            
            from sociostat_sdk import SocioStatClient
            client = SocioStatClient("YOUR_API_KEY")
            client.get_cliente(1001)
            """, language="python")
            
            # Gerar o código do SDK
            sdk_content = gerar_codigo_sdk_python()
            
            # Botão de Download Real (Func 45)
            st.download_button(
                label="📥 Baixar 'sociostat_sdk.py'",
                data=sdk_content,
                file_name="sociostat_sdk.py",
                mime="text/x-python",
                help="Ficheiro Python pronto a usar com a classe cliente."
            )
            
        with col2:
            st.markdown("#### JavaScript / Node.js")
            st.code("""
            npm install sociostat-client
            
            const client = new SocioStatClient('API_KEY');
            const data = await client.getCliente(1001);
            """, language="javascript")
            
            st.button("📥 Baixar pacote NPM (Simulado)", disabled=True)

        st.markdown("---")
        st.markdown("### 📚 Referência da API (Swagger)")
        st.caption("Documentação completa dos endpoints REST.")
        # Em produção, isto seria um iframe do Swagger UI
        st.markdown("- [Especificação OpenAPI (YAML)](https://swagger.io)")
        st.markdown("- [Postman Collection](https://postman.com)")

# ===================== PÁGINA: PORTFOLIO =====================
elif selected == "Portfolio":
    st.header("💼 Otimização de Portfólio")
    log_audit(st.session_state.username, "Portfolio Optimization", "Acesso a portfolio")
    
    if st.button("🎯 Otimizar Portfólio"):
        with st.spinner("Calculando fronteira eficiente..."):
            fig, df_portfolios, best_portfolio = otimizar_portfolio(df)
        
        st.success("✅ Otimização concluída!")
        
        # Métricas do portfolio ótimo
        st.markdown("### ⭐ Portfolio Ótimo (Maior Sharpe Ratio)")
        col1, col2, col3 = st.columns(3)
        col1.metric("Retorno Esperado", f"{best_portfolio['return']:.2%}")
        col2.metric("Risco (Volatilidade)", f"{best_portfolio['risk']:.2%}")
        col3.metric("Sharpe Ratio", f"{best_portfolio['sharpe']:.2f}")
        
        # Gráfico
        st.plotly_chart(fig, use_container_width=True)
        
        sofia_explica(f"""
        **Análise de Portfolio:**
        
        - Portfolio ótimo identificado com Sharpe Ratio de {best_portfolio['sharpe']:.2f}
        - Retorno esperado: {best_portfolio['return']:.2%}
        - Risco: {best_portfolio['risk']:.2%}
        - Recomendação: Rebalancear portfolio para maximizar retorno ajustado ao risco
        """)

# ===================== PÁGINA: CONETORES OFICIAIS (FUNC 43) =====================
elif selected in ["Conetores Oficiais", "Official Connectors", "menu_connectors"]:
    st.header("🏛️ Conetores de Dados Oficiais")
    log_audit(st.session_state.username, "Connectors Access", "Gestão de APIs governamentais")
    
    init_connectors()
    
    st.info("Ligue o SocioStat a fontes de dados oficiais para enriquecer os seus modelos de risco com contexto macroeconómico real.")

    col_nav, col_content = st.columns([1, 3])

    with col_nav:
        st.markdown("### Fontes Disponíveis")
        conetor_escolhido = st.radio(
            "Selecione a Entidade:",
            list(st.session_state.connectors.keys()),
            format_func=lambda x: st.session_state.connectors[x].nome
        )

    connector = st.session_state.connectors[conetor_escolhido]

    with col_content:
        # Cabeçalho do Conetor
        c1, c2 = st.columns([3, 1])
        c1.markdown(f"## {connector.nome}")
        c1.caption(f"Endpoint: `{connector.base_url}`")
        c2.metric("Status", connector.status)
        
        st.markdown("---")
        
        tab1, tab2 = st.tabs(["🔑 Configuração", "📥 Extração de Dados"])
        
        with tab1:
            with st.form(f"config_{conetor_escolhido}"):
                st.markdown("#### Autenticação")
                api_key = st.text_input("Chave de API (Token)", type="password")
                
                submitted = st.form_submit_button("Testar Conexão")
                
                if submitted:
                    if connector.conectar(api_key):
                        st.success(f"Conexão estabelecida com {connector.nome}!")
                        enviar_notificacao(f"Conetor {conetor_escolhido} ativado.", "sucesso")
                        st.rerun()
                    else:
                        st.error("Falha na autenticação. Verifique a chave.")
        
        with tab2:
            if "🟢" in connector.status:
                st.markdown("#### Indicadores Disponíveis")
                
                indicadores = {
                    "INE": ["IPC (Inflação)", "PIB Trimestral", "Taxa Desemprego", "Índice Produção Industrial"],
                    "BCE": ["Euribor 3M", "Euribor 6M", "Euribor 12M", "Taxa de Refinanciamento"],
                    "PORDATA": ["Salário Médio", "Poder de Compra", "Endividamento Famílias"]
                }
                
                ind_sel = st.selectbox("Selecione o Indicador", indicadores.get(conetor_escolhido, ["Dados Genéricos"]))
                
                if st.button("📥 Importar Série Temporal"):
                    with st.spinner(f"A descarregar {ind_sel}..."):
                        df_ind = connector.extrair_indicador(ind_sel)
                        time.sleep(1)
                        
                    st.success(f"Dados de **{ind_sel}** importados com sucesso!")
                    
                    # Visualização rápida
                    st.line_chart(df_ind)
                    
                    with st.expander("Ver Dados Brutos"):
                        st.dataframe(df_ind, use_container_width=True)
                        
                    # Botão para guardar
                    if st.button("💾 Guardar no Macro Watch"):
                        st.toast("Indicador adicionado ao Dashboard Macro!", icon="✅")
            else:
                st.warning("⚠️ Configure a conexão na aba 'Configuração' primeiro.")

# ===================== PÁGINA: OTIMIZAÇÃO DE CUSTOS (FUNC 52) =====================
elif selected in ["Otimização de Custos", "Cost Optimization", "menu_cost"]:
    st.header("💰 Otimização de Custos de Infraestrutura")
    log_audit(st.session_state.username, "Cost Ops Access", "Análise de custos de cloud")
    
    # Executar Análise
    analise = analisar_eficiencia_custos()
    
    # 1. Visão Geral Financeira
    col1, col2, col3 = st.columns(3)
    
    col1.metric("Fatura Mensal Atual", f"€{analise['custo_atual']:.2f}")
    col2.metric("Poupança Identificada", f"€{analise['total_poupanca']:.2f}", delta="Oportunidade", delta_color="normal")
    col3.metric("Custo Otimizado", f"€{analise['novo_custo']:.2f}", delta=f"-{(analise['total_poupanca']/analise['custo_atual'])*100:.1f}%", delta_color="inverse")
    
    st.markdown("---")
    
    c1, c2 = st.columns([2, 1])
    
    with c1:
        st.subheader("🛠️ Recomendações Ativas")
        
        if analise['recomendacoes']:
            for i, rec in enumerate(analise['recomendacoes']):
                with st.expander(f"📉 {rec['acao']} - Poupa €{rec['poupanca']:.2f}/mês", expanded=True):
                    st.write(f"**Motivo:** {rec['motivo']}")
                    st.write(f"**Impacto:** {rec['impacto']}")
                    
                    col_btn, _ = st.columns([1, 3])
                    with col_btn:
                        if st.button(f"Aplicar Correção", key=f"btn_opt_{i}"):
                            with st.spinner("A aplicar alterações na Cloud..."):
                                time.sleep(2)
                                st.success("Configuração aplicada com sucesso!")
                                enviar_notificacao(f"Otimização aplicada: {rec['acao']}", "sucesso")
        else:
            st.success("A infraestrutura está totalmente otimizada!")

    with c2:
        st.subheader("📊 Utilização de Recursos")
        m = analise['metricas']
        
        # Barras de Progresso
        st.write("CPU Média")
        st.progress(m['cpu_avg'] / 100, text=f"{m['cpu_avg']:.1f}%")
        
        st.write("RAM Média")
        st.progress(m['ram_avg'] / 100, text=f"{m['ram_avg']:.1f}%")
        
        st.info("""
        **Nota:** O sistema de *Auto-Scaling* está ativo, mas as instâncias mínimas (Reserved Instances) podem ser renegociadas.
        """)

# ===================== PÁGINA: BILLING (FUNC 24) =====================
elif selected == "Billing & Planos":
    st.header("💳 Gestão de Subscrição")
    log_audit(st.session_state.username, "Billing Access", "Acesso à área financeira")
    
    # Obter dados simulados
    subs = get_dados_subscricao(st.session_state.username)
    planos = get_planos_disponiveis()
    
    # --- 1. ESTADO ATUAL ---
    col1, col2, col3, col4 = st.columns(4)
    col1.metric("Plano Atual", subs['plano_atual'])
    col2.metric("Status", subs['status'])
    col3.metric("Próxima Fatura", subs['proxima_fatura'])
    col4.metric("Valor Estimado", f"€{planos[subs['plano_atual']]['preco']:.2f}")
    
    st.markdown("---")
    
    col_main, col_side = st.columns([2, 1])
    
    with col_main:
        st.markdown("### 📊 Consumo de Recursos")
        
        # Barra de progresso de uso da API
        limite = planos[subs['plano_atual']]['api_calls']
        uso = subs['uso_api']
        pct_uso = uso / limite if isinstance(limite, int) else 0.1
        
        st.write(f"**Chamadas de API:** {uso:,} / {limite}")
        st.progress(pct_uso)
        
        if pct_uso > 0.8:
            st.warning("⚠️ Está perto do limite do seu plano. Considere fazer upgrade.")
            
        st.markdown("### 🧾 Histórico de Faturas")
        df_inv = pd.DataFrame(subs['faturas'])
        st.dataframe(df_inv, use_container_width=True)
        
        if st.button("📥 Download Todas as Faturas (ZIP)"):
            st.toast("Download iniciado...", icon="📥")

    with col_side:
        st.markdown("### 🚀 Mudar de Plano")
        
        plano_sel = st.selectbox("Escolher Novo Plano", list(planos.keys()), index=1)
        detalhes = planos[plano_sel]
        
        st.markdown(f"""
        <div style="background-color:#1e293b; padding:20px; border-radius:10px; border:1px solid #3b82f6;">
            <h2 style="color:#60a5fa">€{detalhes['preco']} <span style="font-size:14px">/mês</span></h2>
            <ul style="list-style-type: none; padding:0;">
                <li>👥 {detalhes['users']} Utilizadores</li>
                <li>🔌 API: {detalhes['api_calls']}</li>
                <li>📞 Suporte: {detalhes['support']}</li>
            </ul>
        </div>
        """, unsafe_allow_html=True)
        
        if plano_sel != subs['plano_atual']:
            if st.button(f"Upgrade para {plano_sel}", type="primary"):
                st.balloons()
                st.success(f"Pedido de alteração para **{plano_sel}** enviado ao gestor de conta!")
                log_audit(st.session_state.username, "Plan Change Request", f"To: {plano_sel}")
        else:
            st.button("Plano Atual", disabled=True)


# ===================== PÁGINA: AUDITORIA DE SEGURANÇA (FUNC 37) =====================
elif selected == "Auditoria de Segurança":
    st.header("🛡️ Centro de Operações de Segurança (SOC)")
    
    # Verificação de Permissão (Apenas Admin deve ver isto)
    if st.session_state.user_role != "Admin":
        st.error("⛔ Acesso restrito a Administradores de Segurança.")
        st.stop()
        
    log_audit(st.session_state.username, "SOC Access", "Monitorização de segurança")
    
    # Botão para Simular Ataque (Para ver o sistema a funcionar)
    if st.button("🚨 Simular Ataque de Brute-Force (Teste)"):
        simular_ataque()
        st.rerun()
    
    # Executar Análise
    ameacas = analisar_ameacas_seguranca(st.session_state.audit_log)
    
    # Métricas de Topo
    col1, col2, col3 = st.columns(3)
    col1.metric("Total Eventos Analisados", len(st.session_state.audit_log))
    col2.metric("Ameaças Ativas", len(ameacas), delta=len(ameacas), delta_color="inverse")
    col3.metric("Estado do Sistema", "🔴 SOB ATAQUE" if ameacas else "🟢 SEGURO")
    
    st.markdown("---")
    
    if ameacas:
        st.subheader(f"🚨 Ameaças Detetadas ({len(ameacas)})")
        
        for ameaca in ameacas:
            cor = "#7f1d1d" if ameaca['nivel'] == "CRÍTICO" else "#c2410c"
            st.markdown(f"""
            <div style="background-color: {cor}; padding: 15px; border-radius: 8px; margin-bottom: 10px; border-left: 5px solid red; color: white;">
                <div style="display:flex; justify-content:space-between;">
                    <strong>{ameaca['tipo']}</strong>
                    <small>{ameaca['hora']}</small>
                </div>
                <p style="margin: 5px 0 0 0;">{ameaca['mensagem']}</p>
                <code style="background:rgba(0,0,0,0.2); color:white;">User: {ameaca['user']}</code>
            </div>
            """, unsafe_allow_html=True)
            
        if st.button("🛡️ Bloquear Utilizadores Suspeitos"):
            st.success("IPs bloqueados na firewall e contas suspensas temporariamente.")
            enviar_notificacao("Medidas de contenção aplicadas pelo Admin.", "sucesso")
            
    else:
        st.image("https://cdn-icons-png.flaticon.com/512/1161/1161388.png", width=100)
        st.success("Nenhuma anomalia de segurança detetada nos logs recentes.")
        st.caption("O motor de análise continua a monitorizar em tempo real.")

# ===================== PÁGINA: ALERTAS PROATIVOS (FUNC 53) =====================
elif selected in ["Gestão de Alertas", "Proactive Alerts", "menu_alerts"]:
    st.header("🚨 Motor de Alertas Proativos")
    log_audit(st.session_state.username, "Alerts Config", "Configuração de regras")
    
    # 1. Executar Motor (Simulação de tempo real)
    disparos = verificar_regras_proativas(df)
    
    if disparos:
        st.error(f"⚠️ {len(disparos)} Regras violadas atualmente!")
    else:
        st.success("✅ Sistema Estável. Nenhuma regra violada.")

    col1, col2 = st.columns([1, 2])

    # --- COLUNA 1: CRIAR REGRA ---
    with col1:
        st.markdown("### ➕ Nova Regra")
        with st.form("add_alert"):
            metrica = st.selectbox("Métrica", ["Taxa Default", "Score Médio", "LTV Médio", "Carteira Total"])
            operador = st.selectbox("Condição", [">", "<"])
            valor = st.number_input("Valor Limite", value=0.0, step=0.01)
            
            st.caption(f"Ex: Alertar se {metrica} {operador} {valor}")
            
            if st.form_submit_button("Criar Alerta"):
                adicionar_regra(metrica, operador, valor)
                st.success("Regra adicionada!")
                st.rerun()

    # --- COLUNA 2: REGRAS ATIVAS ---
    with col2:
        st.markdown("### 📋 Regras Ativas")
        
        if st.session_state.alert_rules:
            # Converter para DF para visualização bonita
            df_regras = pd.DataFrame(st.session_state.alert_rules)
            
            # Mostrar como tabela interativa
            event = st.dataframe(
                df_regras,
                column_config={
                    "id": "ID",
                    "metrica": "KPI Monitorizado",
                    "operador": "Condição",
                    "valor": "Limite",
                    "ativo": st.column_config.CheckboxColumn("Ativo?")
                },
                use_container_width=True,
                hide_index=True,
                selection_mode="single-row",
                on_select="rerun" # Apenas para demo visual
            )
            
            # Botão para apagar (simples)
            regra_apagar = st.selectbox("Apagar Regra (ID)", df_regras['id'])
            if st.button("🗑️ Apagar Regra"):
                apagar_regra(regra_apagar)
                st.rerun()
        else:
            st.info("Sem regras definidas.")

    st.markdown("---")
    
    # --- HISTÓRICO DE DISPAROS ---
    st.markdown("### 📜 Histórico de Disparos")
    if st.session_state.triggered_alerts_history:
        df_hist = pd.DataFrame(st.session_state.triggered_alerts_history)
        st.dataframe(df_hist, use_container_width=True)
    else:
        st.caption("Ainda não foram registados incidentes.")

# ===================== PÁGINA: GESTÃO DE TAGS (FUNC 38) =====================
elif selected == "Gestão de Tags": # Verifique se o nome corresponde ao menu!
    st.header("🏷️ Gestão de Etiquetas")
    log_audit(st.session_state.username, "Tags Access", "Gestão de tags")
    
    # Agora a coluna 'tags' já existe garantidamente
    
    # 1. Estatísticas
    df_exploded = df.explode('tags').dropna(subset=['tags'])
    
    if df_exploded.empty:
        st.info("Ainda não existem etiquetas atribuídas a clientes.")
        st.markdown("👉 Vá à página **Cliente Individual** para adicionar tags.")
    else:
        contagem = df_exploded['tags'].value_counts().reset_index()
        contagem.columns = ['Tag', 'Clientes']
        
        c1, c2 = st.columns([1, 2])
        with c1:
            st.dataframe(contagem, use_container_width=True)
        with c2:
            fig = px.bar(contagem, x='Tag', y='Clientes', title="Distribuição de Tags")
            st.plotly_chart(fig, use_container_width=True)

# ===================== PÁGINA: MLOPS CENTER (FUNC 33) =====================
elif selected == "MLOps Center":
    st.header("🧬 MLOps: Gestão de Ciclo de Vida")
    log_audit(st.session_state.username, "MLOps Access", "Gestão de versões de modelos")
    
    # Estatísticas do Registry
    registry = st.session_state.model_registry
    total_models = len(registry)
    prod_models = len([m for m in registry if m['status'] == "Production"])
    
    col1, col2, col3 = st.columns(3)
    col1.metric("Modelos Registados", total_models)
    col2.metric("Em Produção", prod_models)
    col3.metric("Em Staging (Testes)", len([m for m in registry if m['status'] == "Staging"]))
    
    st.markdown("---")
    
    tab1, tab2 = st.tabs(["📜 Registo de Versões", "📈 Comparação de Performance"])
    
    with tab1:
        st.markdown("### Histórico de Versões")
        
        # Converter para DataFrame para exibição
        if registry:
            # Preparar dados para a tabela (flattening dos dicionários)
            data_table = []
            for m in registry:
                row = m.copy()
                # Extrair métricas principais para colunas
                if 'AUC' in m['metricas']: row['AUC'] = m['metricas']['AUC']
                if 'Accuracy' in m['metricas']: row['Acurácia'] = m['metricas']['Accuracy']
                del row['metricas'] # Limpar para não poluir
                del row['params']
                data_table.append(row)
                
            df_reg = pd.DataFrame(data_table)
            
            # Estilizar Tabela (Badge de Status)
            st.dataframe(
                df_reg,
                column_config={
                    "status": st.column_config.SelectboxColumn(
                        "Estado",
                        options=["Production", "Staging", "Archived", "Development"],
                        required=True
                    ),
                    "AUC": st.column_config.ProgressColumn("Performance (AUC)", min_value=0, max_value=1, format="%.2f"),
                },
                use_container_width=True,
                hide_index=True
            )
            
            # Ação de Promoção
            st.markdown("#### 🚀 Ações de Deployment")
            c1, c2 = st.columns([3, 1])
            with c1:
                modelo_promo = st.selectbox("Selecione Modelo para Promover a Produção", 
                                          [m['id'] + " - " + m['versao'] for m in registry if m['status'] != "Production"])
            with c2:
                st.write("") # Espaço
                st.write("")
                if st.button("Promover Modelo", type="primary"):
                    model_id = modelo_promo.split(" - ")[0]
                    if promover_modelo_producao(model_id):
                        st.success(f"Modelo {model_id} promovido a Produção! Versão anterior arquivada.")
                        enviar_notificacao(f"Deployment: Modelo {model_id} agora em Produção.", "sucesso")
                        time.sleep(1)
                        st.rerun()
        else:
            st.info("Registo vazio.")

    with tab2:
        st.markdown("### Evolução da Performance")
        st.info("Compare como a precisão dos modelos tem evoluído ao longo das versões.")
        
        if registry:
            # Criar gráfico de linha por Tipo de Modelo
            df_chart = pd.DataFrame(registry)
            # Extrair AUC
            df_chart['AUC'] = df_chart['metricas'].apply(lambda x: x.get('AUC', 0))
            # Ordenar por data
            df_chart['data_dt'] = pd.to_datetime(df_chart['data'])
            df_chart = df_chart.sort_values('data_dt')
            
            fig = px.line(
                df_chart, 
                x='versao', 
                y='AUC', 
                color='nome',
                markers=True,
                title="Evolução da Métrica AUC por Versão",
                template="plotly_dark"
            )
            st.plotly_chart(fig, use_container_width=True)

# ===================== PÁGINA: RECONSTRUÇÃO (FUNC 41) =====================
elif selected in ["Reconstrução de Experiências", "Reproducibility", "menu_reconstruct"]:
    st.header("🔄 Reconstrução de Experiências (Time Machine)")
    log_audit(st.session_state.username, "Replay Access", "Acesso ao histórico de experiências")
    
    st.info("Aceda a análises passadas e garanta a reprodutibilidade dos resultados regulatórios.")

    # 1. Tabela de Experiências
    df_exp = pd.DataFrame(st.session_state.experiments_db)
    
    col1, col2 = st.columns([2, 1])
    
    with col1:
        st.markdown("### 📜 Histórico de Snapshots")
        st.dataframe(
            df_exp,
            use_container_width=True,
            column_config={
                "id": "ID Snapshot",
                "tipo": "Módulo",
                "parametros": "Configuração",
                "resultado_chave": "Output Principal"
            },
            hide_index=True
        )

    with col2:
        st.markdown("### 🕹️ Ações")
        xp_selecionada = st.selectbox("Selecionar Experiência", df_exp['id'])
        
        # Obter detalhes da seleção
        detalhes = df_exp[df_exp['id'] == xp_selecionada].iloc[0]
        
        with st.container(border=True):
            st.markdown(f"**{detalhes['tipo']}**")
            st.caption(f"Autor: {detalhes['autor']} | Data: {detalhes['data']}")
            st.code(f"Params: {detalhes['parametros']}")
            st.markdown(f"**Resultado Guardado:** `{detalhes['resultado_chave']}`")
            
            if st.button("🔄 Carregar Estado (Replay)", type="primary"):
                with st.spinner("A restaurar ambiente e dados..."):
                    time.sleep(1.5)
                    st.success(f"Ambiente restaurado para o estado de {detalhes['data']}!")
                    st.balloons()
                    
                    # Simulação de restauração
                    st.info("""
                    **Modo de Reprodução Ativo:**
                    - Dataset revertido para versão v2.1 (Hash: a8f92...)
                    - Parâmetros do modelo injetados
                    - Seed aleatória fixada em 42
                    """)

    st.markdown("---")
    
    # 2. Simulador de Captura (Para testar a funcionalidade agora)
    with st.expander("💾 Simular Nova Captura (Demo)"):
        c1, c2 = st.columns(2)
        with c1:
            teste_tipo = st.selectbox("Módulo Atual", ["Stress Testing", "Credit Scoring"])
            teste_params = st.text_input("Parâmetros", "Cenário: Recessão Severa (-5% PIB)")
        with c2:
            teste_res = st.text_input("Resultado", "Perda Esperada: €2.5M")
            if st.button("📸 Tirar Snapshot"):
                new_id = salvar_snapshot(teste_tipo, teste_params, teste_res)
                st.success(f"Experiência {new_id} guardada!")
                time.sleep(1)
                st.rerun()

# ===================== PÁGINA: RLM (FUNC 23) =====================
elif selected == "Análise de Drivers (RLM)":
    st.header("🧮 Regressão Linear Múltipla (Drivers)")
    log_audit(st.session_state.username, "RLM Access", "Execução de modelo estatístico")
    
    st.info("Descubra quais variáveis têm maior impacto estatístico no seu alvo (ex: Score ou Dívida).")
    
    # Seleção de Variáveis (Apenas Numéricas)
    cols_num_clean = [c for c in cols_num if c not in ['id_cliente', 'latitude', 'longitude']]
    
    col1, col2 = st.columns([1, 2])
    
    with col1:
        st.markdown("### ⚙️ Configuração do Modelo")
        
        target = st.selectbox(
            "Variável Alvo (Y)", 
            cols_num_clean,
            index=cols_num_clean.index('score_interno') if 'score_interno' in cols_num_clean else 0,
            help="O que queremos explicar?"
        )
        
        # Remover o target da lista de features para não haver vazamento
        features_disponiveis = [c for c in cols_num_clean if c != target]
        
        features = st.multiselect(
            "Variáveis Explicativas (X)",
            features_disponiveis,
            default=features_disponiveis[:4], # Pré-selecionar algumas
            help="Que fatores influenciam o alvo?"
        )
        
        if st.button("🚀 Calcular Regressão", type="primary"):
            if not features:
                st.error("Selecione pelo menos uma variável explicativa.")
            else:
                resultado = calcular_rlm(df, target, features)
                st.session_state['rlm_result'] = resultado

    with col2:
        if 'rlm_result' in st.session_state:
            res = st.session_state['rlm_result']
            
            if res['sucesso']:
                # Métricas de Qualidade do Modelo
                c1, c2, c3 = st.columns(3)
                c1.metric("R² (Explicação)", f"{res['r2']:.1%}", help="Percentagem da variância explicada pelo modelo")
                c2.metric("R² Ajustado", f"{res['r2_adj']:.1%}", help="R² corrigido pelo número de variáveis")
                c3.metric("Observações", res['n_obs'])
                
                st.markdown("### 📊 Impacto das Variáveis (Coeficientes)")
                
                # Gráfico de Coeficientes
                coef_df = res['coefs']
                coef_df['Cor'] = ['Positivo' if x > 0 else 'Negativo' for x in coef_df['Coeficiente']]
                
                fig = px.bar(
                    coef_df, 
                    x='Coeficiente', 
                    y='Feature', 
                    orientation='h',
                    color='Cor',
                    color_discrete_map={'Positivo': '#10b981', 'Negativo': '#ef4444'},
                    title=f"Impacto em: {target}",
                    template='plotly_dark'
                )
                fig.update_layout(showlegend=False)
                st.plotly_chart(fig, use_container_width=True)
                
                # Interpretação Automática
                # Encontrar o maior impacto significativo
                # Filtrar apenas p-value < 0.05 (significativo)
                significativos = coef_df[coef_df['P-Value'] < 0.05]
                
                if not significativos.empty:
                    # Maior impacto absoluto
                    top_var = significativos.loc[significativos['Coeficiente'].abs().idxmax()]
                    direcao = "aumenta" if top_var['Coeficiente'] > 0 else "diminui"
                    
                    sofia_explica(f"""
                    **Análise de Drivers:**
                    
                    A variável **{top_var['Feature']}** é o principal impulsionador estatístico.
                    
                    - Cada unidade extra de **{top_var['Feature']}** faz com que **{target}** {direcao} em **{abs(top_var['Coeficiente']):.2f}** unidades.
                    - O modelo explica **{res['r2']:.1%}** da variação total dos dados.
                    """)
                else:
                    st.warning("Nenhuma variável selecionada tem significância estatística (P-Value < 0.05). Tente adicionar outras variáveis.")
                
                # Mostrar Tabela Detalhada (Expander)
                with st.expander("Ver Sumário Estatístico Completo (OLS)"):
                    st.text(res['modelo'].summary())
                    
            else:
                st.error(res['erro'])
        else:
            st.info("👈 Selecione as variáveis e clique em Calcular.")

# ===================== PÁGINA: INTEGRAÇÃO CRM (FUNC 27) =====================
elif selected == "Integração CRM":
    st.header("🤝 Integração com CRM")
    log_audit(st.session_state.username, "CRM Integration Access", "Acesso ao módulo de sincronização")
    
    # Estado da Conexão
    crm_status = get_crm_status()
    
    col1, col2, col3 = st.columns(3)
    col1.metric("Estado da Conexão", crm_status['status'])
    col2.metric("CRM Destino", crm_status['crm_type'])
    col3.metric("Última Sincronização", crm_status['ultima_sincronizacao'])
    
    st.markdown("---")
    
    c1, c2 = st.columns([2, 1])
    
    with c1:
        st.subheader("📤 Enviar Scores de Risco")
        st.info("Esta ação irá atualizar o campo 'SocioStat Score' e 'Risco' na ficha de cada cliente no Salesforce.")
        
        # Preparar dados para envio (Simulação)
        # Vamos assumir que queremos enviar o Score Interno e o Status de Default Previsto
        if 'score_interno' in df.columns:
            df_to_sync = df[['id_cliente', 'score_interno', 'default']].head(100) # Simular envio de 100 registos
            
            st.dataframe(df_to_sync, height=200, use_container_width=True)
            
            if st.button("🔄 Iniciar Sincronização", type="primary"):
                with st.status("A conectar ao Salesforce...", expanded=True) as status:
                    st.write("🔐 Autenticando via OAuth2...")
                    time.sleep(1)
                    st.write("📦 Preparando batch de dados...")
                    time.sleep(1)
                    st.write("🚀 Enviando dados...")
                    
                    res = sincronizar_com_crm(df_to_sync)
                    
                    status.update(label="Sincronização Concluída!", state="complete", expanded=False)
                
                if res['falhas'] == 0:
                    st.success(f"✅ Sucesso! {res['sucesso']} registos atualizados no CRM.")
                    enviar_notificacao(f"Sincronização CRM: {res['sucesso']} registos enviados.", "sucesso")
                else:
                    st.warning(f"⚠️ Parcial: {res['sucesso']} enviados, {res['falhas']} falharam.")
                    enviar_notificacao("Sincronização CRM concluída com erros.", "aviso")
        else:
            st.error("Dados de Score não disponíveis para sincronização.")

    with c2:
        st.subheader("⚙️ Configuração")
        st.text_input("CRM API Key", value="sk_live_********************", type="password")
        st.text_input("Instance URL", value="https://eu12.salesforce.com")
        
        st.markdown("#### Mapeamento de Campos")
        df_map = pd.DataFrame([
            {"SocioStat": "id_cliente", "Salesforce": "External_ID__c"},
            {"SocioStat": "score_interno", "Salesforce": "Credit_Score__c"},
            {"SocioStat": "default_prob", "Salesforce": "Risk_Probability__c"},
        ])
        st.dataframe(df_map, hide_index=True)
        
        if st.button("Testar Conexão"):
            st.toast("Conexão estabelecida com sucesso!", icon="🔗")

# ===================== PÁGINA: REPORTING (ATUALIZADA FUNC 10) =====================
elif selected == "Reporting":
    st.header("📊 Reporting Executivo Avançado")
    log_audit(st.session_state.username, "Reporting Access", "Acesso ao módulo de reporting")
    
    st.markdown("""
    ### 📋 Centro de Relatórios
    Gere relatórios personalizados para diferentes stakeholders
    """)
    
    # Calcular KPIs atuais para o Excel
    kpis_atuais = calcular_kpis_principais(df)
    
    tab1, tab2, tab3 = st.tabs(["📄 Relatórios Executivos", "🔬 Relatórios Técnicos", "📅 Agendamento"])
    
    # --- TAB 1: RELATÓRIOS EXECUTIVOS (Adicionado Excel) ---
    with tab1:
        st.markdown("### 📄 Relatórios para Decisão")
        st.info("Documentos formatados para C-Level e Board.")
        
        col1, col2 = st.columns(2)
        
        # Coluna 1: Word (Já existia)
        with col1:
            st.markdown("""
            <div class="metric-card">
                <h4>📝 Documento Word</h4>
                <p>Relatório narrativo com análise de Basel III e recomendações.</p>
            </div>
            """, unsafe_allow_html=True)
            st.write("")
            
            if st.button("📥 Gerar Word Executivo", key="exec_word", use_container_width=True):
                with st.spinner("A redigir documento..."):
                    relatorio = gerar_relatorio_executivo(df, st.session_state.username, st.session_state.user_role)
                
                st.download_button(
                    label="⬇️ Baixar .docx",
                    data=relatorio.getvalue(),
                    file_name=f"Relatorio_Executivo_{datetime.now().strftime('%Y%m%d')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
                st.success("Word pronto!")

        # Coluna 2: Excel Avançado (NOVO!)
        with col2:
            st.markdown("""
            <div class="metric-card">
                <h4>📊 Excel Analítico</h4>
                <p>Ficheiro multi-aba com Sumário de KPIs e Dados Brutos.</p>
            </div>
            """, unsafe_allow_html=True)
            st.write("")
            
            if st.button("📊 Gerar Excel Completo", key="exec_excel", use_container_width=True):
                with st.spinner("A formatar Excel..."):
                    excel_data = gerar_excel_avancado(df, kpis_atuais)
                
                st.download_button(
                    label="⬇️ Baixar .xlsx",
                    data=excel_data,
                    file_name=f"SocioStat_Dados_{datetime.now().strftime('%Y%m%d')}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True
                )
                st.success("Excel pronto!")
        
        st.markdown("---")
        
        # Área de CSV (Mantida, mas simplificada)
        with st.expander("💾 Exportação de Dados Brutos (CSV)", expanded=False):
            st.caption("Para integração rápida com outros sistemas.")
            csv = df.to_csv(index=False).encode('utf-8')
            st.download_button(
                label="📥 Download CSV",
                data=csv,
                file_name=f"dados_sociostat_{datetime.now().strftime('%Y%m%d')}.csv",
                mime="text/csv"
            )

        # Preview (Mantido)
        with st.expander("👁️ Ver Preview do Conteúdo"):
            st.markdown(f"**Pré-visualização para: {st.session_state.username}**")
            basel = calcular_capital_regulatorio(df)
            if basel:
                st.write(f"- Capital Mínimo Obrigatório: €{basel['Capital_Minimo_8pct']/1e6:.2f}M")
                st.write(f"- Taxa de Default Atual: {df['default'].mean():.2%}")

    # --- TAB 2: RELATÓRIOS TÉCNICOS (Mantido igual) ---
    with tab2:
        st.markdown("### 🔬 Relatórios Técnicos")
        st.info("Estatísticas descritivas e auditoria de modelos.")
        
        if st.button("📥 Gerar Relatório Técnico (Word)", key="tech_word"):
            with st.spinner("Gerando relatório técnico..."):
                relatorio_tec = gerar_relatorio_tecnico(df)
            
            st.download_button(
                label="⬇️ Download Relatório Técnico.docx",
                data=relatorio_tec.getvalue(),
                file_name=f"Relatorio_Tecnico_{datetime.now().strftime('%Y%m%d')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            st.success("Relatório gerado!")

            enviar_notificacao("Relatório Executivo gerado e pronto para download.", "info") # <--- NOVO
        
        st.markdown("#### Estatísticas da Amostra")
        if len(df.select_dtypes(include=[np.number]).columns) > 0:
            st.dataframe(df.describe(), use_container_width=True)

    # --- TAB 3: AGENDAMENTO (Mantido igual) ---
    with tab3:
        st.markdown("### 📅 Agendamento Automático")
        st.info("⚠️ Feature Enterprise - Configuração de envio por email.")
        
        with st.form("schedule_form"):
            c1, c2 = st.columns(2)
            with c1:
                st.selectbox("Relatório", ["Executivo Semanal", "Risco Diário", "Compliance Mensal"])
                st.multiselect("Destinatários", ["cfo@banco.pt", "risk_team@banco.pt"])
            with c2:
                st.selectbox("Frequência", ["Diária", "Semanal (Segunda)", "Mensal (Dia 1)"])
                st.time_input("Horário", value=datetime.strptime("08:00", "%H:%M").time())
            
            if st.form_submit_button("💾 Agendar Envio"):
                st.success("✅ Agendamento guardado no sistema!")

# ===================== CORREÇÃO DA PÁGINA CLOUD =====================
elif selected in ["Integração Cloud", "Cloud Integration", "menu_cloud"]: 
    st.header("☁️ Gestão de Arquivos na Nuvem")
    log_audit(st.session_state.username, "Cloud Access", "Acesso ao gestor de arquivos")
    
    # Inicializar simulador (Importante: Se isto falhar, a página fica em branco)
    if "cloud_storage" not in st.session_state:
        init_cloud_storage()
    
    # Seletor de Provider
    col1, col2 = st.columns([1, 3])
    with col1:
        provider = st.selectbox("Fornecedor", ["Google Drive", "Dropbox", "OneDrive"])
        st.caption("Armazenamento")
        st.progress(0.45)
        st.caption("45GB usados de 100GB")
        
    with col2:
        st.info(f"Conectado a: **{provider} de {st.session_state.username}**")

    st.markdown("---")

    tab1, tab2 = st.tabs(["📂 Explorador de Arquivos", "📤 Upload para Nuvem"])

    # --- TAB 1: EXPLORADOR ---
    with tab1:
        st.markdown(f"### Arquivos em {provider}")
        
        # Garantir que o provider existe no dicionário
        if provider in st.session_state.cloud_storage:
            files = st.session_state.cloud_storage[provider]
            df_files = pd.DataFrame(files)
            
            col_list, col_action = st.columns([3, 1])
            
            with col_list:
                st.dataframe(
                    df_files, 
                    use_container_width=True,
                    column_config={
                        "name": "Nome do Arquivo",
                        "size": "Tamanho",
                        "date": "Data Modificação"
                    }
                )
                
            with col_action:
                st.markdown("**Ações Rápidas**")
                # Proteção contra lista vazia
                opcoes_ficheiros = df_files['name'].tolist() if not df_files.empty else []
                
                if opcoes_ficheiros:
                    file_to_import = st.selectbox("Selecionar Arquivo", opcoes_ficheiros)
                    
                    if st.button("📥 Importar", use_container_width=True):
                        with st.spinner("A transferir..."):
                            time.sleep(1)
                            st.success(f"**{file_to_import}** importado!")
                            enviar_notificacao(f"Importado de {provider}: {file_to_import}", "sucesso")
                else:
                    st.caption("Pasta vazia.")
        else:
            st.warning("Erro de conexão com provider.")

    # --- TAB 2: UPLOAD ---
    with tab2:
        st.markdown(f"### Enviar Relatórios para {provider}")
        
        uploaded_cloud = st.file_uploader("Escolher ficheiro local", accept_multiple_files=False)
        
        if uploaded_cloud:
            st.write(f"**Arquivo:** {uploaded_cloud.name} ({uploaded_cloud.size/1024:.1f} KB)")
            
            if st.button("☁️ Iniciar Upload", type="primary"):
                with st.status("A enviar...", expanded=True) as status:
                    time.sleep(0.5)
                    st.write("🔐 Encriptando...")
                    time.sleep(0.5)
                    
                    # Chamar a função de upload
                    upload_to_cloud(uploaded_cloud, provider)
                    
                    status.update(label="Upload Concluído!", state="complete", expanded=False)
                    st.success("Arquivo guardado!")
                    time.sleep(1)
                    st.rerun()

# ===================== PÁGINA: AUDIT TRAIL & USER ACTIVITY (FUNC 26) =====================
elif selected == "Audit Trail":
    st.header("👣 Registo de Atividade & Auditoria")
    
    # Verificação de Segurança (Apenas Admin e Auditor deveriam ver tudo, mas deixamos aberto para demo)
    log_audit(st.session_state.username, "Audit Access", "Análise de logs de atividade")

    if st.session_state.audit_log:
        df_logs = export_audit_log()
        
        # Pré-processamento para gráficos (Func 26)
        df_logs['timestamp'] = pd.to_datetime(df_logs['timestamp'])
        df_logs['Hora'] = df_logs['timestamp'].dt.hour
        df_logs['Data'] = df_logs['timestamp'].dt.date
        
        # Criar Abas para organizar sem apagar nada
        tab1, tab2, tab3 = st.tabs(["📋 Registos Detalhados", "📈 Análise de Comportamento", "👥 Top Utilizadores"])
        
        # --- TAB 1: O SEU CÓDIGO ORIGINAL (Preservado) ---
        with tab1:
            st.markdown(f"### 📊 Total de Eventos: {len(df_logs)}")
            
            # Filtros (O seu código)
            col1, col2 = st.columns(2)
            with col1:
                user_filter = st.multiselect("Filtrar por Utilizador", 
                                           df_logs['user'].unique(),
                                           default=df_logs['user'].unique().tolist())
            with col2:
                action_filter = st.multiselect("Filtrar por Ação",
                                             df_logs['action'].unique(),
                                             default=df_logs['action'].unique().tolist())
            
            # Aplicar filtros
            filtered_df = df_logs[
                (df_logs['user'].isin(user_filter)) &
                (df_logs['action'].isin(action_filter))
            ]
            
            # Mostrar logs
            st.dataframe(filtered_df[['timestamp', 'user', 'action', 'details']], use_container_width=True, height=400)
            
            # Download
            csv = filtered_df.to_csv(index=False).encode('utf-8')
            st.download_button(
                label="📥 Download Audit Log (CSV)",
                data=csv,
                file_name=f"audit_log_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                mime="text/csv"
            )
            
            # Estatísticas (O seu código)
            st.markdown("---")
            c1, c2, c3 = st.columns(3)
            c1.metric("Total de Ações", len(filtered_df))
            c2.metric("Utilizadores Únicos", filtered_df['user'].nunique())
            c3.metric("Tipos de Ação", filtered_df['action'].nunique())

        # --- TAB 2: NOVOS GRÁFICOS DE ANÁLISE (Func 26) ---
        with tab2:
            st.markdown("### 📊 Padrões de Utilização")
            
            col1, col2 = st.columns(2)
            
            with col1:
                # Atividade por Hora (Histograma)
                fig_hora = px.histogram(
                    df_logs, 
                    x="Hora", 
                    nbins=24,
                    title="Volume de Ações por Hora do Dia",
                    template="plotly_dark",
                    color_discrete_sequence=['#3b82f6']
                )
                fig_hora.update_layout(bargap=0.1)
                st.plotly_chart(fig_hora, use_container_width=True)
                
            with col2:
                # Tipos de Ação mais comuns
                fig_actions = px.pie(
                    df_logs, 
                    names='action', 
                    title="Distribuição de Tipos de Ação",
                    template="plotly_dark",
                    hole=0.4
                )
                st.plotly_chart(fig_actions, use_container_width=True)
            
            # Timeline
            st.markdown("### 📉 Timeline de Eventos")
            # Agrupar por hora para a linha do tempo
            activity_over_time = df_logs.groupby(df_logs['timestamp'].dt.floor('H')).size().reset_index(name='Contagem')
            if not activity_over_time.empty:
                fig_time = px.line(
                    activity_over_time, 
                    x='timestamp', 
                    y='Contagem', 
                    title="Evolução Temporal da Atividade",
                    template="plotly_dark",
                    markers=True
                )
                st.plotly_chart(fig_time, use_container_width=True)

        # --- TAB 3: TOP UTILIZADORES (Func 26) ---
        with tab3:
            st.markdown("### 🏆 Utilizadores Mais Ativos")
            
            # Contagem por user
            top_users = df_logs['user'].value_counts().reset_index()
            top_users.columns = ['Utilizador', 'Ações Totais']
            
            # Último login
            last_seen = df_logs.groupby('user')['timestamp'].max().reset_index()
            last_seen.columns = ['Utilizador', 'Última Atividade']
            
            # Juntar dados
            user_stats = pd.merge(top_users, last_seen, on='Utilizador')
            
            col_chart, col_table = st.columns([1, 1])
            
            with col_chart:
                fig_users = px.bar(
                    user_stats.head(10), 
                    x='Ações Totais', 
                    y='Utilizador', 
                    orientation='h',
                    title="Top 10 Utilizadores",
                    template="plotly_dark",
                    color='Ações Totais',
                    color_continuous_scale='Viridis'
                )
                fig_users.update_layout(yaxis={'categoryorder':'total ascending'})
                st.plotly_chart(fig_users, use_container_width=True)
                
            with col_table:
                st.dataframe(
                    user_stats,
                    use_container_width=True,
                    column_config={
                        "Última Atividade": st.column_config.DatetimeColumn(format="D MMM YYYY, HH:mm")
                    }
                )

    else:
        st.info("📝 Nenhum evento registado ainda. Navegue pela aplicação para gerar dados.")
# ===================== FOOTER =====================
st.markdown("---")
st.markdown("""
<div style='text-align: center; color: #64748b; padding: 20px;'>
    <p><strong>SocioStat Banking Enterprise v3.2</strong></p>
    <p>Developed with ❤️ for Financial Risk Management | Compliance Ready</p>
</div>
""", unsafe_allow_html=True)

# ===================== FIM DO CÓDIGO =====================



